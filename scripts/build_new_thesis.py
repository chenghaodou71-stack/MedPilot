"""Build a fresh MedPilot undergraduate thesis from project facts.

This intentionally does not import or append the previously generated thesis.
It creates a new Word document, using the existing Visio PNG previews as
figures and explicit placeholders for screenshots that require the user's
local runtime.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "MedPilot_毕业设计论文_全新正文版.docx"
FIG_DIR = ROOT / "outputs" / "thesis-images" / "visio-previews"

TITLE = "基于多智能体协同与大语言模型的多专科医疗健康咨询及辅助分诊系统"
EN_TITLE = "A Multi-specialty Medical Health Consultation and Assisted Triage System Based on Multi-agent Collaboration and Large Language Models"

NAVY = "1F4E79"
LIGHT_BLUE = "DDEBF7"
LIGHT_GREEN = "E2F0D9"
LIGHT_YELLOW = "FFF2CC"
LIGHT_RED = "F4CCCC"
LIGHT_GRAY = "F2F2F2"
TEXT = "222222"


def set_font(run, name="宋体", size=12, bold=False, color=TEXT, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table, color="B7C9D6", size="6"):
    tbl_pr = table._tbl.tblPr
    border = tbl_pr.first_child_found_in("w:tblBorders")
    if border is None:
        border = OxmlElement("w:tblBorders")
        tbl_pr.append(border)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = border.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            border.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_widths(table, widths_mm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_mm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 56.6929)))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_mm):
            cell.width = Mm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 56.6929)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, "Times New Roman", 10, color="666666")


def configure(doc):
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(25)
    sec.bottom_margin = Mm(22)
    sec.left_margin = Mm(30)
    sec.right_margin = Mm(25)
    sec.header_distance = Mm(12)
    sec.footer_distance = Mm(12)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in (("Heading 1", 16, NAVY), ("Heading 2", 14, TEXT), ("Heading 3", 12, TEXT)):
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(6)

    if "CaptionCN" not in styles:
        cap = styles.add_style("CaptionCN", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["CaptionCN"]
    cap.font.name = "宋体"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap.font.size = Pt(10.5)
    cap.font.color.rgb = RGBColor.from_string(NAVY)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(10)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)


def paragraph(doc, text, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, size=12, bold=False, color=TEXT, font="宋体", after=6):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(after)
    pf.first_line_indent = Pt(24 if indent and align == WD_ALIGN_PARAGRAPH.JUSTIFY else 0)
    r = p.add_run(text)
    set_font(r, font, size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_font(r, "黑体", 16 if level == 1 else 14 if level == 2 else 12, bold=True, color=NAVY if level == 1 else TEXT)
    return p


def title_page(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本科毕业设计（论文）")
    set_font(r, "黑体", 20, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(35)
    r = p.add_run(TITLE)
    set_font(r, "黑体", 22, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("设计与实现")
    set_font(r, "黑体", 18, bold=True, color=NAVY)
    doc.add_paragraph()
    table = doc.add_table(rows=6, cols=2)
    set_widths(table, [38, 100])
    rows = [("学生姓名", "待填写"), ("学号", "待填写"), ("学院", "待填写"), ("专业", "待填写"), ("指导教师", "待填写"), ("完成日期", "待填写")]
    for row, (label, value) in zip(table.rows, rows):
        shade(row.cells[0], LIGHT_BLUE)
        for cell, text in zip(row.cells, (label, value)):
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = p.add_run(text)
            set_font(rr, "宋体", 12, bold=(cell is row.cells[0]))
    borders(table)
    doc.add_page_break()


def abstract_pages(doc):
    heading(doc, "摘要", 1)
    paragraph(doc, "针对多专科健康咨询中用户描述不规范、危险信号容易遗漏以及大语言模型回答缺少证据链等问题，本文设计并实现了 MedPilot 多专科医疗健康咨询及辅助分诊系统。系统定位为证据增强、规则护栏下的辅助工具，不输出确定性疾病诊断，不替代医生的临床判断。")
    paragraph(doc, "系统采用 Vue 3、Spring Boot 和 FastAPI 的分层架构，AI 服务以 LangGraph 有向工作流组织安全筛查、症状抽取、信息充分性判断、主动追问、医学知识检索、辅助分诊和回答编排等节点。危险信号筛查在大语言模型和向量检索调用之前执行，胸痛、呼吸困难、咯血、便血、晕厥、意识不清和大出血等信号进入高风险快速通道。普通场景使用 bge-m3 生成查询向量，在 FAISS 索引中检索证据，并通过科室加权投票形成辅助分诊结果；证据不足时返回 abstain，并提示补充信息或线下分诊。")
    paragraph(doc, "系统进一步实现了 NDJSON 流式事件、问诊记录、健康档案授权上下文、医生复核队列、知识库审核和索引版本治理、Trace 监控、审计日志以及基于医疗关系、院区和 MFA 的复核权限控制。实验采用人工构造或脱敏病例，软件测试结果只用于验证工程行为，不外推为临床准确率。")
    paragraph(doc, "关键词：多智能体协同；大语言模型；辅助分诊；检索增强生成；医学知识库；安全护栏", indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
    doc.add_page_break()
    heading(doc, "Abstract", 1)
    paragraph(doc, "This thesis presents MedPilot, a multi-specialty medical health consultation and assisted triage system based on multi-agent collaboration and large language models. The system is designed as an evidence-grounded and rule-guarded assistant rather than a definitive diagnostic tool. A layered architecture consisting of Vue 3, Spring Boot, and FastAPI is implemented. A fixed LangGraph workflow coordinates safety screening, symptom extraction, information sufficiency checking, follow-up questioning, medical retrieval, assisted triage, and answer composition. Deterministic red-flag screening is executed before any language-model or embedding call. Ordinary cases use bge-m3 embeddings and a FAISS index, followed by evidence-weighted department voting. Insufficient evidence triggers an abstention result and an offline-care suggestion.")
    paragraph(doc, "The system also provides NDJSON streaming events, consultation records, consent-scoped health context, clinical review queues, knowledge-document governance, versioned index activation, trace monitoring, audit logging, and MFA- and care-relationship-aware authorization. Evaluation data are artificial or de-identified. The reported results validate software behavior and reproducibility only; they do not establish clinical accuracy or replacement of clinicians.")
    paragraph(doc, "Keywords: multi-agent collaboration; large language model; assisted triage; retrieval-augmented generation; medical knowledge base; safety guardrail", indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, font="Times New Roman")
    doc.add_page_break()


def statements(doc):
    heading(doc, "原创性声明", 1)
    paragraph(doc, "本人郑重声明：本论文是在指导教师指导下独立完成的研究与设计工作。除文中已经注明引用的内容外，本论文不包含他人已经发表或撰写的研究成果。论文中的系统代码、测试数据和图表以项目实际提交版本为准。", indent=False)
    paragraph(doc, "签名：待填写　　　　　　　　　日期：待填写", indent=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_page_break()
    heading(doc, "版权使用授权书", 1)
    paragraph(doc, "本人同意按照学校有关规定，在论文管理和学术交流范围内使用本论文。论文中的病例均为人工构造或脱敏数据，不包含可识别的真实患者信息。", indent=False)
    paragraph(doc, "签名：待填写　　　　　　　　　日期：待填写", indent=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_page_break()


def toc(doc):
    heading(doc, "目录", 1)
    items = [
        "第1章 绪论", "第2章 相关技术与理论基础", "第3章 系统需求分析", "第4章 系统总体设计", "第5章 核心算法与系统实现", "第6章 系统测试与结果分析", "第7章 总结与展望", "参考文献", "致谢", "附录A 系统界面截图替换清单"
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(24 if item.startswith("第") else 36)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(item)
        set_font(r, "宋体", 12, bold=item.startswith("第"))
    doc.add_page_break()
    heading(doc, "图清单", 1)
    figures = [
        "图3-1 系统总体架构图", "图3-2 辅助分诊流程图", "图4-1 数据库关系示意图",
        "图4-2 系统功能用例图", "图4-3 角色权限模型图", "图4-4 系统部署架构图",
        "图4-5 医生复核安全闸门图", "图5-1 智能问诊时序图", "图5-2 医学 RAG 流程图",
        "图5-3 Trace 监控与审计流程图", "图6-1 至图6-3 系统界面截图（待替换）",
    ]
    for item in figures:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r, "宋体", 11)
    doc.add_page_break()
    heading(doc, "表清单", 1)
    tables = [
        "表3-1 系统角色与主要职责", "表6-1 测试环境", "表6-2 主要功能测试用例", "表A-1 待替换系统截图",
    ]
    for item in tables:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r, "宋体", 11)
    doc.add_page_break()


def figure(doc, number, caption, filename, width=155):
    path = FIG_DIR / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Mm(width))
        cp = doc.add_paragraph(style="CaptionCN")
        cp.add_run(f"图{number} {caption}")
    else:
        placeholder(doc, f"图{number} {caption}（图件待生成）")


def placeholder(doc, label):
    table = doc.add_table(rows=1, cols=1)
    set_widths(table, [155])
    cell = table.cell(0, 0)
    shade(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    p.paragraph_format.space_after = Pt(55)
    r = p.add_run(label)
    set_font(r, "宋体", 12, bold=True, color="666666")
    borders(table, color="A8A8A8")
    cp = doc.add_paragraph(style="CaptionCN")
    cp.add_run("（待用户运行系统后替换为实际截图）")


def data_table(doc, caption, headers, rows, widths):
    cp = doc.add_paragraph(style="CaptionCN")
    cp.add_run(caption)
    table = doc.add_table(rows=1, cols=len(headers))
    set_widths(table, widths)
    for cell, text in zip(table.rows[0].cells, headers):
        shade(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, "宋体", 10.5, bold=True, color="FFFFFF")
    for row_data in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, row_data):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(text))
            set_font(r, "宋体", 10.2)
    borders(table)
    doc.add_paragraph()


def chapter1(doc):
    heading(doc, "第1章 绪论", 1)
    heading(doc, "1.1 研究背景", 2)
    paragraph(doc, "在线医疗健康咨询降低了用户获取健康信息的门槛，但症状描述通常具有口语化、片段化和上下文缺失等特点。对于胸痛、呼吸困难、意识不清等危险信号，系统如果只追求回答流畅而没有安全优先级，可能造成不适当的延迟。因此，医疗咨询系统需要同时处理信息理解、证据检索、风险分层和人工接管边界。")
    paragraph(doc, "大语言模型能够完成自然语言理解和多轮对话，但其生成结果仍可能出现事实错误、证据不可追溯和提示敏感等问题。本文将大语言模型放在固定职责节点中，采用确定性安全筛查、证据检索和输出约束形成协同机制，使模型承担语言理解与表达任务，而不是直接承担最终诊断决策。")
    heading(doc, "1.2 研究意义", 2)
    paragraph(doc, "理论上，本文将多智能体协同具体化为按职责拆分的有向工作流，讨论安全节点、检索节点、分诊节点和编排节点之间的状态传递关系。工程上，本文给出可运行的前后端分层系统、可追溯的医学证据结构和医生复核安全闸门，为本科阶段研究大语言模型医疗应用提供可复现实例。")
    heading(doc, "1.3 研究内容", 2)
    paragraph(doc, "本文主要工作包括：设计多专科辅助分诊业务流程；实现危险信号优先的安全筛查算法；实现症状结构化抽取、主动追问和 RAG 检索；实现规则优先与证据加权相结合的辅助分诊；实现 NDJSON 流式通信、问诊持久化和 Trace 监控；实现医生复核、知识治理、角色权限和审计控制；基于构造病例与自动化测试验证软件工程行为。")
    heading(doc, "1.4 论文组织结构", 2)
    paragraph(doc, "全文共七章。第1章介绍背景、意义和研究范围；第2章说明相关技术；第3章分析角色、功能和非功能需求；第4章给出系统架构、数据和安全设计；第5章阐述核心算法与主要模块实现；第6章介绍测试方案和结果；第7章总结工作并讨论局限与展望。")


def chapter2(doc):
    heading(doc, "第2章 相关技术与理论基础", 1)
    heading(doc, "2.1 大语言模型", 2)
    paragraph(doc, "大语言模型通过上下文建模完成文本生成，适合承担症状描述抽取、追问生成和回答编排等语言任务。MedPilot 使用本地 Ollama 部署的 qwen2.5:7b，服务通过依赖注入的 chat_fn 调用，便于测试时替换为确定性函数。")
    heading(doc, "2.2 多智能体协同与工作流编排", 2)
    paragraph(doc, "本文所称多智能体协同并非多个自主临床主体自由协商，而是将任务拆分为安全筛查、抽取、追问、检索、分诊和编排等职责节点，并用共享状态和条件边连接。LangGraph 的 StateGraph 保存症状、证据、分诊结果和回答状态，节点开始、完成和错误事件由 EventEmitter 统一输出。")
    heading(doc, "2.3 检索增强生成", 2)
    paragraph(doc, "RAG 先从外部知识库检索与问题相关的证据片段，再将证据交给回答模块。系统使用 bge-m3 生成向量，FAISS IndexFlatIP 执行归一化内积检索，默认 top_k 为3、最低综合分数为0.35，并将向量分数与字符 n-gram 词法重合度按 0.8 和 0.2 融合。")
    heading(doc, "2.4 Spring Boot、FastAPI 与 Vue 3", 2)
    paragraph(doc, "Vue 3 负责交互界面和流式事件展示；Spring Boot 作为业务后端，承担认证、会话归属、记录持久化、事件校验和权限控制；FastAPI 负责 AI 推理工作流。浏览器只访问 Spring Boot，避免把 AI 服务直接暴露给用户端。")
    heading(doc, "2.5 医疗安全与隐私控制", 2)
    paragraph(doc, "系统采用 JWT Cookie、CSRF、RBAC、医疗关系、院区和 MFA 等控制。医疗数据的敏感字段加密保存，知识文档经过审核和有效期控制后才进入索引。系统明确不提供确定性诊断、处方或临床替代结论。")


def chapter3(doc):
    heading(doc, "第3章 系统需求分析", 1)
    heading(doc, "3.1 用户角色分析", 2)
    data_table(doc, "表3-1 系统角色与主要职责", ["角色", "主要功能", "权限边界"], [
        ("患者 USER", "智能问诊、问诊记录、健康档案", "只能访问本人数据"),
        ("医生 DOCTOR", "知识审核、临床复核、复核决定", "需要医院员工档案和 MFA"),
        ("复核员 REVIEWER", "复核队列、临床评测和决定", "禁止自复核，受医疗关系约束"),
        ("知识编辑器", "知识文档入库和索引构建", "不能单独完成临床激活"),
        ("管理员 ADMIN", "用户、配置和治理运维", "不自动获得临床复核权限"),
        ("审计员 AUDITOR", "Trace 和审计只读查看", "不参与临床决定"),
    ], [30, 70, 55])
    heading(doc, "3.2 功能需求", 2)
    paragraph(doc, "系统应支持用户登录、角色导航、症状输入、连续问诊、主动追问、高风险快速通道、证据引用、问诊记录、健康档案、医生复核、医学知识管理、智能体监控和审计日志。")
    heading(doc, "3.3 非功能需求", 2)
    paragraph(doc, "安全性要求危险信号优先、权限失败关闭、敏感字段加密和操作可审计；可靠性要求流式事件顺序严格、异常链路持久化、索引版本可回滚；可维护性要求前后端接口清晰、AI 节点可单元测试、知识库可版本治理；可解释性要求分离规则支持、检索支持和暂缓原因。")
    heading(doc, "3.4 业务流程", 2)
    figure(doc, "3-1", "系统总体架构图", "01_系统总体架构图.png")
    figure(doc, "3-2", "辅助分诊流程图", "02_辅助分诊流程图.png")
    heading(doc, "3.5 需求约束", 2)
    paragraph(doc, "本文只讨论软件原型和工程验证，不将构造病例结果解释为临床准确率，不使用真实患者身份信息，不将管理员权限等同于医生复核权限，也不把 support_score 解释为校准后的临床概率。")


def chapter4(doc):
    heading(doc, "第4章 系统总体设计", 1)
    heading(doc, "4.1 总体架构设计", 2)
    paragraph(doc, "系统采用浏览器、业务后端、AI 服务和数据存储四层结构。前端通过 REST 和流式接口访问 Spring Boot；Spring Boot 负责会话和安全边界，并将经过校验的请求转发给 FastAPI；FastAPI 调用 LangGraph 工作流、Ollama 和 FAISS；MySQL 保存问诊、复核、权限和审计数据。")
    heading(doc, "4.2 数据库设计", 2)
    paragraph(doc, "数据库围绕用户、问诊会话、问诊记录、消息、健康档案、临床复核、Trace、审计日志和医疗关系建立实体关系。问诊结果和 Trace 分开保存，使失败链路也能被追溯，同时避免把不完整结果展示为成功问诊。")
    figure(doc, "4-1", "数据库关系示意图", "03_数据库关系示意图.png")
    heading(doc, "4.3 功能用例与权限设计", 2)
    figure(doc, "4-2", "系统功能用例图", "04_系统功能用例图.png")
    figure(doc, "4-3", "角色权限模型图", "08_角色权限模型图.png")
    heading(doc, "4.4 部署设计", 2)
    figure(doc, "4-4", "系统部署架构图", "09_系统部署架构图.png")
    paragraph(doc, "本地部署使用前端 Vite、Spring Boot 8080、FastAPI 8000、MySQL 8、Ollama 和 FAISS 索引存储。生产预检脚本检查环境变量、服务连通性、索引清单和备份状态，但模型和索引就绪仍需要受保护的 AI readiness 请求验证。")
    heading(doc, "4.5 安全闸门设计", 2)
    figure(doc, "4-5", "医生复核安全闸门图", "07_医生复核安全闸门图.png")
    paragraph(doc, "临床复核服务只允许 DOCTOR 和 REVIEWER 操作，并要求医院员工档案、员工编号和 MFA assurance level 不低于2。复核队列列表同样执行医疗关系和院区过滤，领取和决定操作拒绝本人复核，原始 AI 结果保持不可变并记录复核理由。")


def chapter5(doc):
    heading(doc, "第5章 核心算法与系统实现", 1)
    heading(doc, "5.1 多智能体协同工作流", 2)
    paragraph(doc, "系统将一次咨询建模为状态图。输入状态包含原始文本、历史消息、健康上下文和可选的向量依赖；节点输出写回 symptoms、followup、evidence、triage 和 answer 字段。安全筛查命中时绕过模型和检索，直接进入 classify 与 compose；未命中时执行症状抽取，并依据 is_sufficient 判断进入主动追问或检索分支。")
    figure(doc, "5-1", "智能问诊时序图", "05_智能问诊时序图.png")
    heading(doc, "5.2 危险信号优先算法", 2)
    paragraph(doc, "危险信号算法是确定性规则函数。系统首先调用 match_danger_signs 对胸痛、呼吸困难、气促、咯血、便血、晕厥、意识不清和大出血等关键词及否定表达进行匹配；若命中，则构造带 red_flags 的 StructuredSymptoms，调用同一分诊函数生成高风险结果，并输出急诊或尽快就医建议。该步骤在 LLM 和 embedding 调用之前执行。")
    paragraph(doc, "其核心逻辑可表示为：若存在危险信号 r，则 T(r)=HighRisk，输出对应科室和就医时效；若不存在危险信号，则进入信息充分性与证据检索分支。规则优先保证高风险路径不会被生成式回答覆盖。")
    heading(doc, "5.3 信息充分性与主动追问", 2)
    paragraph(doc, "症状抽取结果包括 symptoms、duration、severity、associated_symptoms、medical_history、medications、red_flags 和 raw_text 等字段。当持续时间、严重程度或关键伴随症状缺失时，系统生成一个有限追问问题；追问结束后回到正常咨询流程。原始文本始终保留，健康档案只作为经过授权的背景上下文，不覆盖本轮症状。")
    heading(doc, "5.4 医学 RAG 检索算法", 2)
    figure(doc, "5-2", "医学 RAG 流程图", "06_医学RAG流程图.png")
    paragraph(doc, "检索查询由原始文本和结构化症状组合而成，并进行同义词扩展。FAISS 返回较宽候选集后，系统计算综合分数：S=0.8S_dense+0.2S_lexical，其中 S_dense 为归一化向量内积，S_lexical 为字符二元组重合度。低于阈值的候选被过滤，最终返回包含 citation_id、doc_id、chunk_id、source、quote、score、index_version 和审核元数据的 RankedEvidence。")
    heading(doc, "5.5 证据加权辅助分诊", 2)
    paragraph(doc, "分诊函数先检查高风险规则；普通场景只保留呼吸内科、消化内科、心血管内科和皮肤科等支持科室的证据。对每个科室累加证据 score，选择权重最高者，并将最高科室权重占总权重的比例作为 support_score。如果没有足够证据，系统返回全科/线下分诊建议并设置 abstained=true。support_score 是内部支持信号，不是临床概率。")
    heading(doc, "5.6 回答编排与 NDJSON 事件", 2)
    figure(doc, "5-3", "Trace 监控与审计流程图", "10_Trace监控与审计流程图.png")
    paragraph(doc, "AI 服务使用 NDJSON 输出 safety_screen、extract、ask_followup、retrieve、classify、compose 等节点事件，事件包含 sequence、status、phase、elapsed_ms 和 data。Spring Boot 校验事件格式和终止状态，只有得到合法 done 事件时才保存成功问诊；异常情况下保存失败 Trace，避免失败请求完全消失。")
    heading(doc, "5.7 医生复核与知识治理实现", 2)
    paragraph(doc, "医生复核队列由后端 ClinicalReviewService 管理，创建、列表、详情、领取和决定操作统一执行临床复核者、员工档案、MFA、医疗关系和自复核禁止检查。知识文档必须经过审核、有效期检查和版本索引构建，失败构建不会覆盖当前激活版本。")


def chapter6(doc):
    heading(doc, "第6章 系统测试与结果分析", 1)
    heading(doc, "6.1 测试环境", 2)
    data_table(doc, "表6-1 测试环境", ["层次", "工具或版本", "测试内容"], [
        ("前端", "Vue 3、Vite、Vitest", "组件、路由、状态和构建"),
        ("后端", "Spring Boot、JUnit 5", "接口、权限、持久化和异常"),
        ("AI 服务", "FastAPI、pytest、FAISS", "规则、抽取、检索和事件"),
        ("运行环境", "MySQL 8、Ollama", "三服务联调和关键路径"),
    ], [30, 55, 70])
    heading(doc, "6.2 功能测试", 2)
    data_table(doc, "表6-2 主要功能测试用例", ["编号", "场景", "预期结果"], [
        ("TC-01", "普通症状咨询", "返回结构化分诊、证据和安全边界"),
        ("TC-02", "信息不足", "生成主动追问，不直接给出分诊结论"),
        ("TC-03", "命中危险信号", "进入高风险快速通道，建议尽快线下就医"),
        ("TC-04", "证据不足", "返回 abstain 和线下分诊建议"),
        ("TC-05", "医生复核越权", "缺少角色、医疗关系或 MFA 时拒绝"),
        ("TC-06", "Trace 异常", "保存失败链路，不生成成功记录"),
    ], [22, 58, 75])
    heading(doc, "6.3 测试结果", 2)
    paragraph(doc, "项目当前自动化验收记录包括 AI 服务 165 项测试、Spring Boot 104 项测试和前端 67 项测试；这些数量用于说明软件门禁覆盖，不代表临床样本量。十例确定性构造病例基线报告安全召回率 1.0、红旗漏检 0、科室 Macro-F1 1.0、风险 Macro-F1 1.0 和错误率 0。由于病例没有 gold_evidence_ids，Recall@K、MRR 和引用可追溯性被标记为 not-evaluated，不能将占位零值解释为检索性能。")
    heading(doc, "6.4 典型场景验证", 2)
    figure(doc, "6-1", "高风险快速通道流程（系统截图待替换）", "05_高风险快速通道.png")
    paragraph(doc, "上图位置需要替换为用户实际运行系统后的高风险页面截图。截图应隐藏真实姓名、手机号、病历号和密钥。")
    placeholder(doc, "图6-2 普通问诊结果页面截图（待用户提供）")
    placeholder(doc, "图6-3 医生复核队列与复核详情截图（待用户提供）")
    heading(doc, "6.5 局限性分析", 2)
    paragraph(doc, "当前系统没有临床专业人员盲评、独立检索金标准、前瞻性临床验证和大规模真实世界数据。监控查询在大数据量下仍需要进一步下推分页和索引优化。管理员角色默认不具备医生复核权限，医生复核页面是否可见还依赖员工档案、MFA 和医疗关系条件。以上限制决定了本文只能给出工程可行性和软件回归结论。")


def chapter7(doc):
    heading(doc, "第7章 总结与展望", 1)
    heading(doc, "7.1 工作总结", 2)
    paragraph(doc, "本文围绕多专科医疗健康咨询和辅助分诊场景，完成了从需求分析、架构设计、数据建模到算法实现和工程测试的完整原型。系统用固定职责节点实现多智能体协同，用危险信号规则和 abstain 机制建立安全护栏，用医学 RAG 和证据字段增强回答可追溯性，并通过医生复核、知识治理、权限和审计形成闭环。")
    heading(doc, "7.2 研究结论边界", 2)
    paragraph(doc, "本文验证的是系统能否按预定规则处理构造病例、输出合法事件、保存问诊和失败 Trace、执行角色权限和完成知识版本治理。测试结果不能推出临床准确率、患者获益、医生替代能力或真实人群安全性。")
    heading(doc, "7.3 后续展望", 2)
    paragraph(doc, "后续工作可从四方面展开：建立经过临床专家标注的检索金标准；开展医生盲评和分层安全性评估；引入更细粒度的机构和科室授权；在隐私合规前提下研究多模态附件的人工确认流程和可控分析能力。")


def references(doc):
    heading(doc, "参考文献", 1)
    refs = [
        "[1] Semigran H L, Linder J A, Gidengil C, et al. Evaluation of symptom checkers for self diagnosis and triage: audit study[J]. BMJ, 2015, 351: h3480.",
        "[2] Chambers D, Cantrell A J, Johnson M, et al. Digital and online symptom checkers and health assessment/triage services for urgent health problems: systematic review[J]. BMJ Open, 2019, 9(8): e027743.",
        "[3] Wallace W, Chan C, Chidambaram S, et al. The diagnostic and triage accuracy of digital and online symptom checker tools: a systematic review[J]. npj Digital Medicine, 2022, 5: 118.",
        "[4] Fraser H, Crossland D, Bacher I, et al. Comparison of diagnostic and triage accuracy of Ada Health and WebMD symptom checkers, ChatGPT, and physicians[J]. JMIR mHealth and uHealth, 2023, 11: e49995.",
        "[5] Singhal K, Azizi S, Tu T, et al. Large language models encode clinical knowledge[J]. Nature, 2023, 620: 172-180.",
        "[6] Hager P, Jungmann F, Holland R, et al. Evaluation and mitigation of the limitations of large language models in clinical decision-making[J]. Nature Medicine, 2024, 30: 2613-2622.",
        "[7] Lewis P, Perez E, Piktus A, et al. Retrieval-augmented generation for knowledge-intensive NLP tasks[C]. NeurIPS, 2020: 9459-9474.",
        "[8] Zakka C, Shad R, Chaurasia A, et al. Almanac-Retrieval-Augmented Language Models for Clinical Medicine[J]. NEJM AI, 2024, 1(2).",
        "[9] Xiong G, Jin Q, Lu Z, et al. Benchmarking Retrieval-Augmented Generation for Medicine[C]. Findings of ACL, 2024: 6233-6251.",
        "[10] Yao S, Zhao J, Yu D, et al. ReAct: Synergizing reasoning and acting in language models[C]. ICLR, 2023.",
        "[11] Chen J, Xiao S, Zhang P, et al. M3-Embedding: Multi-Linguality, Multi-Functionality, Multi-Granularity Text Embeddings Through Self-Knowledge Distillation[C]. Findings of ACL, 2024: 2318-2335.",
        "[12] Qwen Team. Qwen2.5 Technical Report[EB/OL]. arXiv:2412.15115, 2024.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(20)
        p.paragraph_format.first_line_indent = Pt(-20)
        p.paragraph_format.line_spacing = Pt(18)
        r = p.add_run(ref)
        set_font(r, "Times New Roman", 10.5)


def appendix(doc):
    heading(doc, "致谢", 1)
    paragraph(doc, "感谢指导教师在课题设计、系统实现和论文结构方面提供的指导，感谢项目测试和复核过程中给予建议的同学。本文所用病例均为人工构造或脱敏数据，MedPilot 仅作为证据增强、规则护栏下的辅助分诊原型。")
    doc.add_page_break()
    heading(doc, "附录A 系统界面截图替换清单", 1)
    paragraph(doc, "以下位置需要用户运行系统后自行截图替换。建议统一使用 PNG，隐藏真实身份信息。")
    data_table(doc, "表A-1 待替换系统截图", ["序号", "截图内容", "建议放置章节"], [
        (1, "登录页面和 MFA 提示", "第3章/第6章"),
        (2, "患者智能问诊输入页面", "第5章"),
        (3, "正常问诊结果和证据引用", "第6章"),
        (4, "信息不足主动追问页面", "第5章"),
        (5, "高风险危险信号快速通道", "第6章"),
        (6, "问诊记录和健康档案", "第5章"),
        (7, "医生复核队列和复核详情", "第4章/第6章"),
        (8, "医学知识库管理和索引版本", "第4章/第5章"),
        (9, "智能体运行监控和 Trace", "第5章"),
        (10, "用户权限和审计日志", "第4章"),
    ], [18, 100, 37])
    heading(doc, "附录B 术语说明", 1)
    paragraph(doc, "辅助分诊：根据症状描述和知识证据给出建议就诊科室、风险等级和就医时效，不给出确定性疾病诊断。abstain：证据不足或无法安全判断时主动暂缓自动推荐。Trace：记录节点状态、事件顺序、耗时、引用和终止状态的可追溯运行轨迹。")


def build():
    doc = Document()
    configure(doc)
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "MedPilot 全新毕业设计论文正文"
    title_page(doc)
    statements(doc)
    abstract_pages(doc)
    toc(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    chapter6(doc)
    chapter7(doc)
    references(doc)
    appendix(doc)
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
