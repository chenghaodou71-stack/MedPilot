from pathlib import Path
from docx import Document

p = Path(r"C:/Users/Dou-ChengHao/Desktop/7230264110_窦成皓_23软件1_基于多智能体协同与大语言模型的多专科医疗健康咨询及辅助分诊系统.docx")
d = Document(str(p))
out = Path(r"D:/毕设制作/.scratch/paper-polish-source.txt")
with out.open("w", encoding="utf-8") as f:
    for i, para in enumerate(d.paragraphs):
        text = " ".join(para.text.split())
        if text:
            f.write(f"P{i:03d}\t[{para.style.name}]\t{text}\n")
    f.write("\nTABLES\n")
    for ti, table in enumerate(d.tables):
        f.write(f"TABLE {ti}\n")
        for row in table.rows:
            f.write(" | ".join(" ".join(cell.text.split()) for cell in row.cells) + "\n")
print(out)
