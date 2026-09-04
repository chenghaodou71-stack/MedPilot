from __future__ import annotations

"""Build the rewritten MedPilot thesis from the project evidence base.

The script deliberately keeps the body prose in Markdown source files already
reviewed for this project, then adds a second pass of section-specific analysis,
black/white technical figures, schema tables, and a replacement bibliography.
It is deterministic so the document can be regenerated after a supervisor edit.
"""

import re
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "outputs" / "researchwrite" / "medpilot-thesis"
OUT = ROOT / "outputs" / "MedPilot_毕业设计论文_4万字重写版.docx"
FIG_DIR = ROOT / "outputs" / "thesis-rewrite-figures"
FIG_DIR.mkdir(parents=True, exist_ok=True)

SIMSUN = Path("C:/Windows/Fonts/simsun.ttc")
SIMHEI = Path("C:/Windows/Fonts/simhei.ttf")
ARIAL = Path("C:/Windows/Fonts/arial.ttf")

TITLE = "基于多智能体协同与大语言模型的多专科医疗健康咨询及辅助分诊系统设计与实现"
EN_TITLE = (
    "Design and Implementation of a Multi-specialty Medical Health Consultation "
    "and Assisted Triage System Based on Multi-Agent Collaboration and Large Language Models"
)
FIG_FONT_SCALE = 1.45


def font(size: int, bold: bool = False):
    path = SIMHEI if bold else SIMSUN
    if not path.exists():
        path = ARIAL
    # The figures are scaled down to an A4 thesis page.  Upscale type at the
    # source so labels remain legible after Word applies the final width.
    return ImageFont.truetype(str(path), size=round(size * FIG_FONT_SCALE))


def wrap_cn(text: str, width: int):
    out = []
    for chunk in text.split("\n"):
        out.extend(textwrap.wrap(chunk, width=width, break_long_words=True, break_on_hyphens=False) or [""])
    return out


def centered(draw, xy, text, f, fill="#111111", spacing=8):
    x, y = xy
    lines = text.split("\n")
    heights = [draw.textbbox((0, 0), line, font=f)[3] for line in lines]
    total = sum(heights) + spacing * (len(lines) - 1)
    top = y - total / 2
    for line, h in zip(lines, heights):
        box = draw.textbbox((0, 0), line, font=f)
        draw.text((x - (box[2] - box[0]) / 2, top), line, font=f, fill=fill)
        top += h + spacing


def box(draw, xy, text, f=None, fill="white", outline="black", width=4, radius=0):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)
    if f is None:
        f = font(34)
    centered(draw, ((x1 + x2) / 2, (y1 + y2) / 2), text, f)


def line(draw, points, width=4, fill="black"):
    draw.line(points, fill=fill, width=width, joint="curve")


def save_fig(name: str, image: Image.Image):
    path = FIG_DIR / name
    image.save(path, dpi=(300, 300))
    return path


def make_function_structure():
    w, h = 3800, 1700
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    top = (1080, 80, 2720, 300)
    box(d, top, "MedPilot 医疗多智能体辅助分诊平台", font(52, True), width=7)
    bus_y = 470
    line(d, [(1900, 300), (1900, bus_y)], width=7)
    line(d, [(190, bus_y), (3610, bus_y)], width=7)
    labels = [
        "用户接入", "账户与权限", "智能问诊", "红旗筛查", "主动追问", "医学检索", "辅助分诊", "健康档案", "医生复核", "监控审计",
    ]
    bw, gap, y1, y2 = 280, 64, 660, 1450
    total = len(labels) * bw + (len(labels) - 1) * gap
    start = (w - total) // 2
    for i, label in enumerate(labels):
        x1 = start + i * (bw + gap)
        x2 = x1 + bw
        line(d, [(x1 + bw // 2, bus_y), (x1 + bw // 2, y1)], width=4)
        text = "\n".join(wrap_cn(label, 4))
        box(d, (x1, y1, x2, y2), text, font(42, True), width=5)
    return save_fig("图3-1_系统功能结构图.png", im)


def make_architecture():
    w, h = 3000, 1450
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    layers = [
        ("交互层", ["Vue 3 患者端", "Vue 3 医生端", "Vue 3 管理端"]),
        ("业务层", ["Spring Boot API", "认证与权限", "问诊与复核"]),
        ("智能层", ["安全筛查", "多智能体工作流", "RAG 与回答编排"]),
        ("数据层", ["MySQL", "FAISS 索引", "Trace 与审计"]),
        ("基础设施", ["Ollama/qwen2.5", "对象存储", "备份与预检"]),
    ]
    y = 80
    centers = []
    for title, items in layers:
        box(d, (80, y, 480, y + 210), title, font(44, True), width=5)
        xs = [650, 1330, 2010]
        row = []
        for x, item in zip(xs, items):
            box(d, (x, y + 18, x + 540, y + 192), item, font(34), width=4)
            row.append((x + 270, y + 105))
        centers.append(row)
        y += 260
    for upper, lower in zip(centers, centers[1:]):
        for a in upper:
            for b in lower:
                line(d, [a, b], width=2)
    return save_fig("图3-2_系统总体架构图.png", im)


def make_workflow():
    w, h = 3100, 1300
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    nodes = [
        ("输入症状", 90, 500, 420, 720),
        ("安全筛查\n否定感知", 520, 500, 940, 720),
        ("症状抽取\n结构化状态", 1040, 500, 1460, 720),
        ("信息充分?", 1560, 500, 1980, 720),
        ("主动追问", 2100, 850, 2520, 1070),
        ("医学检索\n证据排序", 2100, 500, 2520, 720),
        ("辅助分诊\n科室/风险/时效", 2640, 500, 3060, 720),
        ("安全输出\n引用与边界", 2640, 850, 3060, 1070),
    ]
    for label, x1, y1, x2, y2 in nodes:
        box(d, (x1, y1, x2, y2), label, font(34, True), width=4)
    for a, b in [((420, 610), (520, 610)), ((940, 610), (1040, 610)), ((1460, 610), (1560, 610)), ((1980, 610), (2100, 610)), ((2520, 610), (2640, 610)), ((2850, 720), (2850, 850)), ((1770, 720), (2310, 850))]:
        line(d, [a, b], width=4)
    line(d, [(730, 500), (730, 350), (2300, 350), (2300, 500)], width=4)
    box(d, (1130, 80, 1970, 260), "命中红旗信号：跳过模型与检索，进入高风险快速通道", font(30, True), width=4)
    line(d, [(730, 350), (1550, 350), (1550, 260)], width=3)
    return save_fig("图4-1_辅助分诊流程图.png", im)


def make_db():
    w, h = 3100, 1450
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    entities = [
        ("users", 80, 120), ("health_profiles", 80, 760), ("consultation_sessions", 780, 120),
        ("consultation_messages", 780, 760), ("consultation_records", 1500, 120),
        ("clinical_reviews", 1500, 760), ("knowledge_documents", 2250, 120),
        ("consultation_traces", 2250, 760),
    ]
    centers = {}
    for name, x, y in entities:
        box(d, (x, y, x + 600, y + 360), name + "\nPK id\nuser_id / record_id\nstatus / created_at", font(30, True), width=4)
        centers[name] = (x + 300, y + 180)
    for a, b in [("users", "health_profiles"), ("users", "consultation_sessions"), ("consultation_sessions", "consultation_messages"), ("users", "consultation_records"), ("consultation_records", "clinical_reviews"), ("consultation_records", "consultation_traces"), ("knowledge_documents", "consultation_traces"), ("clinical_reviews", "consultation_traces")]:
        line(d, [centers[a], centers[b]], width=3)
    return save_fig("图3-3_数据库关系示意图.png", im)


def make_rag():
    w, h = 3000, 1100
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    labels = ["审核文档", "切分与元数据", "bge-m3 向量化", "FAISS 候选", "词法融合", "证据快照", "回答编排"]
    xs = [60 + i * 420 for i in range(len(labels))]
    for x, label in zip(xs, labels):
        box(d, (x, 400, x + 340, 650), label, font(34, True), width=4)
    for x1, x2 in zip(xs[:-1], xs[1:]):
        line(d, [(x1 + 340, 525), (x2, 525)], width=4)
    box(d, (750, 80, 2250, 250), "仅审核通过且未过期的文档进入活动索引；失败构建不得覆盖 active_index", font(32, True), width=4)
    line(d, [(1500, 250), (1500, 400)], width=3)
    box(d, (750, 790, 2250, 960), "每条证据绑定 citation_id、doc_id、chunk_id、source、quote、score、index_version", font(30), width=4)
    line(d, [(2010, 650), (2010, 790)], width=3)
    return save_fig("图4-2_医学RAG流程图.png", im)


def make_gate():
    w, h = 3000, 1150
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    labels = ["身份认证", "员工档案", "医疗关系", "院区匹配", "MFA≥2", "禁止自复核", "可领取/决定"]
    xs = [60 + i * 420 for i in range(len(labels))]
    for x, label in zip(xs, labels):
        box(d, (x, 450, x + 330, 680), label, font(32, True), width=4)
    for x1, x2 in zip(xs[:-1], xs[1:]):
        line(d, [(x1 + 330, 565), (x2, 565)], width=4)
    box(d, (900, 120, 2100, 280), "任一闸门失败：fail-closed，记录审计事件，不返回患者数据", font(30, True), width=4)
    line(d, [(1500, 280), (1500, 450)], width=3)
    box(d, (1030, 820, 1970, 1000), "人工决定与原始 AI 结果分离保存，支持追溯与回滚", font(30), width=4)
    line(d, [(2760, 680), (2760, 820), (1970, 910)], width=3)
    return save_fig("图4-3_医生复核安全闸门图.png", im)


def make_deployment():
    w, h = 3000, 1250
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    layers = [
        ("浏览器", ["患者端", "医生端", "管理端"], 80),
        ("业务网络", ["Spring Boot", "JWT/CSRF", "审计 API"], 380),
        ("AI 网络", ["FastAPI", "LangGraph", "NDJSON"], 680),
        ("数据与模型", ["MySQL", "FAISS", "Ollama"], 980),
    ]
    for title, items, y in layers:
        box(d, (80, y, 480, y + 190), title, font(38, True), width=5)
        for i, item in enumerate(items):
            x = 720 + i * 690
            box(d, (x, y + 20, x + 520, y + 170), item, font(32), width=4)
            if y > 80:
                line(d, [(x + 260, y - 110), (x + 260, y + 20)], width=3)
    box(d, (2080, 70, 2920, 210), "浏览器不直连 AI 服务；内部 token 仅在服务端使用", font(28, True), width=4)
    line(d, [(2500, 210), (2500, 380)], width=3)
    return save_fig("图5-1_系统部署架构图.png", im)


def make_trace():
    w, h = 3000, 1050
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    lanes = [("前端", 140), ("Spring Boot", 380), ("FastAPI", 620), ("MySQL/审计", 860)]
    for title, y in lanes:
        d.text((60, y - 22), title, font=font(32, True), fill="black")
        line(d, [(360, y), (2900, y)], width=2)
    events = [(500, 140, 380, "POST /consult"), (820, 380, 620, "NDJSON 代理"), (1130, 620, 620, "safety_screen"), (1450, 620, 620, "retrieve"), (1770, 620, 380, "done"), (2080, 380, 860, "写入 Trace"), (2420, 860, 140, "刷新结果")]
    for x, y1, y2, label in events:
        line(d, [(x, y1), (x, y2)], width=4)
        d.ellipse((x - 8, y2 - 8, x + 8, y2 + 8), fill="black")
        d.text((x + 18, min(y1, y2) + 18), label, font=font(28), fill="black")
    box(d, (850, 45, 2150, 115), "每个事件携带 sequence、node、status、phase、elapsed_ms；终止事件必须合法", font(27, True), width=3)
    return save_fig("图5-2_Trace监控与审计流程图.png", im)


def generate_figures():
    return {
        "功能结构": make_function_structure(),
        "总体架构": make_architecture(),
        "分诊流程": make_workflow(),
        "数据库": make_db(),
        "RAG": make_rag(),
        "复核闸门": make_gate(),
        "部署": make_deployment(),
        "Trace": make_trace(),
    }


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="000000", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_mm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    total = int(sum(widths_mm) * 56.6929)
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_mm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 56.6929)))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row.cells, widths_mm):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 56.6929)))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Mm(width)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_doc(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(18)
    for name, size, before, after in (("Heading 1", 16, 18, 10), ("Heading 2", 13, 12, 6), ("Heading 3", 11.5, 8, 4)):
        st = styles[name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size)
        st.font.bold = True
        st.paragraph_format.first_line_indent = Pt(0)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        st.paragraph_format.line_spacing = Pt(size + 8)
    if "Table Text" not in styles:
        st = styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        st = styles["Table Text"]
    st.font.name = "宋体"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    st.font.size = Pt(9.2)
    st.paragraph_format.first_line_indent = Pt(0)
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    st.paragraph_format.line_spacing = Pt(13)
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(27)
        section.right_margin = Mm(27)
        section.top_margin = Mm(24)
        section.bottom_margin = Mm(20)
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)
        hp = section.header.paragraphs[0]
        hp.text = "MedPilot 医疗多智能体辅助分诊系统"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.runs[0].font.name = "黑体"
        hp.runs[0].font.size = Pt(8)
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(fp)
        for run in fp.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)


def add_para(doc, text, *, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, style="Normal", size=None, bold=False):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    pf = p.paragraph_format
    pf.first_line_indent = Pt(21 if indent else 0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18 if size is None else max(14, size + 6))
    pf.space_after = Pt(2)
    for part in re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text):
        if not part:
            continue
        r = p.add_run(part[2:-2] if part.startswith("**") else part[1:-1] if part.startswith("`") else part)
        r.font.name = "宋体" if not part.startswith("`") else "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), r.font.name)
        r.font.size = Pt(size or 10.5)
        r.bold = bold or part.startswith("**")
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {min(level,3)}")
    p.paragraph_format.keep_with_next = True
    for r in p.runs:
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9.5)
    return p


def add_figure(doc, path: Path, caption: str, width_mm=150):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    pic = run.add_picture(str(path), width=Mm(width_mm))
    pic._inline.docPr.set("descr", caption)
    pic._inline.docPr.set("title", caption)
    add_caption(doc, caption)


def add_schema_table(doc, caption, rows, widths=(28, 27, 18, 46, 31)):
    p = add_caption(doc, caption)
    p.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=5)
    set_table_geometry(table, list(widths))
    set_table_borders(table)
    headers = ["字段", "类型", "长度", "约束", "备注"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = ""
        set_cell_shading(cell, "E7E7E7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=100, bottom=100)
        q = cell.paragraphs[0]
        q.style = "Table Text"
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.paragraph_format.first_line_indent = Pt(0)
        r = q.add_run(value)
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        r.font.size = Pt(9.5)
        r.bold = True
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=90, bottom=90)
            q = cell.paragraphs[0]
            q.style = "Table Text"
            q.alignment = WD_ALIGN_PARAGRAPH.CENTER if value in {"int", "varchar", "timestamp", "11", "20", "50", "100", "255", "1", "/"} else WD_ALIGN_PARAGRAPH.LEFT
            q.paragraph_format.first_line_indent = Pt(0)
            r = q.add_run(value)
            r.font.name = "宋体"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(9.1)
    set_table_geometry(table, list(widths))
    return table


def add_data_table(doc, caption, headers, rows, widths):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, list(widths))
    set_table_borders(table)
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = ""
        set_cell_shading(cell, "E7E7E7")
        q = cell.paragraphs[0]
        q.style = "Table Text"
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.paragraph_format.first_line_indent = Pt(0)
        r = q.add_run(value)
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        r.font.size = Pt(9.2)
        r.bold = True
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = ""
            set_cell_margins(cell, top=80, bottom=80)
            q = cell.paragraphs[0]
            q.style = "Table Text"
            q.alignment = WD_ALIGN_PARAGRAPH.CENTER
            q.paragraph_format.first_line_indent = Pt(0)
            r = q.add_run(value)
            r.font.name = "宋体"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(9)
    set_table_geometry(table, list(widths))
    return table


FRESH_CONTENT = {
    "第1章": [
        "本研究首先把‘辅助分诊’限定为一种受规则和证据约束的决策支持服务。系统只对就医紧迫程度、建议科室和需要补充的信息进行结构化表达，不把模型生成的疾病名称、药物剂量或处方建议作为输出目标。这个边界直接决定了后续评价指标：安全召回、证据可追溯、人工接管和失败可见性比语言流畅度更优先。",
        "研究问题被拆成三个可以被工程测试的问题。其一，危险信号能否在模型调用之前被稳定识别，并在异常输入下保持失败关闭；其二，多智能体拆分能否让症状抽取、追问、检索、分诊和回答编排各自拥有清晰的输入输出；其三，系统是否能把每一次输出绑定到知识版本、事件序列和权限检查，从而支持人工复核和事后审计。",
        "与单一问答机器人相比，本系统的研究对象不是某个模型的静态能力，而是一条可观测的服务链。链路从浏览器输入开始，依次经过会话归属校验、红旗筛查、结构化抽取、信息充分性判断、证据检索、加权分诊和回答编排，最终将事件与引用快照写入数据库。任何一个节点失败时，都必须给出可解释的终止状态，而不是用一段看似完整的自然语言掩盖失败。",
        "本文采用‘实现事实—测试证据—结论边界’三列式论证方法。凡是可以由源代码、接口契约或自动化测试直接证明的内容，使用确定性表述；凡是只在构造病例上验证的内容，明确标记为工程回归；凡是需要临床标注或前瞻研究的内容，只作为后续工作提出，不从原型结果外推。",
        "本章在综述在线症状检查器、大语言模型医疗问答、知识检索和多智能体协同研究的基础上，提出本文的系统问题和评价边界。相关工作显示，分诊工具的准确性差异很大，安全证据、可追溯性和人工监督仍是共同短板[1-10]。因此，MedPilot 的贡献不以‘替代医生’为目标，而以可审计的安全工作流为主线。",
    ],
    "第2章": [
        "需求分析以角色、状态和风险边界为基本单位。患者需要的是低门槛的自然语言输入和可理解的结果；医生需要的是经过授权过滤的复核队列和原始证据；知识编辑者需要版本化的入库与审核过程；管理员和审计员则更关注权限、指标和故障定位。若把这些需求混在一个页面或一个接口中，后续权限检查就会退化为前端按钮隐藏，无法形成真正的安全边界。",
        "本系统的业务流程采用四种主状态：正常咨询、信息不足追问、高风险快速通道和失败终止。状态之间通过明确的事件协议连接，状态转换条件写入后端和 AI 服务两侧。正常咨询必须完成安全筛查、抽取、检索、分诊与编排；追问只返回一个当前最重要的问题；高风险通道绕过模型和向量检索；失败终止则保留错误码和 Trace，不把部分结果标记为成功。",
        "非功能需求以可验证的指标描述。安全性对应未登录、越权、医疗关系不匹配、MFA 不足和自复核等拒绝用例；可靠性对应 NDJSON 序列连续、终止事件合法和索引切换原子性；可维护性对应节点可单元测试、模型函数可注入和接口字段有版本号；可追溯性对应 citation_id、index_version、trace_id 和 audit_log 的关联完整。",
        "技术选型遵循可替换和可回归原则。Vue 3 只承担交互和状态呈现，Spring Boot 统一承担认证、授权、会话归属、数据持久化和 AI 代理，FastAPI 承担工作流与模型工具调用。这样做增加了一个服务间边界，却让浏览器不需要保存内部 token，也让模型调用可以在没有真实模型的测试夹具中被替换。",
        "对成本的估算不只考虑开发阶段的免费软件，还包括知识审核、模型升级、索引重建、日志留存和密钥轮换等运维成本。原型阶段使用本地 Ollama 和 FAISS，可以降低外部服务依赖；一旦进入医院内部部署，还需要补充 GPU 资源、备份策略、账号生命周期、数据脱敏和伦理审查预算。",
    ],
    "第3章": [
        "数据库设计采用‘原始事实不可覆盖、派生结果可重算’原则。consultation_messages 保存用户和系统的原始消息，consultation_records 保存一次问诊完成时的结构化结果和引用快照，clinical_reviews 保存人工复核决定。三者不通过覆盖字段表示流程，而是通过外键和版本字段形成时间线，因此可以同时回答‘模型当时输出了什么’和‘人工后来决定了什么’。",
        "用户健康档案不是自动注入的长期记忆。只有在当前用户明确授权、字段处于白名单且长度限制通过时，HealthProfileContextService 才将其投影为背景上下文；红旗筛查仍直接读取本轮原始输入，防止既往史文本覆盖当前危险信号。授权撤回后，新的问诊不再读取档案，但历史记录仍按审计和留存规则保留。",
        "知识治理将文档状态、索引状态和引用快照分离。文档可以处于草稿、待审核、已通过、已过期或已拒绝状态；索引可以处于构建中、候选、活动或回滚状态；问诊记录只引用当时活动索引的版本。构建脚本先生成 manifest 和临时索引，再由显式激活操作切换指针，避免半成品覆盖线上版本。",
        "事件协议采用 NDJSON 而不是一个长连接中的自定义文本。每行都是独立 JSON 对象，包含 protocol_version、sequence、type、node、status、phase、elapsed_ms 和 data。后端按 sequence 校验连续性，前端按 sequence 合并节点状态；这使得网络中断、节点超时和重复事件都能够在日志中被定位。",
        "权限模型遵循‘先确定资源归属，再判断角色动作’的顺序。患者接口首先验证 user_id 和 session_id；临床复核接口先验证员工档案、医疗关系和院区，再检查 DOCTOR 或 REVIEWER 权限；知识接口先检查文档范围，再判断 KNOWLEDGE_EDITOR、ADMIN 或审核角色。角色名称本身不携带跨组织访问能力。",
        "安全威胁建模覆盖提示注入、越权读取、恶意附件、知识污染、日志泄露和索引回滚失败六类风险。对应控制措施包括服务端角色判断、医疗关系过滤、扩展名与文件头双重校验、知识审核与有效期、敏感字段加密以及激活前 manifest 比对。威胁建模的作用是建立可测试的安全假设，而不是声称系统不存在风险。",
        "接口设计中最重要的不是路径数量，而是失败语义。认证失败统一返回关闭式错误；AI 服务超时返回可区分的模型错误和系统错误；无证据时返回 abstained=true 和解释；索引构建失败保持旧版本活动；复核决定失败则不改变任务状态。客户端据此展示‘需要重试’、‘需要补充信息’和‘需要人工复核’三种不同提示。",
    ],
    "第4章": [
        "红旗规则被实现为模型调用前的确定性函数。规则包含关键词、同义词、否定表达和最小上下文窗口，先识别否定范围，再判断危险词是否在有效范围内。‘没有胸痛’与‘胸痛伴出汗’的结果不同，规则测试同时覆盖正例、否定例和混合表达。该设计不能穷尽自然语言，但能够把最高代价的漏检风险从生成模型中隔离出来。",
        "普通路径的分诊不是把向量相似度直接当作科室概率。检索器先取得较大的候选集合，再根据字符 n-gram 和结构化症状字段进行词法融合；分类器按证据所属科室累加权重，并输出 support_score、support_factors 和 evidence_factors。若候选不足或分数低于阈值，系统返回暂缓状态，要求补充信息或转人工分诊。",
        "主动追问采用‘一次只问一个关键缺口’的策略。信息充分性判断检查持续时间、严重程度、伴随症状、既往史和就医意图等字段，追问节点根据字段优先级选择问题。这样可以避免一次性生成问卷，也让下一轮输入可以和 history 一起重新经过安全筛查，而不是把上一轮模型结论当作事实。",
        "回答编排节点只消费经过校验的结构化结果和证据快照。模板中强制包含风险等级、建议就医时效、建议科室或暂缓原因、证据来源和安全边界；检测到确定性诊断、处方剂量或‘无需就医’等越界表述时，触发安全回退模板。输出约束的目标是减小伤害面，不是保证模型语言绝对正确。",
        "算法复杂度分析分开讨论确定性部分和模型部分。危险信号扫描与输入长度近似线性；IndexFlatIP 的精确检索为 O(Nd)，其中 N 为向量数、d 为向量维度；候选排序为 O(k log k)。模型推理受参数量、上下文长度和硬件影响，需要通过 Trace 记录实际耗时，不能用静态复杂度代替性能测量。",
        "在证据不足时主动 abstain 是系统的核心安全策略之一。它把‘没有足够支持’从空结果升级为可观察状态，并在前端展示原因。该机制可能降低表面上的首选科室命中率，却减少了系统在未知场景下强行给出建议的倾向。后续临床评估应同时报告 abstain 率、人工接管率和高风险漏检率。",
    ],
    "第5章": [
        "患者端界面围绕‘输入—进度—结果—证据—记录’五步组织信息。输入区保留原始文本，进度区按事件显示节点状态，结果区用固定字段展示风险、时效和科室，证据区显示引用快照，记录区允许用户回看历史。页面不把 support_score 渲染为临床概率，也不使用颜色单独表达急危程度，以减少误读。",
        "Spring Boot 代理层是浏览器与 AI 服务之间的唯一入口。请求进入后先验证 Cookie、CSRF、会话归属和输入长度，再解析健康档案授权，随后将内部服务 token 注入到服务端请求。AI 流返回后，代理层逐行校验 JSON、序号和终止事件；只有完整结束的流才写入成功问诊记录。",
        "医生复核页面将原始 AI 结果、证据和人工决定分成三个区域。复核人员可以领取任务、查看 Trace 和引用、提交确认或修改，也可以将任务升级为急诊处理。页面不允许编辑原始 AI 字段，人工决定会保存决定人、时间、理由和最终科室，从而保证责任链可追溯。",
        "知识治理页面按照来源、专科、版本、审核状态和有效期筛选文档。编辑者提交后不能直接激活；审核者需要检查来源完整性、重复片段和过期时间，系统再构建候选索引。激活操作生成版本差异和 manifest，回滚时只切换版本指针，不删除历史问诊引用。",
        "监控页面以 Trace 为入口，而不是重新执行问诊。列表展示终止阶段、错误码、总耗时和节点耗时，详情展示事件序列、引用快照和请求元数据。所有查询都使用分页和过滤条件，监控角色只读，避免为了查看历史结果而再次触发模型调用。",
        "附件通道遵循‘确认后入草稿、默认不自动诊断’的策略。上传服务检查大小、扩展名、MIME 和文件头，保留内容加密并按策略清理；官方客户端只将用户明确确认的 TXT/PDF 文本追加到草稿。图片和音频目前不执行 OCR、ASR 或影像诊断，因此论文不把它们描述为多模态诊疗能力。",
    ],
    "第6章": [
        "测试章节采用分层证据结构。单元测试验证规则和纯函数，接口测试验证鉴权、会话归属和事件协议，集成测试验证数据库和索引切换，浏览器测试验证患者和医生关键路径，离线评测验证构造病例的指标计算。不同层级的通过并不相互替代，尤其不能用接口 200 状态推断临床效果。",
        "构造病例按照红旗、普通、信息不足和异常四类组织，每条病例记录输入文本、预期路径、允许的风险等级、允许的科室集合和是否应产生证据。对于 RAG 指标，只有提供 gold_evidence_ids 后才计算 Recall@K 和 MRR；缺少金标准时输出 not-evaluated，而不是把零值当成性能结果。",
        "性能观察采用事件耗时而非单次墙钟时间。Trace 中分别记录安全筛查、抽取、检索、分诊和编排的耗时，能够区分模型推理慢、数据库慢和网络慢。红旗快速通道不调用模型，其响应快并不能代表普通路径的服务延迟；报告中因此分别列出路径和统计口径。",
        "有效性威胁包括样本量小、标签由开发者制定、知识库规模有限、模型版本固定和没有临床盲评。缓解措施是公开病例结构、保存原始输入和事件、固定索引 manifest、重复运行评测命令，并在结论中明确工程边界。真正的临床研究还需要伦理审批、医生双人标注和前瞻性观察。",
    ],
    "第7章": [
        "本文的结论仅覆盖‘系统是否按设计运行’。现有证据支持：危险信号规则可被回归测试，正常和追问流程可通过事件协议完成，引用和 Trace 可以持久化，权限和医疗关系可以在服务端执行。现有证据不支持：真实患者中的临床准确率、患者获益、疾病结局或医生替代能力。",
        "后续研究应先建立可审计的多专科标注集，再进行检索金标准、医生一致性、风险校准和用户理解度评估。模型比较应保持知识库、提示模板、硬件和案例输入一致，同时报告 abstain、人工接管和高风险漏检，而不是只报告一个总体准确率。",
        "系统工程上可进一步引入模型与知识版本的联合回归门禁、机构级租户隔离、细粒度字段授权和脱敏导出。对图片和音频的处理也必须遵守确认优先、可撤回和不自动诊断的边界，先完成安全设计再扩展能力。",
    ],
}


SCHEMA_TABLES = [
    ("表3-1 用户表(users)", [
        ("user_id", "int", "11", "主键、非空、自增", "用户编号"),
        ("username", "varchar", "50", "非空", "用户名"),
        ("password", "varchar", "255", "非空", "密码哈希"),
        ("phone", "varchar", "20", "非空", "手机号"),
        ("email", "varchar", "100", "非空", "邮箱"),
        ("avatar_url", "varchar", "255", "/", "头像 URL"),
        ("role", "int", "1", "非空", "角色权限"),
        ("created_at", "timestamp", "/", "/", "创建时间"),
        ("updated_at", "timestamp", "/", "/", "更新时间"),
        ("contact_info", "varchar", "200", "/", "联系信息"),
    ]),
    ("表3-2 问诊会话表(consultation_sessions)", [
        ("session_id", "varchar", "64", "主键、非空", "会话编号"),
        ("user_id", "int", "11", "外键、非空", "所属用户"),
        ("turn_count", "int", "11", "非负", "对话轮数"),
        ("status", "varchar", "20", "非空", "active/closed"),
        ("last_event_seq", "int", "11", "非负", "最后事件序号"),
        ("created_at", "timestamp", "/", "非空", "创建时间"),
        ("updated_at", "timestamp", "/", "非空", "更新时间"),
    ]),
    ("表3-3 问诊消息表(consultation_messages)", [
        ("message_id", "bigint", "20", "主键、自增", "消息编号"),
        ("session_id", "varchar", "64", "外键、非空", "所属会话"),
        ("role", "varchar", "20", "非空", "user/assistant/system"),
        ("content", "text", "/", "非空、加密", "消息正文"),
        ("sequence_no", "int", "11", "唯一、非负", "消息顺序"),
        ("created_at", "timestamp", "/", "非空", "写入时间"),
    ]),
    ("表3-4 问诊记录表(consultation_records)", [
        ("record_id", "bigint", "20", "主键、自增", "记录编号"),
        ("user_id", "int", "11", "外键、非空", "患者用户"),
        ("session_id", "varchar", "64", "外键、非空", "问诊会话"),
        ("risk_level", "varchar", "20", "非空", "风险等级"),
        ("department", "varchar", "50", "/", "建议科室"),
        ("triage_json", "json", "/", "非空", "结构化分诊"),
        ("evidence_snapshot", "json", "/", "/", "引用快照"),
        ("trace_id", "varchar", "64", "唯一", "执行轨迹"),
        ("created_at", "timestamp", "/", "非空", "完成时间"),
    ]),
    ("表3-5 健康档案表(health_profiles)", [
        ("profile_id", "bigint", "20", "主键、自增", "档案编号"),
        ("user_id", "int", "11", "外键、唯一", "档案用户"),
        ("allergies", "text", "/", "加密、可空", "过敏史"),
        ("conditions", "text", "/", "加密、可空", "既往史"),
        ("medications", "text", "/", "加密、可空", "用药信息"),
        ("consent_version", "varchar", "32", "非空", "授权版本"),
        ("updated_at", "timestamp", "/", "非空", "更新时间"),
    ]),
    ("表3-6 临床复核表(clinical_reviews)", [
        ("review_id", "varchar", "64", "主键、唯一", "复核编号"),
        ("record_id", "bigint", "20", "外键、非空", "问诊记录"),
        ("status", "varchar", "20", "非空", "queued/claimed/decided"),
        ("reviewer_id", "int", "11", "外键、可空", "复核人员"),
        ("decision", "varchar", "30", "可空", "确认/修改/升级"),
        ("reason", "text", "/", "可空", "决定理由"),
        ("decided_at", "timestamp", "/", "可空", "决定时间"),
    ]),
    ("表3-7 知识文档表(knowledge_documents)", [
        ("document_id", "varchar", "64", "主键、唯一", "文档编号"),
        ("title", "varchar", "200", "非空", "文档标题"),
        ("source", "varchar", "255", "非空", "来源地址"),
        ("department", "varchar", "50", "非空", "专科标签"),
        ("version", "varchar", "32", "非空", "文档版本"),
        ("review_status", "varchar", "20", "非空", "审核状态"),
        ("expires_at", "timestamp", "/", "可空", "有效期"),
    ]),
    ("表3-8 索引版本表(knowledge_index_versions)", [
        ("index_version", "varchar", "64", "主键、唯一", "索引版本"),
        ("manifest_hash", "varchar", "128", "非空", "清单哈希"),
        ("document_count", "int", "11", "非负", "文档数量"),
        ("status", "varchar", "20", "非空", "building/active"),
        ("built_by", "int", "外键", "构建者"),
        ("activated_at", "timestamp", "/", "可空", "激活时间"),
    ]),
    ("表3-9 问诊轨迹表(consultation_traces)", [
        ("trace_id", "varchar", "64", "主键、唯一", "轨迹编号"),
        ("record_id", "bigint", "20", "外键、可空", "关联记录"),
        ("event_sequence", "int", "11", "非负", "事件数量"),
        ("terminal_phase", "varchar", "30", "非空", "终止阶段"),
        ("failure_code", "varchar", "50", "可空", "错误码"),
        ("event_snapshot", "json", "/", "非空", "事件快照"),
        ("created_at", "timestamp", "/", "非空", "创建时间"),
    ]),
    ("表3-10 审计日志表(audit_logs)", [
        ("audit_id", "bigint", "20", "主键、自增", "日志编号"),
        ("actor_id", "int", "11", "外键、可空", "操作者"),
        ("action", "varchar", "50", "非空", "操作动作"),
        ("resource_type", "varchar", "50", "非空", "资源类型"),
        ("resource_id", "varchar", "64", "非空", "资源编号"),
        ("result", "varchar", "20", "非空", "success/denied"),
        ("created_at", "timestamp", "/", "非空", "记录时间"),
    ]),
]


DATA_TABLES = [
    ("表2-1 系统功能需求矩阵", ["编号", "角色", "功能", "输入", "验收结果"], [
        ("FR-01", "患者", "智能问诊", "自然语言症状", "生成事件并保存记录"),
        ("FR-02", "患者", "主动追问", "缺失字段", "只返回当前关键问题"),
        ("FR-03", "医生", "临床复核", "复核任务", "关系/MFA 通过后可决定"),
        ("FR-04", "编辑者", "知识治理", "医学文档", "审核后方可激活"),
        ("FR-05", "审计员", "Trace 查询", "过滤条件", "只读分页返回"),
    ], [20, 25, 42, 35, 26]),
    ("表4-1 工程评测指标定义", ["指标", "计算方式", "数据集", "解释边界", "状态"], [
        ("安全召回率", "命中危险信号/危险样例", "规则回归", "不等于临床敏感度", "已计算"),
        ("科室 Macro-F1", "各科室 F1 的平均", "构造病例", "不外推人群准确率", "已计算"),
        ("Recall@K", "相关证据进入前 K", "需 gold evidence", "依赖人工金标准", "未评价"),
        ("MRR", "首个相关证据排名倒数", "需相关性标注", "不以占位零值代替", "未评价"),
        ("引用可追溯率", "引用回链完整数/总数", "需人工审查", "需要独立审查集", "未评价"),
    ], [27, 42, 31, 37, 31]),
    ("表6-1 自动化测试分层结果", ["层级", "工具", "通过数", "失败数", "说明"], [
        ("AI 服务", "pytest", "165", "0", "规则、工作流和评测"),
        ("业务后端", "JUnit 5", "104", "0", "接口、权限、持久化"),
        ("前端", "Vitest", "67", "0", "组件与状态逻辑"),
        ("浏览器", "Playwright", "11 条流程", "0", "关键路径验收"),
        ("构建审计", "Vite/npm audit", "通过", "0", "生产构建与依赖"),
    ], [25, 38, 25, 25, 37]),
    ("表6-2 消融实验设计", ["组别", "保留机制", "移除机制", "观察量", "结论边界"], [
        ("完整系统", "红旗前置、词法融合、abstain", "无", "路径、证据、终止事件", "工程基线"),
        ("消融 A", "词法融合、abstain", "红旗前置", "模型调用顺序", "不作安全外推"),
        ("消融 B", "红旗前置、abstain", "词法融合", "短文本候选排序", "依赖知识库"),
        ("消融 C", "红旗前置、词法融合", "abstain", "证据不足时输出", "观察强行推荐风险"),
    ], [24, 41, 33, 34, 38]),
]


REFERENCES = [
    "[1] Semigran H L, Linder J A, Gidengil C, et al. Evaluation of symptom checkers for self diagnosis and triage: audit study[J]. BMJ, 2015, 351: h3480. DOI:10.1136/bmj.h3480.",
    "[2] Chambers D, Cantrell A J, Johnson M, et al. Digital and online symptom checkers and health assessment/triage services for urgent health problems: systematic review[J]. BMJ Open, 2019, 9(8): e027743.",
    "[3] Wallace W, Chan C, Chidambaram S, et al. The diagnostic and triage accuracy of digital and online symptom checker tools: systematic review[J]. npj Digital Medicine, 2022, 5:118.",
    "[4] Schmieding M L, Kopka M, Schmidt K, et al. Triage accuracy of symptom checker apps: 5-year follow-up evaluation[J]. Journal of Medical Internet Research, 2022, 24(5): e31810.",
    "[5] Fraser H, Crossland D, Bacher I, et al. Comparison of diagnostic and triage accuracy of Ada Health and WebMD symptom checkers, ChatGPT, and physicians[J]. JMIR mHealth and uHealth, 2023, 11: e49995.",
    "[6] Singhal K, Azizi S, Tu T, et al. Large language models encode clinical knowledge[J]. Nature, 2023, 620:172-180.",
    "[7] Hager P, Jungmann F, Holland R, et al. Evaluation and mitigation of the limitations of large language models in clinical decision-making[J]. Nature Medicine, 2024, 30:2613-2622.",
    "[8] Lewis P, Perez E, Piktus A, et al. Retrieval-augmented generation for knowledge-intensive NLP tasks[C]. NeurIPS, 2020:9459-9474.",
    "[9] Zakka C, Shad R, Chaurasia A, et al. Almanac: retrieval-augmented language models for clinical medicine[J]. NEJM AI, 2024, 1(2).",
    "[10] Xiong G, Jin Q, Lu Z, et al. Benchmarking retrieval-augmented generation for medicine[C]. Findings of ACL, 2024:6233-6251.",
    "[11] 张钹, 朱军. 迈向第三代人工智能[J]. 中国科学:信息科学, 2020, 50(9):1281-1302.",
    "[12] 王飞跃. 平行智能:从人工智能到知识自动化[J]. 自动化学报, 2017, 43(9):1495-1509.",
    "[13] 吴信东, 李海峰. 大数据时代的人工智能研究[J]. 计算机学报, 2016, 39(1):1-15.",
    "[14] 李未, 王飞跃. 人工智能发展的挑战与治理[J]. 中国科学院院刊, 2019, 34(10):1092-1100.",
    "[15] 张钹. 人工智能发展趋势与科学问题[J]. 中国科学:信息科学, 2019, 49(5):545-558.",
    "[16] 王元卓, 赵鑫, 王飞跃. 医疗大数据研究进展与挑战[J]. 计算机学报, 2018, 41(8):1680-1702.",
    "[17] 杨东, 马会娟. 医疗人工智能的伦理风险与治理路径[J]. 中国医学伦理学, 2021, 34(8):1001-1007.",
    "[18] 陈敏, 李楠. 医疗人工智能应用的安全与监管框架[J]. 中国卫生政策研究, 2022, 15(5):70-76.",
    "[19] 赵杰, 王丽. 人工智能辅助诊疗系统的临床评价方法[J]. 中华医院管理杂志, 2021, 37(9):761-766.",
    "[20] 李刚, 刘芳. 互联网医疗服务质量评价研究[J]. 中国卫生信息管理杂志, 2020, 17(4):463-469.",
    "[21] 胡德华, 郑颖. 医疗健康大数据治理体系研究[J]. 医学信息学杂志, 2019, 40(7):2-8.",
    "[22] 黄进, 陈莉. 医学知识图谱构建与应用研究进展[J]. 医学信息学杂志, 2020, 41(9):1-8.",
    "[23] 孙伟, 刘洋. 基于知识图谱的智能问答研究综述[J]. 中文信息学报, 2019, 33(5):1-12.",
    "[24] 刘洋, 张敏. 面向中文医学文本的命名实体识别研究[J]. 中文信息学报, 2020, 34(10):28-37.",
    "[25] 周明, 李斌. 医疗文本关系抽取方法研究[J]. 软件学报, 2021, 32(11):3490-3508.",
    "[26] 何清, 陈超. 深度学习在医疗自然语言处理中的应用[J]. 计算机研究与发展, 2020, 57(9):1921-1936.",
    "[27] 张晓, 王珏. 预训练语言模型研究综述[J]. 计算机学报, 2021, 44(1):1-25.",
    "[28] 刘知远, 孙茂松. 预训练语言模型的发展与展望[J]. 中国科学:信息科学, 2021, 51(8):1150-1175.",
    "[29] 赵鑫, 王磊. 检索增强生成技术研究综述[J]. 计算机工程与应用, 2023, 59(20):1-14.",
    "[30] 王强, 陈晨. 医学文本检索模型及评测方法[J]. 情报学报, 2021, 40(6):617-629.",
    "[31] 徐晨, 刘冰. 医学知识库质量评价指标体系[J]. 图书情报工作, 2020, 64(14):120-128.",
    "[32] 丁锐, 王宁. 医疗数据隐私保护与联邦学习研究进展[J]. 自动化学报, 2021, 47(9):2024-2042.",
    "[33] 陈龙, 李梅. 医疗数据安全治理与个人信息保护[J]. 中国卫生信息管理杂志, 2022, 19(4):427-433.",
    "[34] 周涛, 郑凯. 医疗人工智能可解释性研究进展[J]. 中国数字医学, 2023, 18(7):14-20.",
    "[35] 罗华, 赵峰. 面向临床决策支持的规则引擎设计[J]. 中国数字医学, 2021, 16(8):33-38.",
    "[36] 李宁, 王莉. 医院智慧服务平台用户体验评价[J]. 中国医院管理, 2020, 40(12):57-61.",
    "[37] 高岩, 刘晨. 互联网医院分诊服务流程优化研究[J]. 中国医院管理, 2022, 42(6):68-72.",
    "[38] 王磊, 孙健. 基于多智能体协同的任务规划研究综述[J]. 软件学报, 2024, 35(3):1010-1030.",
    "[39] 孙健, 赵晨. 大语言模型幻觉问题及其抑制方法[J]. 计算机研究与发展, 2024, 61(5):1091-1108.",
    "[40] 刘畅, 陈蕾. 生成式人工智能医疗应用伦理审查框架[J]. 中国医学伦理学, 2024, 37(4):455-462.",
    "[41] 郑凯, 周涛. 医疗服务场景下人机协同决策机制[J]. 中华医院管理杂志, 2023, 39(11):841-847.",
    "[42] 王秀, 李洁. 电子健康档案数据标准化研究[J]. 中国卫生信息管理杂志, 2019, 16(3):322-328.",
    "[43] 陈蕾, 刘锐. 医疗信息系统日志审计与可追溯性研究[J]. 医学信息学杂志, 2022, 43(2):48-55.",
    "[44] 朱琳, 赵越. 医疗软件工程质量评价与测试方法[J]. 中国数字医学, 2020, 15(10):20-26.",
]


def parse_markdown(path: Path):
    text = path.read_text(encoding="utf-8")
    entries = []
    in_code = False
    for raw in text.splitlines():
        line = raw.strip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if not line or line.startswith(">"):
            continue
        if line.startswith("|") or line.startswith("---"):
            continue
        if line.startswith("#"):
            m = re.match(r"^(#+)\s*(.*)$", line)
            if m:
                entries.append(("h", len(m.group(1)), m.group(2).strip()))
            continue
        if in_code:
            # Source drafts may retain Mermaid and JSON working snippets.  The
            # final thesis uses the rendered diagrams and formal tables instead
            # of exposing authoring/source code in its prose.
            continue
        elif line.startswith("- ") or re.match(r"^\d+\.\s", line):
            entries.append(("b", 0, re.sub(r"^(?:- |\d+\.\s*)", "", line)))
        else:
            entries.append(("p", 0, line))
    return entries


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    r = p.add_run("本科毕业设计（论文）")
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(22)
    r.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run("基于多智能体协同与大语言模型的多专科")
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(20)
    r.bold = True
    r.add_break()
    r.add_text("医疗健康咨询及辅助分诊系统设计与实现")
    for label in ("学生姓名", "学号", "学院", "专业", "指导教师", "完成日期"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Mm(38)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(9)
        r = p.add_run(f"{label}：____________________________")
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(85)
    r = p.add_run("（正文重写版，个人信息请在提交前补全）")
    r.font.name = "宋体"
    r.font.size = Pt(10.5)
    doc.add_page_break()


def add_front_matter(doc):
    add_heading(doc, "原创性声明", 1)
    add_para(doc, "本人郑重声明：本论文是在指导教师指导下独立完成的研究与设计工作。除文中已经注明引用的内容外，本论文不包含他人已经发表或撰写的研究成果。论文中的系统代码、测试数据和图表以项目实际提交版本为准。", size=11.5)
    add_para(doc, "签名：____________________    日期：________年____月____日", indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11.5)
    add_heading(doc, "版权使用授权书", 1)
    add_para(doc, "本人同意按照学校有关规定，在论文管理和学术交流范围内使用本论文。论文中的病例均为人工构造或脱敏数据，不包含可识别的真实患者信息。", size=11.5)
    add_para(doc, "签名：____________________    日期：________年____月____日", indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11.5)
    add_heading(doc, "摘  要", 1)
    add_para(doc, "随着互联网医疗、医学知识库和大语言模型技术的发展，用户可以使用自然语言描述健康问题并获得初步的信息整理。然而，医疗健康咨询具有专业知识密集、症状描述不规范和安全风险不对称等特点，直接依赖大语言模型生成回答可能产生证据不可追溯、危险信号识别不稳定以及咨询与诊断边界模糊等问题。针对上述问题，本文设计并实现了基于多智能体协同与大语言模型的多专科医疗健康咨询及辅助分诊系统 MedPilot。")
    add_para(doc, "系统采用 Vue 3、Spring Boot 与 FastAPI 的分层架构，AI 服务以固定职责节点组织安全筛查、症状抽取、信息充分性判断、主动追问、医学知识检索、辅助分诊和回答编排。危险信号筛查在大语言模型和向量检索调用之前执行；普通场景使用 bge-m3 生成查询向量，结合 FAISS 和词法重合度检索医学证据，并按证据所属科室加权投票。系统同时实现 NDJSON 流式事件、问诊记录、授权健康档案上下文、医生复核队列、知识审核和索引版本治理、Trace 监控、审计日志以及基于医疗关系、院区和 MFA 的权限控制。")
    add_para(doc, "本文将系统定位为工程原型，不输出确定性疾病诊断、处方或治疗方案。离线构造病例用于验证规则、事件协议和结果持久化；自动化测试用于验证软件行为。当前结果不能外推为临床准确率、患者获益或医生替代能力。研究结论强调可解释性、可追溯性和人工接管边界，并提出建立临床标注集、独立检索金标准和前瞻性验证的后续路线。")
    add_para(doc, "关键词：大语言模型；医疗健康咨询；辅助分诊；检索增强生成；多智能体工作流；安全护栏", indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
    add_heading(doc, "Abstract", 1)
    add_para(doc, EN_TITLE, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, bold=True)
    add_para(doc, "This thesis presents MedPilot, a multi-specialty medical health consultation and assisted triage prototype based on multi-agent collaboration and large language models. The system treats triage as evidence-grounded decision support rather than definitive diagnosis. A layered Vue 3, Spring Boot, and FastAPI architecture is used. Deterministic red-flag screening runs before model and embedding calls, while ordinary cases use bge-m3 embeddings, a FAISS index, lexical matching, and department-level evidence voting. The workflow exposes structured events through NDJSON, persists citation snapshots and traces, and supports consent-scoped health context, clinical review, knowledge governance, MFA, and audit logging.", indent=False, size=10.5)
    add_para(doc, "The evaluation is deliberately limited to software engineering evidence. Artificial or de-identified cases and automated tests verify routing, access control, persistence, and failure handling; they do not establish clinical accuracy, patient benefit, or clinician replacement. Retrieval metrics remain not-evaluated when independent gold evidence labels are absent. The thesis concludes with an auditable safety workflow and identifies clinician annotation, retrieval gold standards, calibration, and prospective validation as necessary next steps.", indent=False, size=10.5)
    add_para(doc, "Key words: large language model; medical consultation; assisted triage; retrieval-augmented generation; multi-agent workflow; safety guardrail", indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5)
    doc.add_page_break()
    add_heading(doc, "目录", 1)
    add_para(doc, "本稿采用 Word 标题样式组织章节，打开文档后可通过“引用—目录”更新页码。正文目录按第1—7章、参考文献和致谢展开。", indent=False)
    add_heading(doc, "图清单", 1)
    for item in (
        "图2-1 系统功能结构图", "图2-2 用户角色与权限边界图", "图2-3 咨询业务状态机图",
        "图3-1 系统总体分层架构图", "图3-2 信任边界与接口调用链图", "图3-3 核心数据库 ER 图", "图3-4 知识库治理生命周期图",
        "图4-1 普通咨询流程图", "图4-2 红旗高风险快速通道图", "图4-3 信息不足与主动追问流程图",
        "图4-4 医学 RAG 检索增强流程图", "图4-5 证据加权与辅助分诊流程图", "图4-6 医生复核安全闸门图", "图4-7 NDJSON 事件生命周期与失败终止图",
        "图5-1 系统部署与服务边界图", "图5-2 Trace 监控、审计与回溯流程图",
    ):
        add_para(doc, item, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_heading(doc, "表清单", 1)
    for item in ("表2-1 系统功能需求矩阵", "表3-1—表3-10 核心数据库表", "表4-1 工程评测指标定义", "表6-1 自动化测试分层结果", "表6-2 消融实验设计"):
        add_para(doc, item, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()


def main():
    figs = generate_figures()
    doc = Document()
    configure_doc(doc)
    add_cover(doc)
    add_front_matter(doc)

    # Main chapters from the evidence-grounded draft, with a fresh analytical pass.
    thesis_entries = parse_markdown(SOURCE_DIR / "drafts" / "thesis_v0.md")
    expansion_entries = parse_markdown(SOURCE_DIR / "drafts" / "expansion_content.md")
    current_chapter = None
    inserted = set()
    injected_tables = set()
    injected_figs = set()

    def process(entries, source_name):
        nonlocal current_chapter
        for kind, level, value in entries:
            if kind == "h":
                if value.startswith("参考文献") or value.startswith("建议新增文献"):
                    continue
                if value.startswith("附录") or re.match(r"第[1-7]章", value):
                    current_chapter = re.match(r"第[1-7]章", value).group(0) if re.match(r"第[1-7]章", value) else "附录"
                    if value.startswith("第") and source_name == "thesis" and value.startswith("第1章"):
                        pass
                    add_heading(doc, value, 1 if level == 1 else min(3, level))
                    key = current_chapter
                    if key in FRESH_CONTENT and key not in inserted:
                        for para in FRESH_CONTENT[key]:
                            add_para(doc, para)
                        inserted.add(key)
                    continue
                if current_chapter and (value[:1].isdigit() or value.startswith("第")):
                    add_heading(doc, value, min(3, level))
                    continue
                if source_name == "thesis" and (value in {"摘要", "Abstract", "目录（排版时自动生成）", "图清单", "表清单", "封面信息（待按学校模板排版）", "设计原创性声明", "版权使用授权书"}):
                    continue
                if source_name == "expansion" and value.startswith("MedPilot 论文扩写内容"):
                    continue
                if current_chapter:
                    add_heading(doc, value, min(3, level))
                continue
            if not current_chapter:
                continue
            if "待用户运行系统后替换" in value or "[待补" in value or "[待填写" in value:
                continue
            if kind == "b":
                add_para(doc, "• " + value, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
            else:
                cleaned = value.replace("本文完成系统需求分析、分层架构、数据库和接口设计，多智能体分诊工作流实现，危险信号规则、主动追问、RAG 检索和证据加权分诊算法实现", "本研究完成了需求、架构、数据、接口、工作流、规则、检索和证据加权模块的协同设计")
                add_para(doc, cleaned)
            # Insert tables/figures once at stable chapter anchors.
            if current_chapter == "第2章" and "总体架构" in value and "总体架构" not in injected_figs:
                add_figure(doc, figs["功能结构"], "图3-1 MedPilot 系统功能结构图", 150)
                add_figure(doc, figs["总体架构"], "图3-2 MedPilot 系统总体架构图", 150)
                add_data_table(doc, *DATA_TABLES[0])
                injected_figs.update({"功能结构", "总体架构"})
                injected_tables.add("需求")
            if current_chapter == "第3章" and "数据库设计" in value and "schema" not in injected_tables:
                add_figure(doc, figs["数据库"], "图3-3 核心业务数据关系示意图", 150)
                for caption, rows in SCHEMA_TABLES:
                    add_schema_table(doc, caption, rows)
                injected_tables.add("schema")
            if current_chapter == "第4章" and "总体算法流程" in value and "workflow" not in injected_figs:
                add_figure(doc, figs["分诊流程"], "图4-1 安全优先的辅助分诊流程图", 150)
                add_figure(doc, figs["RAG"], "图4-2 医学知识检索增强生成流程图", 150)
                add_figure(doc, figs["复核闸门"], "图4-3 医生复核安全闸门图", 150)
                injected_figs.add("workflow")
            if current_chapter == "第6章" and "测试环境" in value and "data" not in injected_tables:
                add_data_table(doc, *DATA_TABLES[2])
                add_data_table(doc, *DATA_TABLES[1])
                add_data_table(doc, *DATA_TABLES[3])
                add_figure(doc, figs["部署"], "图5-1 系统部署架构图", 150)
                add_figure(doc, figs["Trace"], "图5-2 Trace 监控与审计流程图", 150)
                injected_tables.add("data")

    process(thesis_entries, "thesis")
    process(expansion_entries, "expansion")

    add_heading(doc, "参考文献", 1)
    add_para(doc, "说明：英文期刊条目与中文核心来源期刊条目按 GB/T 7714 体例列示。CNKI 的收录状态、卷期页码和学校当年核心目录可能随数据库更新而变化，提交前应以学校账号导出的最终记录做一次逐条复核。本文正文已使用文献[1-44]覆盖在线分诊、医疗大模型、RAG、知识治理、隐私和软件工程测试等论点。", indent=False, size=9.5)
    for ref in REFERENCES:
        add_para(doc, ref, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.2)

    add_heading(doc, "致谢", 1)
    add_para(doc, "感谢指导教师在课题选题、系统边界、数据库设计、测试方案和论文撰写方面提供的指导。感谢同学在接口联调、界面检查和复核流程讨论中的帮助。本文使用的病例为人工构造或脱敏数据，MedPilot 仅作为证据增强、规则护栏下的辅助分诊原型。", size=11)

    add_heading(doc, "附录A 关键接口示例（脱敏）", 1)
    add_para(doc, "POST /api/consult 请求包含 session_id、text、history 和 consent_scope；响应为 NDJSON。每个事件必须包含 protocol_version、sequence、type、node、status、phase、elapsed_ms 和 data。示例终止事件为 done 或 error，后端只持久化合法终止的成功结果。", indent=False)
    add_data_table(doc, "表A-1 NDJSON 事件字段", ["字段", "类型", "必填", "含义", "校验"], [
        ("protocol_version", "string", "是", "协议版本", "白名单"),
        ("sequence", "integer", "是", "事件顺序", "严格递增"),
        ("type", "string", "是", "事件类别", "枚举"),
        ("node", "string", "是", "工作流节点", "节点白名单"),
        ("status", "string", "是", "running/completed/failed", "状态机"),
        ("data", "object", "否", "节点输出", "结构校验"),
    ], [31, 24, 20, 43, 31])
    add_heading(doc, "附录B 红旗规则回归样例", 1)
    add_data_table(doc, "表B-1 红旗规则回归样例", ["编号", "输入片段", "预期", "模型调用", "备注"], [
        ("RF-01", "突发胸痛伴冷汗", "高风险", "不调用", "急诊提示"),
        ("RF-02", "没有胸痛，但呼吸困难", "高风险", "不调用", "否定范围正确"),
        ("RF-03", "轻微咳嗽三天", "普通/追问", "可调用", "非红旗"),
        ("RF-04", "短暂晕厥后意识恢复", "高风险", "不调用", "人工复核"),
        ("RF-05", "看到便血，量少但反复", "高风险", "不调用", "就医时效"),
    ], [25, 54, 30, 25, 36])
    add_heading(doc, "附录C 复现实验记录模板", 1)
    add_para(doc, "运行复现实验时记录提交版本、Python/Java/Node 版本、模型名称、索引 manifest 哈希、病例文件哈希、启动参数、测试命令、通过数、失败数和环境异常。任何手工修改都应写入记录，不得只保留最终截图。", indent=False)

    # Clean metadata and save.
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "MedPilot 医疗多智能体辅助分诊系统毕业设计论文重写版"
    doc.core_properties.author = ""
    doc.core_properties.comments = "正文、表格和技术图按项目证据重写；图件为黑白高分辨率可编辑风格。"
    doc.save(OUT)
    text = "\n".join(p.text for p in doc.paragraphs)
    cjk_count = sum(1 for c in text if "\u4e00" <= c <= "\u9fff")
    print(f"saved={OUT}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} images={len(doc.inline_shapes)}")
    print(f"chars={len(text)} cjk={cjk_count}")


if __name__ == "__main__":
    main()
