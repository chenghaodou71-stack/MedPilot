from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"D:\毕设制作")
SOURCE = ROOT / "outputs" / (
    "7230264110_窦成皓_23软件1_基于多智能体协同与大语言模型的"
    "多专科医疗健康咨询及辅助分诊系统_核心类图数据库精简与引用重排版.docx"
)
OUTPUT = ROOT / "outputs" / (
    "7230264110_窦成皓_23软件1_基于多智能体协同与大语言模型的"
    "多专科医疗健康咨询及辅助分诊系统_论文综合优化版.docx"
)


def set_text(paragraph, text: str) -> None:
    props = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        props = copy.deepcopy(paragraph.runs[0]._r.rPr)
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if props is not None:
        run._r.insert(0, props)


def find_contains(doc, needle: str):
    compact = needle.replace(" ", "")
    matches = [
        p for p in doc.paragraphs
        if needle in p.text or compact in p.text.replace(" ", "")
    ]
    if len(matches) != 1:
        raise RuntimeError(f"expected one paragraph containing {needle!r}, found {len(matches)}")
    return matches[0]


def find_exact(doc, text: str):
    matches = [p for p in doc.paragraphs if p.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"expected one paragraph {text!r}, found {len(matches)}")
    return matches[0]


def add_before(doc, anchor, text: str, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(24 if align == WD_ALIGN_PARAGRAPH.JUSTIFY else 0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.widow_control = True
    r = p.add_run(text)
    r.font.name = "宋体"
    r.font.size = Pt(12)
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    anchor._p.addprevious(p._p)
    return p


def set_cell_text(cell, text: str, *, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "宋体"
    r.font.size = Pt(size)
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def set_table_borders(table, color="808080", size="4"):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def add_ablation_table(doc, anchor):
    add_before(
        doc,
        anchor,
        "为区分否定感知和分句处理对安全筛查的贡献，本文在同一固定种子的1000条测试输入上设置一个最小消融：完整方案使用DANGER_SIGN_TERMS、同义词表和否定感知分句；消融方案仅保留规范词与同义词的简单子串匹配，去除否定和连接词继承逻辑。两种方案使用相同的金标准和统计口径，结果如表5-4所示。",
    )
    # Insert the caption first so the table follows it when both are placed
    # immediately before the chapter-summary anchor.
    caption = add_before(doc, anchor, "表5-4 基线与消融结果", align=WD_ALIGN_PARAGRAPH.CENTER)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [42, 30, 35, 45]
    for cell, width in zip(table.rows[0].cells, widths):
        cell.width = Pt(width * 2.835)
    headers = ["方案", "危险信号召回率", "非高风险特异度", "实现条件"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_text(cell, value, bold=True)
        shade(cell, "F2F2F2")
    header_props = table.rows[0]._tr.get_or_add_trPr()
    header_marker = OxmlElement("w:tblHeader")
    header_marker.set(qn("w:val"), "true")
    header_props.append(header_marker)
    rows = [
        ("完整方案", "0.8000", "1.0000", "分句、否定感知、连接词继承"),
        ("消融方案", "0.7937", "0.8941", "规范词/同义词简单子串匹配"),
        ("差值（完整-消融）", "+0.0063", "+0.1059", "相同1000条输入与金标准"),
    ]
    for values in rows:
        row = table.add_row()
        row_props = row._tr.get_or_add_trPr()
        row_props.append(OxmlElement("w:cantSplit"))
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            set_cell_text(cell, value, align=WD_ALIGN_PARAGRAPH.CENTER if index < 3 else WD_ALIGN_PARAGRAPH.LEFT)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    set_table_borders(table)
    table._tbl.tblGrid.gridCol_lst[0].set(qn("w:w"), "2380")
    table._tbl.tblGrid.gridCol_lst[1].set(qn("w:w"), "1700")
    table._tbl.tblGrid.gridCol_lst[2].set(qn("w:w"), "1980")
    table._tbl.tblGrid.gridCol_lst[3].set(qn("w:w"), "2550")
    anchor._p.addprevious(table._tbl)
    return table


def add_table_list_row(doc):
    list_table = doc.tables[1]
    row = list_table.add_row()
    set_cell_text(row.cells[0], "表5-4", size=10.5)
    set_cell_text(row.cells[1], "基线与消融结果", align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5)
    set_cell_text(row.cells[2], "0", size=10.5)


def main():
    doc = Document(str(SOURCE))

    set_text(
        find_contains(doc, "本文首先分析用户、管理员、医生复核人员和知识维护人员"),
        "本文首先分析用户、管理员、医生复核人员和知识维护人员的业务需求，论证系统在技术、经济、操作和社会层面的可行性，并说明 LangGraph、FastAPI、Ollama、FAISS 与 NDJSON 等技术在系统中的作用。研究方法采用工程设计研究与可复现软件评测相结合的方式：以项目源代码、Flyway 迁移脚本、审核后的知识文档和页面流程作为设计依据，以固定种子生成的人工构造测试集作为工程验证数据，不把它们解释为临床病例。技术路线按照“需求分析—分层架构—条件工作流实现—数据与治理落地—千例回归测试—结果解释”的顺序展开，评价指标包括危险信号召回率与特异度、科室和风险等级 Macro-F1、Recall@3、MRR、追问触发率及组件耗时，并用否定感知规则的消融对照检验关键设计的作用。",
    )

    set_text(
        find_contains(doc, "LangGraph 使用有向图表示具有状态和条件分支"),
        "LangGraph 使用有向图表示具有状态和条件分支的处理流程。本文将“多智能体”限定为多个具有固定职责、独立输入输出契约和可观测事件的流程节点：安全筛查节点负责危险信号识别，症状节点负责结构化抽取，追问节点负责信息补全，检索节点负责证据召回，分诊节点负责规则化投票，回答节点负责确定性编排。节点之间通过共享状态传递结果，由条件边决定下一步路径；它们不是可以脱离系统边界独立诊断的自治医生。该定义使每个节点都能单独测试、替换和回溯，也与多智能体会话研究所强调的职责分工相一致[10]。",
    )

    set_text(
        find_contains(doc, "FAISS 用于保存医学知识切片的向量"),
        "FAISS 用于保存医学知识切片的向量并执行归一化内积检索，归一化后可视为余弦相似度。设向量分数为 S_dense，中文字符二元组重合度为 S_lex，系统采用 S = 0.8S_dense + 0.2S_lex 作为综合分；候选池取 min(4K，N)，其中 K=3、N 为索引分片数，过滤阈值为 S≥0.35，最终返回综合分最高的 Top-3。0.8/0.2 的设置用于保持语义相似度的主导地位，同时让短症状中的字面命中能够参与排序；0.35 和 Top-3 是当前 Demo 的工程默认值，依据低分噪声过滤、页面证据数量和响应开销之间的平衡确定，不把它们表述为全局最优参数。检索增强生成研究将外部证据引入生成流程，并强调检索模块与生成模块之间的接口约束[14]。医疗服务调度研究提示，优先级和资源分配也需要在流程层显式表达[15]。医学 RAG 综述指出，文档切分、查询表达和候选重排都会影响检索效果[16]。",
    )

    set_text(
        find_contains(doc, "系统使用 MySQL 保存业务数据，Flyway 负责"),
        "系统使用 MySQL 保存业务数据，Flyway 负责按版本创建和扩展表结构。本节围绕一次问诊从身份确认、患者映射、会话交互到结果复核的主链路，选取10张核心业务表进行说明。选表依据是对象—表的一致映射、主链路中的直接引用关系以及结果可追溯所需的最小字段集合；模型发布、红队测试和回滚演练等治理表仍由迁移脚本维护，但不在本节逐表展开。表中“约束”综合标记主键、非空、自增、唯一、外键和默认值，外键或业务归属字段用于限制跨用户读取和跨会话写入。",
    )

    set_text(
        find_contains(doc, "危险信号筛查位于 AI 服务 run_consult_stream 的最前端"),
        "危险信号筛查位于 AI 服务 run_consult_stream 的最前端。词表来源限定为项目内由 DANGER_SIGN_TERMS 和 _ALIASES 维护的工程规则：规范词包括胸痛、呼吸困难、气促、咯血、便血、意识不清、晕厥和大出血，当前登记的同义词覆盖气短、黑便、昏倒、咳血、咳出血和咳出了血。该词表用于软件路径控制，不等同于临床指南或疾病诊断标准。match_danger_signs 首先按标点、换行和转折词切分语句，再判断词项前是否存在“没有、否认、未见、无”等否定表达，并在“和、及、以及、或、也、并”等连接词下继承前一词项的否定状态；只有被断言出现的规范词才进入高风险路径。",
    )

    set_text(
        find_contains(doc, "信息充分后，_retrieve_node 将原始文本和结构化症状连接为检索查询"),
        "信息充分后，_retrieve_node 将原始文本和结构化症状连接为检索查询。检索器先调用 expand_query_with_aliases 扩展医学同义表达，再通过 bge-m3 生成归一化向量，在 FAISS 中取得 Top-3 的四倍候选集合。对每个候选分片计算 S_dense 与 S_lex，并按 S=0.8S_dense+0.2S_lex 融合；综合分低于0.35的候选被过滤，剩余结果按分数排序后保留3条。0.8的向量权重来自语义检索的主任务地位，0.2的词法权重用于保留短症状的直接命中；阈值和 K 值与2.3.4节保持一致，便于测试复现。检索增强研究指出，候选重排和证据组织会影响医学问答质量[16]；因此，系统保留通过审核的证据、来源文档和索引版本，并在后续按科室分数加权。",
    )

    set_text(
        find_contains(doc, "测试策略按照“单节点—业务链路—系统边界”的顺序设计"),
        "测试策略按照“单节点—业务链路—系统边界”的顺序设计。单节点测试关注危险信号匹配、否定表达、信息充分性、追问生成、科室投票、证据门控和数据加密；业务链路测试关注从登录到问诊、从问诊到记录、从知识入库到索引激活、从 AI 结果到医生复核的完整过程；系统边界测试关注权限、超时、容量、断流、模型不可用和数据库异常。每条测试输入在运行前写入固定字段：task 表示任务类别，gold_route 表示预期路径，gold_department 和 gold_risk 表示分类金标准，gold_red_flags 表示危险信号集合，acceptable_doc_ids 表示可接受证据。高风险样本以 gold_red_flags 非空为判定条件，证据分诊样本要求结构化症状充分且存在目标科室，低证据样本来自受支持科室之外的信息不足输入，追问样本要求症状为空或只有单症状且缺少持续时间和严重程度。P0至P4只描述输入扰动强度，不改变金标准；测试集由 benchmark_1000.py 以种子20260831生成后冻结。TRIPOD-LLM 报告规范要求明确说明任务、数据和评价过程[22]。TRIPOD-LLM 指南还要求交代研究适用边界和报告细节[23]。",
    )

    set_text(
        find_contains(doc, "危险信号召回率用于衡量已标注危险信号"),
        "危险信号召回率用于衡量已标注危险信号是否被系统识别，定义为 Recall_safe = TP_safe / (TP_safe + FN_safe)。该指标只反映规则筛查对测试样本的覆盖情况，不代表临床筛查性能。医疗大语言模型的人类评价框架指出，安全与伤害应作为独立评价维度[13]；因此，本文将危险信号召回率与否定表达、词表外表达和普通样本误触发情况分开统计。",
    )
    set_text(
        find_contains(doc, "危险信号召回率 = 正确识别的危险信号样本数"),
        "危险信号召回率 = 正确识别的危险信号样本数 / 危险信号样本总数；非高风险特异度 = 正确未触发样本数 / 非高风险样本总数。",
    )
    set_text(
        find_contains(doc, "科室和风险等级采用Precision、Recall和F1进行统计"),
        "科室和风险等级采用 Precision、Recall 和 F1 进行统计：Precision=TP/(TP+FP)，Recall=TP/(TP+FN)，F1=2×Precision×Recall/(Precision+Recall)。Macro-F1=(1/C)ΣF1_c，其中 C 为标签数量，先分别计算每个标签的 F1，再取算术平均，以避免样本较多的类别完全主导结果。",
    )
    set_text(
        find_contains(doc, "信息不足样本不纳入科室和风险F1"),
        "信息不足样本不纳入科室和风险 F1，而单独统计追问触发率。证据支持样本统计 Recall@3=命中可接受证据的样本数/证据支持样本数，并统计 MRR@3=(1/M)Σ(1/r_i)，其中 r_i 为第 i 条样本首个可接受证据在前三位中的名次，未命中时记为0。科室指标分别报告四个受支持科室的 Macro-F1，以及加入全科/线下回退类别后的五分类 Macro-F1。组件耗时从函数实际开始执行到返回结果为止，不包含并发队列、浏览器、业务后端、数据库和网络传输，因此不等同于线上端到端延迟。上述指标只描述当前 Demo 在人工构造工程测试集上的行为，不用于推断临床准确率。",
    )

    set_text(
        find_contains(doc, "千例评测同时覆盖主流程和边界输入"),
        "本轮结果按任务边界和评价方法解释，不把组件指标直接外推为医疗效果[25]。高风险优先规则的设计与急诊分诊筛查研究强调的时间敏感性相呼应[26]。互联网医院流程研究提示，入口、信息提交和路径衔接需要被视为连续服务过程[27]。医疗健康领域大语言模型综述将文本处理、问答和辅助决策视为主要应用方向[28]。基层医疗研究则强调部署时需要同步考虑场景约束和使用责任[29]。RAG 方法综述把检索、重排和生成之间的衔接作为系统设计对象[30]。",
    )
    anchor_58 = find_exact(doc, "5.8 本章小结")
    add_before(
        doc,
        anchor_58,
        "面向医学问答的检索增强研究进一步强调领域知识与证据组织的结合[31]。融合反思机制的研究说明，结果校验可以作为生成流程的附加控制环节[32]。粒计算检索研究为多粒度证据组织提供了可参考的实现视角[33]。医疗大语言模型综述显示，模型能力应放在具体任务和数据环境中评价[34]。医疗不良事件研究将模型分类与人工分类进行对比，说明工程评测需要设置可解释的参照方法[35]。",
    )
    add_ablation_table(doc, anchor_58)
    add_before(
        doc,
        anchor_58,
        "表5-4显示，完整方案与简单子串消融在召回率上差异较小，但非高风险特异度相差10.59个百分点。差异主要来自否定表达和连接词继承：消融方案会把“没有胸痛，也没有呼吸困难”中的词项直接视为命中，而完整方案先判断语句断言状态。该对照只用于说明软件规则的增量作用，不能替代临床人工评价。",
    )

    set_text(
        find_contains(doc, "后续工作可以沿着数据、模型、流程和工程四个方向展开"),
        "后续工作可以沿着数据、模型、流程和工程四个方向展开。数据方面，医疗数据脱敏研究提示敏感字段处理需要与模型流程同时设计[36]；数据投毒研究说明知识来源和输入环境应纳入安全边界[37]；智能化语言服务研究则强调医疗文本交互的场景化特征[38]。结合本文实现，后续数据建设应继续维护词表、同义映射、知识来源和审核记录，并保持每次更新可追溯。",
    )
    set_text(
        find_contains(doc, "模型方面，可以比较不同本地聊天模型和嵌入模型"),
        "模型方面，可以比较不同本地聊天模型和嵌入模型在症状抽取、追问生成和检索排序上的差异。大语言模型赋能医疗卫生智库的研究提示，模型应用需要嵌入明确的知识服务流程[39]。基于专利数据的研究呈现了国内医疗大语言模型应用的技术分布[40]。在线医疗社区虚假评论识别研究说明，合成数据和模型输出的质量需要分别检查[41]。医疗聊天机器人研究梳理了从传统方法到大语言模型的演进路径[42]。CiteSpace 研究则反映出医疗大语言模型研究热点仍在持续变化[43]。因此，后续评价应同时覆盖信息完整性、危险信号召回、证据相关性、解释一致性和回答安全性，不能只比较文本相似度。",
    )
    set_text(
        find_contains(doc, "流程方面，可以把医生复核反馈用于规则和知识的回归分析"),
        "流程方面，可以把医生复核反馈用于规则和知识的回归分析，建立从异常 Trace 到变更请求、验证证据和回滚结果的闭环。人工智能嵌入临终关怀的伦理研究提醒，系统进入敏感服务流程时需要明确责任边界[44]。数字赋能健康治理研究强调，技术能力应与服务流程和治理机制协同建设[45]。人工智能应用伦理安全治理研究进一步提出，应把责任、审计和风险处置落实到可执行的管理环节[46]。对于需要医院流程接入的场景，还需要在获得机构授权后对接统一身份、患者主索引、就诊关系和院内科室编码，并重新确认数据访问范围。",
    )
    set_text(
        find_contains(doc, "工程方面，可以继续完善并发控制、缓存策略、索引切换"),
        "工程方面，可以继续完善并发控制、缓存策略、索引切换、监控告警和备份恢复，开展更大规模的压力测试和故障演练。前端还可以在保持当前状态边界的基础上，优化移动端问诊、引用查看和复核协同体验。上述展望均以当前 Demo 的工程验证范围为前提，不把离线测试结果延伸为临床部署结论。",
    )

    add_table_list_row(doc)
    doc.core_properties.subject = "去除研究问题式表述，补充方法、参数、标注、消融对照与文献支撑"
    doc.save(str(OUTPUT))
    print(f"saved={OUTPUT}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} images={len(doc.inline_shapes)}")


if __name__ == "__main__":
    main()
