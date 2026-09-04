from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re

from PIL import Image
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"D:\毕设制作")
SOURCE = Path(
    r"C:\Users\Dou-ChengHao\Desktop\7230264110_窦成皓_23软件1_"
    r"基于多智能体协同与大语言模型的多专科医疗健康咨询及辅助分诊系统.docx"
)
OUTPUT = ROOT / "outputs" / (
    "7230264110_窦成皓_23软件1_基于多智能体协同与大语言模型的"
    "多专科医疗健康咨询及辅助分诊系统_第四章图文增强版.docx"
)
ASSET_DIR = ROOT / ".scratch" / "chapter4-assets"


FIGURES = {
    "4-1": ("用户权限管理页面", "26_用户权限管理_裁剪.png"),
    "4-2": ("高风险快速通道结果", "08_危险信号预警结果.png"),
    "4-3": ("主动追问交互页面", "06_问诊追问交互.png"),
    "4-4": ("检索证据支持的辅助分诊结果", "07_辅助分诊结果.png"),
    "4-5": ("多智能体协同分析过程", "05_多智能体协同分析过程.png"),
    "4-6": ("问诊记录的风险与证据链可视化", "12_智能体执行轨迹与证据链.png"),
    "4-7": ("医学知识库与索引版本管理页面", "21_医学知识库管理.png"),
    "4-8": ("智能体运行监控页面", "24_智能体运行监控.png"),
    "4-9": ("问诊调用链路详情", "25_调用链路详情.png"),
    "4-10": ("审计日志查询页面", "28_审计日志查询.png"),
}


LISTINGS = {
    "4-1": {
        "title": "关键资源角色授权规则",
        "path": ROOT / "backend" / "src" / "main" / "java" / "com" / "medpilot" / "config" / "SecurityConfig.java",
        "start": 102,
        "end": 117,
    },
    "4-2": {
        "title": "否定感知危险信号匹配",
        "path": ROOT / "ai-service" / "app" / "agents" / "danger.py",
        "start": 59,
        "end": 79,
    },
    "4-3": {
        "title": "信息充分性判断与追问生成",
        "path": ROOT / "ai-service" / "app" / "agents" / "followup.py",
        "start": 7,
        "end": 42,
    },
    "4-4": {
        "title": "向量与词法融合检索",
        "path": ROOT / "ai-service" / "app" / "rag" / "retriever.py",
        "start": 115,
        "end": 154,
    },
    "4-5": {
        "title": "NDJSON事件封装与递增序号",
        "path": ROOT / "ai-service" / "app" / "events.py",
        "start": 69,
        "end": 107,
    },
    "4-6": {
        "title": "知识入库载荷清洗",
        "path": ROOT / "backend" / "src" / "main" / "java" / "com" / "medpilot" / "knowledge" / "KnowledgeController.java",
        "start": 194,
        "end": 203,
    },
}


FIGURE_LIST = [
    ("3-1", "系统总体架构图"),
    ("3-2", "多智能体协同分诊流程图"),
    ("3-3", "数据库E-R图"),
    *[(number, title) for number, (title, _) in FIGURES.items()],
    ("5-1", "各类危险信号召回率"),
    ("5-2", "各科室分诊F1值"),
    ("5-3", "各风险等级F1值"),
]


TABLE_LIST = [
    ("2-1", "系统角色及职责"),
    ("2-2", "功能需求"),
    ("2-3", "非功能需求"),
    ("2-4", "关键技术及用途"),
    ("3-1", "多智能体节点职责与输入输出"),
    ("3-2", "核心接口设计"),
    ("3-3", "用户表(users)"),
    ("3-4", "问诊记录表(consultation_records)"),
    ("3-5", "Trace表(consultation_traces)"),
    ("3-6", "健康档案表(health_profiles)"),
    ("3-7", "知识文档表(knowledge_documents)"),
    ("3-8", "临床复核表(clinical_reviews)"),
    ("3-9", "审计日志表(audit_logs)"),
    ("3-10", "角色权限矩阵"),
    ("5-1", "主要测试环境配置"),
    ("5-2", "扩展工程测试集分布"),
    ("5-3", "核心功能测试用例"),
]


def set_run_font(run, latin: str, east_asia: str, size: float, bold: bool | None = None) -> None:
    run.font.name = latin
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def insert_after(cursor, element):
    cursor.addnext(element)
    return element


def find_paragraph(doc: Document, needle: str):
    matches = [p for p in doc.paragraphs if needle in p.text]
    if len(matches) != 1:
        raise RuntimeError(f"expected one paragraph containing {needle!r}, found {len(matches)}")
    return matches[0]


def exact_paragraph(doc: Document, text: str):
    matches = [p for p in doc.paragraphs if p.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"expected one exact paragraph {text!r}, found {len(matches)}")
    return matches[0]


def add_body_paragraph(doc: Document, cursor, text: str, template):
    paragraph = doc.add_paragraph(style=template.style)
    if template._p.pPr is not None:
        if paragraph._p.pPr is not None:
            paragraph._p.remove(paragraph._p.pPr)
        paragraph._p.insert(0, deepcopy(template._p.pPr))
    run = paragraph.add_run(text)
    if template.runs and template.runs[0]._r.rPr is not None:
        run._r.insert(0, deepcopy(template.runs[0]._r.rPr))
    set_run_font(run, "Times New Roman", "宋体", 12)
    paragraph.paragraph_format.first_line_indent = Pt(21.6)
    insert_after(cursor, paragraph._p)
    return paragraph._p


def add_centered_caption(doc: Document, cursor, text: str, template, keep_with_next: bool):
    paragraph = doc.add_paragraph(style=template.style)
    if template._p.pPr is not None:
        if paragraph._p.pPr is not None:
            paragraph._p.remove(paragraph._p.pPr)
        paragraph._p.insert(0, deepcopy(template._p.pPr))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    run = paragraph.add_run(text)
    set_run_font(run, "Times New Roman", "宋体", 10.5)
    insert_after(cursor, paragraph._p)
    return paragraph._p


def source_excerpt(spec: dict) -> str:
    lines = spec["path"].read_text(encoding="utf-8").splitlines()
    excerpt = lines[spec["start"] - 1 : spec["end"]]
    return "\n".join(line.expandtabs(4).rstrip() for line in excerpt)


def set_cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_code_table(doc: Document, cursor, code: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "8352")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "BFBFBF")
    row_pr = table.rows[0]._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    row_pr.append(cant_split)
    cell = table.cell(0, 0)
    cell.width = Inches(5.8)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell, top=90, start=150, bottom=90, end=150)
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), "F2F2F2")
    tc_pr.append(shading)
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(10.2)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(code)
    set_run_font(run, "Consolas", "宋体", 8)
    insert_after(cursor, table._tbl)
    return table._tbl


def add_listing(doc: Document, cursor, number: str, caption_template):
    spec = LISTINGS[number]
    cursor = add_centered_caption(
        doc,
        cursor,
        f"程序清单{number} {spec['title']}",
        caption_template,
        keep_with_next=True,
    )
    return add_code_table(doc, cursor, source_excerpt(spec))


def picture_path(filename: str) -> Path:
    candidate = ASSET_DIR / filename
    if candidate.exists():
        return candidate
    return ROOT / "picture" / filename


def add_figure(doc: Document, cursor, number: str, caption_template, width: float = 5.8):
    title, filename = FIGURES[number]
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(picture_path(filename)), width=Inches(width))
    inline = doc.inline_shapes[-1]._inline
    description = f"图{number} {title}"
    inline.docPr.set("title", description)
    inline.docPr.set("descr", description)
    cursor = insert_after(cursor, paragraph._p)
    return add_centered_caption(
        doc,
        cursor,
        description,
        caption_template,
        keep_with_next=False,
    )


def prepare_assets() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    source = ROOT / "picture" / "26_用户权限管理.png"
    target = ASSET_DIR / "26_用户权限管理_裁剪.png"
    with Image.open(source) as image:
        crop_bottom = min(1020, image.height)
        image.crop((0, 0, image.width, crop_bottom)).save(target, optimize=True)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    for start in paragraph._p.findall(qn("w:bookmarkStart")):
        if start.get(qn("w:name")) == name:
            return
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    position = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(position, start)
    paragraph._p.append(end)


def max_bookmark_id(doc: Document) -> int:
    values = []
    for start in doc.element.body.iter(qn("w:bookmarkStart")):
        value = start.get(qn("w:id"))
        if value and value.isdigit():
            values.append(int(value))
    return max(values, default=0)


def add_pageref_field(paragraph, bookmark: str) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), f" PAGEREF {bookmark} \\h ")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "21")
    r_pr.append(r_fonts)
    r_pr.append(size)
    run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = "0"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def set_list_cell_text(cell, text: str, align: WD_ALIGN_PARAGRAPH) -> None:
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_run_font(run, "Times New Roman", "宋体", 10.5)


def populate_list_table(table, rows, prefix: str) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for number, title in rows:
        row = table.add_row()
        row.height = None
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=55, start=90, bottom=55, end=90)
        set_list_cell_text(row.cells[0], f"{prefix}{number}", WD_ALIGN_PARAGRAPH.CENTER)
        set_list_cell_text(row.cells[1], title, WD_ALIGN_PARAGRAPH.LEFT)
        page_paragraph = row.cells[2].paragraphs[0]
        clear_paragraph(page_paragraph)
        page_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        page_paragraph.paragraph_format.first_line_indent = Pt(0)
        page_paragraph.paragraph_format.space_before = Pt(0)
        page_paragraph.paragraph_format.space_after = Pt(0)
        add_pageref_field(page_paragraph, f"{prefix.lower()}_{number.replace('-', '_')}")
        tr_pr = row._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))


def add_caption_bookmarks_and_lists(doc: Document) -> None:
    next_id = max_bookmark_id(doc) + 1
    for prefix, rows in (("图", FIGURE_LIST), ("表", TABLE_LIST)):
        for number, title in rows:
            caption = exact_paragraph(doc, f"{prefix}{number} {title}")
            bookmark = f"{prefix.lower()}_{number.replace('-', '_')}"
            add_bookmark(caption, bookmark, next_id)
            next_id += 1
    if len(doc.tables) < 2:
        raise RuntimeError("figure/table list tables were not found")
    populate_list_table(doc.tables[0], FIGURE_LIST, "图")
    populate_list_table(doc.tables[1], TABLE_LIST, "表")


def enable_field_updates(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def build() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    prepare_assets()
    doc = Document(str(SOURCE))
    original_inline_shapes = len(doc.inline_shapes)
    body_template = find_paragraph(doc, "路由守卫只负责改善用户体验")
    caption_template = exact_paragraph(doc, "图3-1 系统总体架构图")

    anchor = find_paragraph(doc, "路由守卫只负责改善用户体验")
    cursor = add_body_paragraph(
        doc,
        anchor._p,
        "为使权限控制规则与页面角色划分能够相互核对，程序清单4-1截取了知识库、监控和审计资源的服务端授权配置。",
        body_template,
    )
    cursor = add_listing(doc, cursor, "4-1", caption_template)
    cursor = add_body_paragraph(
        doc,
        cursor,
        "代码按HTTP方法和资源路径区分写入、读取与审计权限，所有未单独放行的请求仍要求完成认证。对应的用户权限管理页面如图4-1所示。",
        body_template,
    )
    cursor = add_figure(doc, cursor, "4-1", caption_template)
    add_body_paragraph(
        doc,
        cursor,
        "图4-1集中展示系统管理员、患者用户、知识编辑、医学审核员、医生顾问和审计员等角色账号，便于管理员核对账号状态与职责分配；具体访问是否允许仍由后端授权规则进行最终校验。",
        body_template,
    )

    anchor = find_paragraph(doc, "减少了模型或索引异常对高风险提示的影响")
    cursor = add_body_paragraph(
        doc,
        anchor._p,
        "危险信号处理的关键在于区分用户确认出现的症状与明确否认的症状，程序清单4-2给出了否定感知匹配逻辑。",
        body_template,
    )
    cursor = add_listing(doc, cursor, "4-2", caption_template)
    cursor = add_body_paragraph(
        doc,
        cursor,
        "匹配器在每个语句片段内计算直接否定和连接词继承否定，仅将未被否定的规范危险信号加入结果集合。高风险规则的运行结果如图4-2所示。",
        body_template,
    )
    cursor = add_figure(doc, cursor, "4-2", caption_template)
    add_body_paragraph(
        doc,
        cursor,
        "图4-2显示胸痛触发安全快速通道后，界面同时给出高风险标识、建议科室和立即就医时效，并明确该路径未执行知识检索。",
        body_template,
    )

    anchor = find_paragraph(doc, "系统不把等待追问的会话写成低风险分诊结果")
    cursor = add_body_paragraph(
        doc,
        anchor._p,
        "程序清单4-3展示了信息充分性判断与追问生成逻辑。该逻辑先判断症状数量、持续时间和严重程度，再按缺失字段优先级生成一个追问。",
        body_template,
    )
    cursor = add_listing(doc, cursor, "4-3", caption_template)
    cursor = add_body_paragraph(
        doc,
        cursor,
        "主动追问在运行界面中的交互状态如图4-3所示。",
        body_template,
    )
    cursor = add_figure(doc, cursor, "4-3", caption_template)
    add_body_paragraph(
        doc,
        cursor,
        "图4-3中，用户仅输入咳嗽后，系统询问症状持续时间，知识检索和辅助分诊节点保持等待，避免在信息不足时提前输出科室结论。",
        body_template,
    )

    anchor = find_paragraph(doc, "解释文本提示补充信息或线下分诊")
    cursor = add_body_paragraph(
        doc,
        anchor._p,
        "为兼顾医学语义相近和短中文症状的字面命中，检索器采用向量分数与词法重合度融合排序，核心实现如程序清单4-4所示。",
        body_template,
    )
    cursor = add_listing(doc, cursor, "4-4", caption_template)
    cursor = add_body_paragraph(
        doc,
        cursor,
        "检索器先扩大向量候选池，再按0.8和0.2的权重融合向量分数与词法分数，并在阈值过滤后返回Top-K证据。对应结果页面如图4-4所示。",
        body_template,
    )
    cursor = add_figure(doc, cursor, "4-4", caption_template)
    add_body_paragraph(
        doc,
        cursor,
        "图4-4同时展示呼吸内科建议、风险等级、就医时效与三条医学来源，体现候选证据既参与科室评分，也作为用户查看判定依据的入口。",
        body_template,
    )

    anchor = find_paragraph(doc, "流提前结束则生成INCOMPLETE_STREAM错误")
    cursor = add_body_paragraph(
        doc,
        anchor._p,
        "程序清单4-5给出了统一事件封装逻辑。事件发射器为每条事件写入协议版本、Trace、会话、递增序号、阶段状态和节点数据，使前端进度展示与后端持久化使用同一事件依据。",
        body_template,
    )
    cursor = add_listing(doc, cursor, "4-5", caption_template)
    cursor = add_body_paragraph(
        doc,
        cursor,
        "基于该事件序列，前端能够把多智能体协同过程转换为连续的节点进度，运行状态如图4-5所示。",
        body_template,
    )
    cursor = add_figure(doc, cursor, "4-5", caption_template)
    add_body_paragraph(
        doc,
        cursor,
        "图4-5将安全筛查、信息采集、知识检索、辅助分诊和建议生成依次呈现，并同步显示当前完成比例与会话内容，便于用户理解系统仍处于分析过程。",
        body_template,
    )

    anchor = find_paragraph(doc, "避免用户误以为系统已完成临床诊断")
    cursor = add_body_paragraph(
        doc,
        anchor._p,
        "问诊记录详情不仅保存最终文本，还把风险等级、就医时效、检索支持度和证据关系进行可视化，具体页面如图4-6所示。",
        body_template,
    )
    cursor = add_figure(doc, cursor, "4-6", caption_template, width=5.45)
    add_body_paragraph(
        doc,
        cursor,
        "图4-6把症状节点、检索依据与科室、风险和时效结论连接起来，同时在支持度区域标注该数值不代表临床诊断结论，从展示层面保留辅助分诊边界。",
        body_template,
    )

    anchor = find_paragraph(doc, "具体分片和向量处理由AI服务完成")
    cursor = add_body_paragraph(
        doc,
        anchor._p,
        "知识文档进入AI服务前，业务后端会清洗审核字段并重新计算摘要，关键实现如程序清单4-6所示。",
        body_template,
    )
    cursor = add_listing(doc, cursor, "4-6", caption_template)
    add_body_paragraph(
        doc,
        cursor,
        "该处理移除客户端提交的审核人、审核时间和摘要，强制新文档进入pending状态，从入口处防止前端伪造审核结果或内容校验值。",
        body_template,
    )

    anchor = find_paragraph(doc, "避免“已向量化”被误认为“已批准使用”")
    cursor = add_body_paragraph(
        doc,
        anchor._p,
        "医学知识库页面将文档统计、解析与向量状态、索引版本及激活动作集中展示，运行界面如图4-7所示。",
        body_template,
    )
    cursor = add_figure(doc, cursor, "4-7", caption_template)
    add_body_paragraph(
        doc,
        cursor,
        "图4-7能够同时查看知识文档数量、切片数量、科室覆盖、可用索引版本和当前激活状态，使文档审核、版本构建与版本切换形成可核对的管理流程。",
        body_template,
    )

    anchor = find_paragraph(doc, "不把样本量或耗时直接解释为医疗质量指标")
    cursor = add_body_paragraph(
        doc,
        anchor._p,
        "监控模块先提供系统与Trace总览，展示模型服务、活动会话、知识文档、向量索引及链路统计，页面如图4-8所示。",
        body_template,
    )
    cursor = add_figure(doc, cursor, "4-8", caption_template)
    cursor = add_body_paragraph(
        doc,
        cursor,
        "图4-8用于定位服务可用性和链路运行状态。选择具体Trace后，系统进一步显示事件级调用链路，如图4-9所示。",
        body_template,
    )
    cursor = add_figure(doc, cursor, "4-9", caption_template)
    add_body_paragraph(
        doc,
        cursor,
        "图4-9按事件序号列出安全筛查、症状采集、主动追问和流程完成状态，并保留各事件耗时及输出摘要，便于从单次会话追踪条件路由结果。",
        body_template,
    )

    anchor = find_paragraph(doc, "其他角色不能读取全量日志")
    cursor = add_body_paragraph(
        doc,
        anchor._p,
        "审计日志页面支持按操作者和响应状态筛选，并以时间、角色、动作、结果和耗时为主要字段展示访问记录，如图4-10所示。",
        body_template,
    )
    cursor = add_figure(doc, cursor, "4-10", caption_template)
    add_body_paragraph(
        doc,
        cursor,
        "图4-10中的记录覆盖用户管理、监控统计与Trace查询等操作；页面只保留访问边界、结果和耗时，不保存请求正文、令牌或医疗文本。",
        body_template,
    )

    add_caption_bookmarks_and_lists(doc)
    enable_field_updates(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)

    check = Document(str(OUTPUT))
    expected_shapes = original_inline_shapes + len(FIGURES)
    if len(check.inline_shapes) != expected_shapes:
        raise RuntimeError(f"inline shape count {len(check.inline_shapes)} != {expected_shapes}")
    if len(check.tables[0].rows) != len(FIGURE_LIST) + 1:
        raise RuntimeError("figure list row count mismatch")
    if len(check.tables[1].rows) != len(TABLE_LIST) + 1:
        raise RuntimeError("table list row count mismatch")
    for number, (title, _) in FIGURES.items():
        exact_paragraph(check, f"图{number} {title}")
    for number, spec in LISTINGS.items():
        exact_paragraph(check, f"程序清单{number} {spec['title']}")
    print(
        f"saved={OUTPUT}; figures_added={len(FIGURES)}; listings_added={len(LISTINGS)}; "
        f"inline_shapes={len(check.inline_shapes)}; tables={len(check.tables)}"
    )


if __name__ == "__main__":
    build()
