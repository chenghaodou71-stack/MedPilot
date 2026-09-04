"""Build the thesis using the user's exact chapter directory.

The module imports only layout helpers from ``build_new_thesis.py``. No prior
thesis prose is read or appended.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_new_thesis as base  # noqa: E402


OUT = base.ROOT / "outputs" / "MedPilot_毕业设计论文_按目录全新正文版.docx"


def h(doc, text, level=1):
    return base.heading(doc, text, level)


def p(doc, text, **kwargs):
    return base.paragraph(doc, text, **kwargs)


def figure(doc, number, caption, filename):
    before = len(doc.inline_shapes)
    result = base.figure(doc, number, caption, filename)
    if len(doc.inline_shapes) > before:
        # Keep figure descriptions in the OOXML for screen readers and audits.
        doc_pr = doc.inline_shapes[-1]._inline.docPr
        description = f"图{number} {caption}"
        doc_pr.set("title", description)
        doc_pr.set("descr", description)
    return result


def table(doc, caption, headers, rows, widths):
    before = len(doc.tables)
    result = base.data_table(doc, caption, headers, rows, widths)
    if len(doc.tables) > before:
        # Mark the first row as a semantic header for assistive technologies.
        header_row = doc.tables[-1].rows[0]._tr
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        header_row.get_or_add_trPr().append(tbl_header)
    return result


def placeholder(doc, label):
    return base.placeholder(doc, label)


def add_front_title(doc, title):
    h(doc, title, 1)


def front_matter(doc):
    base.title_page(doc)
    base.statements(doc)

    add_front_title(doc, "摘  要")
    p(doc, "随着互联网医疗、医学知识库和大语言模型技术的发展，用户可以使用自然语言描述健康问题并获得初步的信息整理。然而，医疗健康咨询具有专业知识密集、症状描述不规范和安全风险不对称等特点，直接依赖大语言模型生成回答可能产生证据不可追溯、危险信号识别不稳定以及咨询与诊断边界模糊等问题。针对上述问题，本文设计并实现了基于多智能体协同与大语言模型的多专科医疗健康咨询及辅助分诊系统 MedPilot。")
    p(doc, "系统采用 Vue 3、Spring Boot 与 FastAPI 的分层架构，AI 服务使用 LangGraph 组织安全筛查、症状抽取、信息充分性判断、主动追问、医学知识检索、辅助分诊和回答编排等职责节点。危险信号筛查在大语言模型和向量检索调用之前执行，胸痛、呼吸困难、气促、咯血、便血、晕厥、意识不清和大出血等信号进入高风险快速通道；普通场景使用 bge-m3 生成查询向量，结合 FAISS 向量检索与字符 n-gram 词法匹配获得医学证据，再通过科室加权投票形成辅助分诊结果。证据不足时系统主动 abstain，建议补充信息或线下分诊。")
    p(doc, "系统还实现了 NDJSON 流式事件、问诊记录、授权健康档案上下文、医生复核队列、医学知识审核和索引版本治理、Trace 监控、审计日志以及基于医疗关系、院区和 MFA 的权限控制。测试数据采用人工构造或脱敏病例，软件测试结果仅用于说明工程行为和回归稳定性，不外推为临床准确率、患者获益或医生替代能力。")
    add_front_title(doc, "关键词")
    p(doc, "多智能体协同；大语言模型；医疗健康咨询；辅助分诊；检索增强生成；医学知识库；安全护栏", indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
    doc.add_page_break()

    add_front_title(doc, "Abstract")
    p(doc, "This thesis presents MedPilot, a multi-specialty medical health consultation and assisted triage system based on multi-agent collaboration and large language models. The system is designed as an evidence-grounded and rule-guarded assistant rather than a definitive diagnostic tool. A layered architecture consisting of Vue 3, Spring Boot, and FastAPI is implemented. A fixed LangGraph workflow coordinates safety screening, symptom extraction, information sufficiency checking, follow-up questioning, medical retrieval, assisted triage, and answer composition. Deterministic red-flag screening is executed before any language-model or embedding call. Ordinary cases use bge-m3 embeddings, FAISS retrieval, lexical overlap fusion, and evidence-weighted department voting. Insufficient evidence triggers an abstention result and an offline-care suggestion.", font="Times New Roman")
    p(doc, "The system also provides NDJSON streaming events, consultation records, consent-scoped health context, clinical review queues, knowledge-document governance, versioned index activation, trace monitoring, audit logging, and MFA- and care-relationship-aware authorization. Evaluation data are artificial or de-identified. The reported results validate software behavior and reproducibility only; they do not establish clinical accuracy, patient outcomes, or replacement of clinicians.", font="Times New Roman")
    add_front_title(doc, "Key Words")
    p(doc, "multi-agent collaboration; large language model; medical health consultation; assisted triage; retrieval-augmented generation; medical knowledge base; safety guardrail", indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, font="Times New Roman")
    doc.add_page_break()

    add_front_title(doc, "目 录")
    toc_items = [
        "摘  要", "关键词", "Abstract", "Key Words", "目 录", "图清单", "表清单",
        "第1章 绪论", "1.1 研究背景", "1.2 研究目的与意义", "1.3 国内外研究现状", "1.3.1 医疗健康咨询研究现状", "1.3.2 智能辅助分诊研究现状", "1.3.3 多智能体与大语言模型应用现状", "1.3.4 现有研究不足", "1.4 研究内容", "1.5 技术路线与论文组织", "1.6 本章小结",
        "第2章 系统分析与相关技术", "2.1 系统可行性分析", "2.1.1 技术可行性", "2.1.2 经济可行性", "2.1.3 操作可行性", "2.1.4 医疗安全与伦理可行性", "2.2 系统角色与业务流程分析", "2.3 系统功能需求分析", "2.3.1 患者智能问诊需求", "2.3.2 辅助分诊需求", "2.3.3 医生复核需求", "2.3.4 医学知识库管理需求", "2.3.5 监控、审计与治理需求", "2.4 系统非功能需求分析", "2.5 相关技术介绍", "2.5.1 大语言模型与 Ollama", "2.5.2 多智能体协同与 LangGraph", "2.5.3 RAG、Embedding 与 FAISS", "2.5.4 Vue 3、Spring Boot 与 FastAPI", "2.5.5 NDJSON 流式通信与 Trace", "2.6 本章小结",
        "第3章 系统总体设计", "3.1 系统设计目标与原则", "3.2 系统总体架构设计", "3.3 系统功能模块设计", "3.4 多智能体分诊工作流设计", "3.5 用户角色与权限设计", "3.6 医疗关系与医生复核流程设计", "3.7 数据库设计", "3.7.1 问诊记录相关数据表", "3.7.2 患者与医疗关系数据表", "3.7.3 医生复核数据表", "3.7.4 知识治理、审计与监控数据表", "3.8 接口与事件协议设计", "3.9 本章小结",
        "第4章 核心算法设计与实现", "4.1 系统辅助分诊总体算法流程", "4.2 多智能体协同算法", "4.2.1 智能体角色划分", "4.2.2 结构化状态与消息传递", "4.2.3 条件有向图路由机制", "4.2.4 主动追问与信息充分性判断", "4.3 医学知识检索增强生成算法", "4.3.1 医学知识库构建与文本切分", "4.3.2 医学文本向量化与 FAISS 索引", "4.3.3 向量检索与词法匹配融合", "4.3.4 基于证据的科室推荐", "4.3.5 引用可追溯与回答约束", "4.4 医疗安全与兜底机制", "4.4.1 危险信号识别与否定表达处理", "4.4.2 高风险快速通道", "4.4.3 证据不足时的 abstain 机制", "4.4.4 医生复核与人工接管机制", "4.5 算法复杂度与局限性分析", "4.6 本章小结",
        "第5章 系统详细设计与实现", "5.1 患者端智能问诊模块", "5.2 症状抽取与主动追问模块", "5.3 医学知识检索与辅助分诊模块", "5.4 问诊记录、健康档案与复诊提醒模块", "5.5 医生复核队列模块", "5.6 医学知识治理模块", "5.7 智能体运行监控与审计模块", "5.8 用户认证、MFA 与医疗数据安全模块", "5.9 本章小结",
        "第6章 系统测试与实验分析", "6.1 测试环境与测试数据", "6.2 系统功能测试", "6.2.1 用户登录与权限测试", "6.2.2 正常问诊流程测试", "6.2.3 信息不足追问测试", "6.2.4 高风险快速通道测试", "6.2.5 医生复核流程测试", "6.3 辅助分诊算法评测", "6.3.1 安全召回率", "6.3.2 科室 Macro-F1", "6.3.3 风险等级 Macro-F1", "6.3.4 错误率与响应延迟", "6.4 医学 RAG 检索评测", "6.4.1 Recall@K", "6.4.2 MRR", "6.4.3 引用可追溯率", "6.5 对比实验与消融实验", "6.6 性能、安全性与鲁棒性测试", "6.7 实验结果分析", "6.8 本章小结",
        "第7章 总结与展望", "7.1 研究工作总结", "7.2 系统主要贡献", "7.3 系统不足", "7.4 未来展望", "参考文献", "致谢", "附录"
    ]
    for item in toc_items:
        pp = doc.add_paragraph()
        pp.paragraph_format.left_indent = Pt(12 if item.startswith("第") else 30 if item[:1].isdigit() else 42)
        pp.paragraph_format.space_after = Pt(2)
        rr = pp.add_run(item)
        base.set_font(rr, "宋体", 10.5, bold=item.startswith("第"))
    doc.add_page_break()

    add_front_title(doc, "图清单")
    for item in [
        "图3-1 系统总体架构图", "图3-2 辅助分诊流程图", "图3-3 系统功能用例图", "图3-4 角色权限模型图", "图3-5 医生复核安全闸门图", "图3-6 数据库关系示意图", "图3-7 智能问诊时序图", "图3-8 医学 RAG 流程图", "图3-9 Trace 监控与审计流程图", "图3-10 系统部署架构图", "图5-1 至图5-5 系统界面截图（待用户替换）", "图6-1 至图6-3 测试场景截图（待用户替换）"
    ]:
        p(doc, item, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11)
    doc.add_page_break()

    add_front_title(doc, "表清单")
    for item in [
        "表2-1 系统可行性分析", "表2-2 系统功能需求", "表3-1 系统角色与权限", "表3-2 核心数据表", "表3-3 NDJSON 事件字段", "表6-1 测试环境与数据", "表6-2 系统功能测试用例", "表6-3 辅助分诊算法评测结果", "表6-4 RAG 检索评测状态", "表6-5 对比与消融实验设计", "表A-1 系统截图替换清单"
    ]:
        p(doc, item, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11)
    doc.add_page_break()


def chapter1(doc):
    h(doc, "第1章 绪论")
    h(doc, "1.1 研究背景", 2)
    p(doc, "在线医疗健康咨询能够降低用户获取健康信息的门槛，但用户输入往往是口语化、片段化且缺少持续时间和严重程度等关键信息。对胸痛、呼吸困难、意识不清等危险信号而言，系统如果只关注回答的语言流畅性而没有安全优先级，可能导致用户延误线下就医。")
    p(doc, "大语言模型在医学问答中具有较强的语言理解和组织能力，但仍存在事实错误、提示敏感、证据不可追溯和边界表达不清等风险。因此，医疗咨询系统需要将模型放入可解释、可审计和可人工接管的工作流，而不是把模型输出直接当作诊断结论。")
    h(doc, "1.2 研究目的与意义", 2)
    p(doc, "本文的研究目的，是设计一个面向多专科健康咨询的工程原型，验证多智能体协同、大语言模型和医学 RAG 能否在安全规则约束下完成症状整理、主动追问、证据检索和辅助分诊。")
    p(doc, "研究意义体现在三个方面：第一，将危险信号筛查前置并形成高风险快速通道；第二，将检索证据、科室推荐和回答引用绑定，提升结果的可追溯性；第三，通过医生复核、知识治理、权限控制和 Trace 监控建立人工接管与持续治理机制。")
    h(doc, "1.3 国内外研究现状", 2)
    h(doc, "1.3.1 医疗健康咨询研究现状", 3)
    p(doc, "现有在线症状检查器通常采用问答树、规则匹配或统计模型完成健康信息整理。Semigran 等人的审计研究指出，不同工具在自我诊断和分诊任务上的表现存在明显差异；系统综述也表明，医疗场景的安全证据和真实世界验证仍然不足。")
    h(doc, "1.3.2 智能辅助分诊研究现状", 3)
    p(doc, "智能辅助分诊研究关注就诊科室、风险等级和就医时效等输出。近年来，研究逐渐从单一准确率转向危险信号召回、可解释因素、人工复核和数据漂移等问题。本文将分诊定位为辅助建议，不输出确定性疾病名称和处方。")
    h(doc, "1.3.3 多智能体与大语言模型应用现状", 3)
    p(doc, "大语言模型能够支持医学文本抽取、对话追问和知识问答，多智能体研究则尝试通过任务拆分、角色协作和工具调用完成复杂任务。医疗场景更需要职责边界清晰的协同方式，避免多个模型自由生成导致责任不清。")
    h(doc, "1.3.4 现有研究不足", 3)
    p(doc, "现有研究的不足主要包括：临床验证样本有限；检索证据与最终回答的绑定不充分；高风险规则与生成式模型的优先级关系不明确；模型、知识库、权限和审计治理较少被放在同一个系统闭环中。")
    h(doc, "1.4 研究内容", 2)
    p(doc, "本文完成系统需求分析、总体架构设计、数据库和接口设计，多智能体分诊工作流实现，危险信号规则、主动追问、RAG 检索和证据加权分诊算法实现，医生复核和知识治理模块实现，以及自动化测试和构造病例评测。")
    h(doc, "1.5 技术路线与论文组织", 2)
    p(doc, "技术路线为：需求分析与安全边界确定，分层架构和角色权限设计，LangGraph 工作流和核心算法实现，Spring Boot 业务接口与 MySQL 持久化，Vue 3 交互和 NDJSON 流展示，最后通过单元测试、接口测试、浏览器关键路径和离线构造病例回归进行验证。全文按照系统分析、总体设计、算法实现、详细实现、测试分析和总结展望展开。")
    h(doc, "1.6 本章小结", 2)
    p(doc, "本章说明了医疗健康咨询和辅助分诊的研究背景、研究目的、国内外研究进展与不足，明确了本文的工程研究范围和论文组织结构，为后续系统设计和算法实现奠定基础。")


def chapter2(doc):
    h(doc, "第2章 系统分析与相关技术")
    h(doc, "2.1 系统可行性分析", 2)
    h(doc, "2.1.1 技术可行性", 3)
    p(doc, "项目采用的 Vue 3、Spring Boot、FastAPI、LangGraph、Ollama、bge-m3 和 FAISS 均可在本地开发环境部署。前后端通过明确接口通信，AI 节点使用依赖注入的模型函数，便于离线测试和替换模型，因此技术上具备实现条件。")
    h(doc, "2.1.2 经济可行性", 3)
    p(doc, "系统采用开源框架和本地推理方式，开发阶段不依赖商业云端推理服务，主要成本为普通开发计算机、存储和维护时间。对于本科毕业设计原型而言，投入可控；正式生产仍需评估 GPU、运维和合规成本。")
    h(doc, "2.1.3 操作可行性", 3)
    p(doc, "患者端以对话式输入降低使用门槛，医生端提供复核队列和详情页，管理员和审计员分别进入治理和只读监控页面。系统提供统一启动脚本和预检脚本，能够支持开发人员按固定步骤运行。")
    h(doc, "2.1.4 医疗安全与伦理可行性", 3)
    p(doc, "系统明确不提供确定性诊断、处方和急救替代服务；危险信号优先进入线下就医建议；证据不足时自动暂缓；医生复核需要员工身份、医疗关系和 MFA。测试病例使用人工构造或脱敏数据，因此满足毕业设计原型阶段的安全和伦理约束。")
    table(doc, "表2-1 系统可行性分析", ["维度", "分析结论", "主要依据"], [("技术", "可行", "开源技术栈、本地模型和可替换节点"), ("经济", "可行", "开源组件和本地推理降低服务成本"), ("操作", "可行", "患者、医生和治理角色界面分离"), ("医疗安全与伦理", "有条件可行", "规则护栏、人工复核、脱敏病例和明确边界")], [35, 35, 85])
    h(doc, "2.2 系统角色与业务流程分析", 2)
    p(doc, "系统包含患者 USER、医生 DOCTOR、复核员 REVIEWER、知识编辑器 KNOWLEDGE_EDITOR、管理员 ADMIN 和审计员 AUDITOR 六类角色。患者发起咨询后，系统先进行安全筛查；未命中危险信号时抽取症状并判断信息是否充分，必要时主动追问，信息充分后检索知识并完成辅助分诊。高风险结果可以进入医生复核队列，原始 AI 结果不被覆盖。")
    figure(doc, "2-1", "辅助分诊流程图", "02_辅助分诊流程图.png")
    h(doc, "2.3 系统功能需求分析", 2)
    h(doc, "2.3.1 患者智能问诊需求", 3)
    p(doc, "患者可以创建问诊会话，输入自然语言症状，查看结构化症状、主动追问、分诊建议、风险等级、就医时效、证据引用和安全边界，并访问本人问诊记录和健康档案。")
    h(doc, "2.3.2 辅助分诊需求", 3)
    p(doc, "系统需要区分高风险和普通路径。高风险路径优先执行规则并给出尽快线下就医建议；普通路径根据证据科室加权投票输出建议科室；证据不足时应返回 abstain，而不是强行推荐。")
    h(doc, "2.3.3 医生复核需求", 3)
    p(doc, "医生或复核员可以查看符合医疗关系和院区条件的复核任务，领取任务、查看原始 AI 结果和证据，并提交确认、修改、退回或急诊升级决定。系统禁止复核员审核本人产生的问诊记录。")
    h(doc, "2.3.4 医学知识库管理需求", 3)
    p(doc, "知识编辑器负责文档入库和索引构建，医生或复核员负责审核和激活，管理员负责治理和回滚。文档需要记录来源、版本、审核状态、有效期和索引版本。")
    h(doc, "2.3.5 监控、审计与治理需求", 3)
    p(doc, "系统应记录节点开始、完成和错误事件，保存成功和失败 Trace，支持按终止阶段、错误码、节点和时间过滤，并记录用户、资源、动作和时间戳等审计信息。")
    h(doc, "2.4 系统非功能需求分析", 2)
    p(doc, "系统的非功能需求包括安全性、可用性、可靠性、可维护性、可追溯性和可扩展性。安全性要求权限失败关闭、敏感字段加密和 CSRF 防护；可靠性要求事件序列连续、错误链路持久化和索引原子切换；可维护性要求节点职责单一、接口契约明确和测试可注入。")
    h(doc, "2.5 相关技术介绍", 2)
    h(doc, "2.5.1 大语言模型与 Ollama", 3)
    p(doc, "Ollama 为本地大语言模型提供统一调用接口，项目使用 qwen2.5:7b 承担症状抽取、追问和回答编排。模型调用被限制在指定节点，危险信号规则和输出结构校验不依赖模型判断。")
    h(doc, "2.5.2 多智能体协同与 LangGraph", 3)
    p(doc, "LangGraph 以 StateGraph 描述节点、状态和条件边。系统将安全筛查、抽取、追问、检索、分诊和编排作为有明确输入输出的协同节点，使流程能够流式观测和单元测试。")
    h(doc, "2.5.3 RAG、Embedding 与 FAISS", 3)
    p(doc, "RAG 先检索外部医学证据再生成回答。bge-m3 将文本映射为向量，FAISS IndexFlatIP 在归一化向量空间中执行内积检索，系统再融合字符 n-gram 词法重合度，提高短中文症状的匹配稳定性。")
    h(doc, "2.5.4 Vue 3、Spring Boot 与 FastAPI", 3)
    p(doc, "Vue 3 负责界面和流式状态展示；Spring Boot 负责认证、授权、请求代理、事件校验和数据库持久化；FastAPI 负责 AI 工作流和 NDJSON 输出。浏览器不直接访问 AI 服务。")
    h(doc, "2.5.5 NDJSON 流式通信与 Trace", 3)
    p(doc, "NDJSON 以一行一个 JSON 事件传递节点状态，适合在问诊过程中展示实时进度。Trace 保存事件快照、引用快照、终止阶段、失败码和节点耗时，用于调试和审计。")
    h(doc, "2.6 本章小结", 2)
    p(doc, "本章从技术、经济、操作以及医疗安全伦理角度分析了系统可行性，明确了六类角色和五类功能需求，并介绍了大语言模型、LangGraph、RAG、FAISS、前后端分层及 Trace 通信等技术基础。")


def chapter3(doc):
    h(doc, "第3章 系统总体设计")
    h(doc, "3.1 系统设计目标与原则", 2)
    p(doc, "系统设计目标是形成一个可运行、可解释、可复核和可治理的多专科医疗健康咨询原型。设计遵循安全优先、证据增强、职责分离、最小权限、失败可追溯和结果边界明确六项原则。")
    h(doc, "3.2 系统总体架构设计", 2)
    p(doc, "系统由用户端、业务后端、AI 服务和数据与知识基础设施组成。用户端使用 Vue 3；Spring Boot 负责业务 API、JWT、审计和医疗关系访问；FastAPI 负责 LangGraph 工作流；MySQL 保存业务数据，Ollama 提供本地大语言模型，FAISS 保存医学检索索引。")
    figure(doc, "3-1", "系统总体架构图", "01_系统总体架构图.png")
    h(doc, "3.3 系统功能模块设计", 2)
    table(doc, "表3-1 系统功能模块", ["模块", "主要功能", "核心输出"], [("患者问诊", "输入症状、追问、查看结果", "会话、回答、证据"), ("辅助分诊", "风险筛查、科室推荐、就医时效", "TriageResult"), ("医生复核", "领取、确认、修改、升级", "复核决定"), ("知识治理", "入库、审核、构建、激活", "索引版本"), ("监控审计", "Trace、统计、日志", "运行轨迹和审计记录")], [35, 75, 45])
    h(doc, "3.4 多智能体分诊工作流设计", 2)
    p(doc, "工作流的入口为安全筛查节点。未命中危险信号时进入症状抽取，抽取结果经信息充分性判断后分流到主动追问或医学检索；检索结果进入辅助分诊，再由回答编排节点生成带引用和安全边界的结果。")
    figure(doc, "3-2", "多智能体分诊工作流", "02_辅助分诊流程图.png")
    figure(doc, "3-3", "系统功能用例图", "04_系统功能用例图.png")
    h(doc, "3.5 用户角色与权限设计", 2)
    p(doc, "患者权限以本人 user_id 和 session_id 为边界；知识编辑器只能提交和构建知识；医生和复核员承担临床复核；管理员承担用户和系统治理；审计员只读访问监控和审计接口。ADMIN 不自动获得临床复核权限，医生复核还需要医院员工档案和 MFA。")
    figure(doc, "3-4", "角色权限模型图", "08_角色权限模型图.png")
    h(doc, "3.6 医疗关系与医生复核流程设计", 2)
    p(doc, "ClinicalReviewService 对列表、详情、领取和决定操作统一执行角色、员工档案、员工编号、MFA 和医疗关系检查。复核任务只向当前医疗关系和院区匹配的临床人员展示，复核员不得处理自己创建的问诊结果；原始 AI 结果保留，最终决定另行记录。")
    figure(doc, "3-5", "医生复核安全闸门图", "07_医生复核安全闸门图.png")
    h(doc, "3.7 数据库设计", 2)
    h(doc, "3.7.1 问诊记录相关数据表", 3)
    p(doc, "问诊会话表 consultation_sessions 保存用户、轮次和状态；问诊消息表 consultation_messages 保存角色和内容；问诊记录表 consultation_records 保存最终结构化分诊、引用快照和 Trace 关联。")
    h(doc, "3.7.2 患者与医疗关系数据表", 3)
    p(doc, "users 保存账号、角色、员工编号和 MFA 状态；health_profiles 保存过敏史、既往史、用药和备注；patient_care_relationships 保存患者、临床人员、组织和院区关系。")
    h(doc, "3.7.3 医生复核数据表", 3)
    p(doc, "clinical_reviews 保存复核编号、问诊记录编号、状态、领取人、决定、最终科室、风险、时效、理由和时间戳，用于实现原始结果和人工决定分离。")
    h(doc, "3.7.4 知识治理、审计与监控数据表", 3)
    p(doc, "knowledge_documents 保存来源和审核状态；索引版本表保存构建清单与激活版本；consultation_traces 保存节点事件和引用快照；audit_logs 保存操作者、动作、资源和时间。")
    figure(doc, "3-6", "数据库关系示意图", "03_数据库关系示意图.png")
    table(doc, "表3-2 核心数据表", ["数据表", "关键字段", "用途"], [("consultation_records", "user_id、session_id、triage_json", "保存问诊结果"), ("clinical_reviews", "review_id、status、decision", "保存医生复核"), ("consultation_traces", "trace_id、events_json、failure_code", "保存运行轨迹"), ("knowledge_documents", "source、review_status、expires_at", "知识治理"), ("audit_logs", "actor、action、resource", "审计追踪")], [45, 65, 45])
    h(doc, "3.8 接口与事件协议设计", 2)
    p(doc, "接口按认证、问诊、记录、临床复核、知识和监控分类。问诊接口返回 NDJSON 事件，事件至少包含 protocol_version、sequence、type、node、status、phase、elapsed_ms 和 data 字段。前端按照 sequence 严格递增合并节点状态，后端只接受合法终止事件。")
    figure(doc, "3-7", "智能问诊时序图", "05_智能问诊时序图.png")
    figure(doc, "3-8", "Trace 监控与审计流程图", "10_Trace监控与审计流程图.png")
    h(doc, "3.9 本章小结", 2)
    p(doc, "本章完成了系统总体架构、功能模块、协同工作流、权限与医疗关系、数据库和事件协议设计，为第4章算法实现和第5章详细模块实现提供了结构基础。")


def chapter4(doc):
    h(doc, "第4章 核心算法设计与实现")
    h(doc, "4.1 系统辅助分诊总体算法流程", 2)
    p(doc, "一次咨询的总体计算流程可表示为：输入文本 X 经过安全筛查 R；若 R 命中危险信号则直接输出高风险结果；否则由抽取节点得到结构化症状 S，并判断信息充分性 Q；Q 不充分时输出追问问题，否则检索证据 E；最后由分诊函数 f(S,E) 得到 TriageResult，再由回答编排函数生成自然语言结果 A。")
    p(doc, "该流程的关键是把安全筛查置于模型调用之前，把模型输出限制在结构化字段和回答编排节点内，把证据不足显式建模为 abstain 状态。")
    h(doc, "4.2 多智能体协同算法", 2)
    h(doc, "4.2.1 智能体角色划分", 3)
    p(doc, "系统将协同职责划分为安全筛查智能体、症状抽取智能体、主动追问智能体、医学检索智能体、辅助分诊智能体和回答编排智能体。前四类可以使用模型或检索工具，安全筛查和分诊规则具有确定性优先级。")
    h(doc, "4.2.2 结构化状态与消息传递", 3)
    p(doc, "ConsultState 保存 text、history、health_context、symptoms、followup、evidence、triage、answer 和 event_sink 等字段。每个节点只写入自己负责的字段，后续节点从共享状态读取，避免使用难以校验的自由格式长文本传递结果。")
    h(doc, "4.2.3 条件有向图路由机制", 3)
    p(doc, "LangGraph 图从 START 进入 extract，再由 _route_after_extract 判断 is_sufficient。充分时路由至 retrieve→classify→compose；不充分时路由至 ask_followup→END。安全快速通道在图执行前完成，并直接生成 classify 和 compose 事件。")
    h(doc, "4.2.4 主动追问与信息充分性判断", 3)
    p(doc, "is_sufficient 根据症状数量、持续时间、严重程度和伴随症状等字段判断信息是否足够。build_followup 从缺失字段中选择一个优先问题，避免一次返回过多问题，下一轮输入和 history 再进入抽取节点。")
    figure(doc, "4-1", "智能问诊时序图", "05_智能问诊时序图.png")
    h(doc, "4.3 医学知识检索增强生成算法", 2)
    h(doc, "4.3.1 医学知识库构建与文本切分", 3)
    p(doc, "知识文档需要记录标题、来源、机构、专科、版本、审核状态和有效期。构建阶段只选取审核通过且未过期的文档，按句子和语义边界切分为 chunk，并生成 doc_id 与 chunk_id。")
    h(doc, "4.3.2 医学文本向量化与 FAISS 索引", 3)
    p(doc, "bge-m3 为每个知识片段生成向量，FAISS IndexFlatIP 保存归一化向量。索引版本与文档清单绑定，构建失败时不替换当前 active_index，只有显式激活成功的新版本才参与检索。")
    h(doc, "4.3.3 向量检索与词法匹配融合", 3)
    p(doc, "系统将原始文本与结构化症状拼接并进行同义词扩展，先取得较宽的向量候选集，再计算字符 n-gram 重合度。综合分数为 S=0.8S_dense+0.2S_lexical，低于 MEDPILOT_RAG_MIN_SCORE 的候选被过滤。")
    h(doc, "4.3.4 基于证据的科室推荐", 3)
    p(doc, "对检索结果按 department 累加 score，选择总权重最高的科室。support_score 等于最高科室权重除以总权重，并限制在系统支持范围内；如果没有有效证据，返回全科/线下分诊建议并设置 abstained=true。")
    h(doc, "4.3.5 引用可追溯与回答约束", 3)
    p(doc, "每个 RankedEvidence 都携带 citation_id、doc_id、chunk_id、source、quote、score、index_version 和审核元数据。回答编排只引用检索返回的证据，并附带安全边界；支持分数不解释为临床概率。")
    figure(doc, "4-2", "医学 RAG 流程图", "06_医学RAG流程图.png")
    h(doc, "4.4 医疗安全与兜底机制", 2)
    h(doc, "4.4.1 危险信号识别与否定表达处理", 3)
    p(doc, "danger.py 将危险信号词和否定表达结合处理，区分“没有胸痛”与“出现胸痛”，避免简单字符串包含造成误报。当前规则覆盖胸痛、呼吸困难、气促、咯血、便血、晕厥、意识不清和大出血等场景。")
    h(doc, "4.4.2 高风险快速通道", 3)
    p(doc, "screen_for_emergency 在任何 LLM 和 embedding 调用之前运行。命中后构造高风险 TriageResult，输出急诊或尽快就医建议，并通过 NDJSON 发出 safety_screen、classify、compose 和 done 事件。")
    h(doc, "4.4.3 证据不足时的 abstain 机制", 3)
    p(doc, "普通路径没有足够的支持证据时，系统不强行选择专科，将 abstained 设置为 true，support_score 设为0，并提示补充信息或线下分诊。该机制用于控制不确定性，不能保证临床安全。")
    h(doc, "4.4.4 医生复核与人工接管机制", 3)
    p(doc, "高风险或治理规则要求的问诊结果可以进入临床复核队列。医生和复核员在医疗关系、院区、员工档案和 MFA 检查通过后才能领取和决定，原始 AI 结果不可变，人工决定单独留痕。")
    figure(doc, "4-3", "医生复核安全闸门图", "07_医生复核安全闸门图.png")
    h(doc, "4.5 算法复杂度与局限性分析", 2)
    p(doc, "安全筛查和规则分诊的复杂度与规则数量和输入长度近似线性；FAISS IndexFlatIP 的精确检索复杂度约为 O(Nd)，N 为知识向量数量，d 为向量维度；候选排序复杂度约为 O(k log k)。大语言模型推理耗时取决于模型规模和硬件，无法仅由代码复杂度表示。")
    p(doc, "算法局限性包括：规则覆盖有限，否定表达和复杂上下文仍可能产生边界错误；词法融合参数需要在更多标注数据上校准；科室投票依赖知识库来源和专科标签；当前 support_score 不是临床校准概率；没有临床盲评和前瞻验证。")
    h(doc, "4.6 本章小结", 2)
    p(doc, "本章详细设计并实现了多智能体协同、RAG 检索增强、证据加权分诊、危险信号快速通道、abstain 兜底和医生人工接管机制，并分析了算法复杂度与研究边界。")


def chapter5(doc):
    h(doc, "第5章 系统详细设计与实现")
    h(doc, "5.1 患者端智能问诊模块", 2)
    p(doc, "患者端提供问诊入口、输入框、会话历史、节点进度、回答内容、风险等级、建议科室、就医时效和证据引用。前端通过 Spring Boot 建立流式请求，接收 NDJSON 后更新节点状态，完成事件后再刷新问诊记录。")
    placeholder(doc, "图5-1 患者端登录与智能问诊页面截图（待用户运行系统后替换）")
    h(doc, "5.2 症状抽取与主动追问模块", 2)
    p(doc, "症状抽取模块将自然语言转换为 StructuredSymptoms，保留原始文本并提取持续时间、严重程度、伴随症状、既往史、用药和危险信号。信息不足时主动追问模块只提出当前最重要的问题，下一轮回答与 history 一起参与判断。")
    placeholder(doc, "图5-2 信息不足主动追问页面截图（待用户运行系统后替换）")
    h(doc, "5.3 医学知识检索与辅助分诊模块", 2)
    p(doc, "前端结果页展示科室、风险等级、就医时效、support_score、规则因素、证据因素、引用和安全边界。后端保存 TriageResult 和 RankedEvidence 快照，保证用户看到的证据与当时激活的索引版本一致。")
    placeholder(doc, "图5-3 正常问诊结果与医学证据页面截图（待用户运行系统后替换）")
    h(doc, "5.4 问诊记录、健康档案与复诊提醒模块", 2)
    p(doc, "问诊记录按照用户归属过滤，健康档案字段经过授权后才作为背景上下文注入 AI 服务，原始症状文本不被覆盖。复诊提醒支持用户创建带到期时间的任务，患者端轮询到期任务并可以完成或取消；当前未实现外部短信或微信推送。")
    placeholder(doc, "图5-4 问诊记录、健康档案与复诊提醒页面截图（待用户运行系统后替换）")
    h(doc, "5.5 医生复核队列模块", 2)
    p(doc, "医生复核队列支持按状态查询、领取任务、查看详情和提交决定。队列接口在列表阶段也执行医疗关系和院区过滤，防止通过列表发现无权访问的患者记录。")
    placeholder(doc, "图5-5 医生复核队列和复核详情页面截图（待用户运行系统后替换）")
    h(doc, "5.6 医学知识治理模块", 2)
    p(doc, "知识治理页面展示文档来源、专科、审核状态、有效期、索引版本和版本差异。知识编辑器负责提交和构建，医生、复核员或管理员按权限完成审核和激活；构建失败不会覆盖正在使用的索引。")
    h(doc, "5.7 智能体运行监控与审计模块", 2)
    p(doc, "监控模块从 consultation_traces 读取成功或失败轨迹，展示 trace_id、终止阶段、错误码、总耗时和节点耗时。审计模块记录操作者、资源、动作和时间戳，监控界面只读，不会重新触发模型推理。")
    figure(doc, "5-6", "Trace 监控与审计流程图", "10_Trace监控与审计流程图.png")
    h(doc, "5.8 用户认证、MFA 与医疗数据安全模块", 2)
    p(doc, "认证模块使用 JWT Cookie 和 CSRF 防护，角色由后端统一判断。医生复核要求 MFA assurance level 不低于2和医院员工档案；医疗数据敏感字段加密保存；附件校验扩展名、MIME 和文件头，限制大小为10 MB，图像和音频不自动执行 OCR、ASR 或诊断。")
    figure(doc, "5-7", "角色权限模型图", "08_角色权限模型图.png")
    h(doc, "5.9 本章小结", 2)
    p(doc, "本章从前端患者问诊到后端记录、AI 检索、医生复核、知识治理、Trace 监控和认证安全，说明了系统各功能模块的详细实现方式，并明确了需要运行系统后补充的界面截图。")


def chapter6(doc):
    h(doc, "第6章 系统测试与实验分析")
    h(doc, "6.1 测试环境与测试数据", 2)
    p(doc, "测试环境包含 Vue 3/Vite/Vitest 前端、Spring Boot/JUnit 5 后端、FastAPI/pytest/FAISS AI 服务、MySQL 8 和 Ollama。测试数据分为危险信号规则回归、主动追问夹具、RAG 排序回归、NDJSON 事件夹具和浏览器关键路径，均为人工构造或脱敏数据。")
    table(doc, "表6-1 测试环境与测试数据", ["测试层", "工具或数据", "说明"], [("前端", "Vitest、Playwright", "组件和浏览器关键路径"), ("后端", "JUnit 5、Spring Boot", "接口、权限和持久化"), ("AI 服务", "pytest、FAISS、构造病例", "规则、抽取、检索和事件"), ("业务数据", "人工构造或脱敏病例", "不用于临床效果推断")], [35, 60, 50])
    h(doc, "6.2 系统功能测试", 2)
    h(doc, "6.2.1 用户登录与权限测试", 3)
    p(doc, "测试不同角色访问患者接口、知识接口、监控接口和临床复核接口，验证未登录、角色不匹配、医疗关系不匹配和 MFA 不足时均失败关闭。管理员可进行用户和治理运维，但不因 ADMIN 角色自动获得临床复核权限。")
    h(doc, "6.2.2 正常问诊流程测试", 3)
    p(doc, "输入包含症状、持续时间和严重程度的普通咨询，验证安全筛查、抽取、检索、分诊、回答编排和 done 事件顺序完整，问诊记录和引用快照能够持久化。")
    h(doc, "6.2.3 信息不足追问测试", 3)
    p(doc, "输入只有单一症状或缺少持续时间时，验证系统进入 ask_followup 并返回问题，不提前生成完整科室推荐；后续补充信息后才能进入正常检索和分诊。")
    h(doc, "6.2.4 高风险快速通道测试", 3)
    p(doc, "输入胸痛、呼吸困难、咯血、便血、晕厥、意识不清和大出血等危险信号，验证未调用 LLM 和 embedding 即进入高风险结果，输出急诊或尽快就医建议。")
    h(doc, "6.2.5 医生复核流程测试", 3)
    p(doc, "验证 DOCTOR/REVIEWER 在员工档案、医疗关系和 MFA 满足时可以领取和决定；ADMIN、AUDITOR、患者以及自复核请求被拒绝；原始 AI 结果和人工决定分别保存。")
    table(doc, "表6-2 系统功能测试用例", ["编号", "测试场景", "预期结果"], [("FT-01", "登录和角色访问", "合法角色成功，越权请求拒绝"), ("FT-02", "正常咨询", "事件完整并保存结果"), ("FT-03", "信息不足", "返回主动追问"), ("FT-04", "危险信号", "高风险快速通道"), ("FT-05", "医生复核", "医疗关系和 MFA 校验"), ("FT-06", "AI 流失败", "保存失败 Trace，不生成成功记录")], [25, 65, 55])
    h(doc, "6.3 辅助分诊算法评测", 2)
    h(doc, "6.3.1 安全召回率", 3)
    p(doc, "安全召回率定义为危险信号被正确识别的病例数除以危险信号病例总数。当前12条红旗规则回归样例中，预期危险信号均被识别，否定表达样例未误报，工程规则回归安全召回率为1.0。该结果仅适用于构造规则样例。")
    h(doc, "6.3.2 科室 Macro-F1", 3)
    p(doc, "科室 Macro-F1 在构造病例标签上统计各科室 F1 的宏平均。当前确定性十例基线的科室 Macro-F1 为1.0；该指标来自软件回归样例，不能解释为临床人群准确率。")
    h(doc, "6.3.3 风险等级 Macro-F1", 3)
    p(doc, "风险等级 Macro-F1 统计高、中、低或未定等风险类别的宏平均。当前确定性基线风险 Macro-F1 为1.0，主要由规则优先路径和构造样例组成，仍需真实数据和临床盲评验证。")
    h(doc, "6.3.4 错误率与响应延迟", 3)
    p(doc, "确定性十例基线错误率为0。真实普通流程的节点事件 elapsed_ms 求和为2826 ms；红旗浏览器路径观察到的墙钟时间为146 ms；追问夹具节点耗时为41 ms。三类数据来源不同，不能混合解释为统一服务延迟指标。")
    table(doc, "表6-3 辅助分诊算法评测结果", ["指标", "结果", "解释边界"], [("安全召回率", "1.0", "12条规则回归样例"), ("科室 Macro-F1", "1.0", "确定性构造病例"), ("风险 Macro-F1", "1.0", "确定性构造病例"), ("错误率", "0", "确定性构造病例"), ("普通流程节点耗时和", "2826 ms", "事件耗时和，不是统一 p95")], [50, 35, 60])
    h(doc, "6.4 医学 RAG 检索评测", 2)
    h(doc, "6.4.1 Recall@K", 3)
    p(doc, "Recall@K 需要每个病例提供 gold_evidence_ids。当前正式病例没有该字段，因此 Recall@K 标记为 not-evaluated。四条 in-memory 排序回归样例的 top1 均命中预期科室，只能说明测试索引排序行为。")
    h(doc, "6.4.2 MRR", 3)
    p(doc, "MRR 同样依赖人工标注的相关证据集合。当前没有独立 gold evidence 标注，不能用占位零值解释为 MRR 为0，也不能声称 RAG 已达到某种临床检索水平。")
    h(doc, "6.4.3 引用可追溯率", 3)
    p(doc, "引用可追溯率应检查回答中的 citation_id 是否能回到 doc_id、chunk_id、source、quote 和 index_version。系统已实现字段链路和快照保存，但当前没有独立人工标注集，因此总体引用可追溯率暂记为 not-evaluated。")
    table(doc, "表6-4 RAG 检索评测状态", ["指标", "当前状态", "原因"], [("Recall@K", "not-evaluated", "缺少 gold_evidence_ids"), ("MRR", "not-evaluated", "缺少相关证据排序标注"), ("引用可追溯率", "not-evaluated", "缺少人工审核样本"), ("受控排序回归", "4/4 top1 命中", "仅验证 in-memory 测试索引")], [45, 45, 55])
    h(doc, "6.5 对比实验与消融实验", 2)
    p(doc, "本文设计三组工程消融：移除危险信号前置筛查，观察高风险请求是否仍能快速进入安全路径；移除词法融合，观察短中文症状的候选排序变化；移除 abstain，观察证据不足时是否出现强行推荐。由于当前没有独立临床标注集，消融结果只作为工程行为对比，不能给出临床效果结论。")
    table(doc, "表6-5 对比与消融实验设计", ["实验组", "移除或保留机制", "观察指标"], [("完整系统", "规则、词法融合、abstain 均保留", "安全事件、证据和输出状态"), ("消融A", "移除危险信号前置筛查", "红旗路径和模型调用顺序"), ("消融B", "移除字符 n-gram 融合", "短症状检索排序"), ("消融C", "移除 abstain", "证据不足时的输出边界")], [45, 65, 35])
    h(doc, "6.6 性能、安全性与鲁棒性测试", 2)
    p(doc, "自动化统计显示 AI service 150项、Spring Boot backend 70项、Vue frontend 52项和 Playwright browser flows 11项测试均通过。安全测试覆盖 JWT Cookie、CSRF、RBAC、会话归属、医疗关系、MFA、自复核禁止、敏感字段加密和附件大小及文件头校验。鲁棒性测试覆盖空输入、信息不足、模型异常、无证据、无效 NDJSON 和索引构建失败。")
    figure(doc, "6-1", "系统部署架构图", "09_系统部署架构图.png")
    h(doc, "6.7 实验结果分析", 2)
    p(doc, "从工程结果看，确定性危险信号规则和事件协议能够稳定回归，普通路径能够产生结构化结果和证据快照，错误链路能够被持久化。另一方面，RAG 关键指标尚未完成独立人工标注，临床安全性尚未完成专业人员盲评，因此本文结果只能支持“系统按设计运行”的结论，不能支持“系统在真实患者中准确有效”的结论。")
    placeholder(doc, "图6-2 普通问诊结果测试截图（待用户运行系统后替换）")
    placeholder(doc, "图6-3 高风险快速通道和医生复核测试截图（待用户运行系统后替换）")
    h(doc, "6.8 本章小结", 2)
    p(doc, "本章从测试环境、功能测试、辅助分诊指标、RAG 指标、消融设计、性能安全鲁棒性和结果边界等方面验证系统。结果表明工程回归路径稳定，但临床效果和检索金标准仍需要后续研究补充。")


def chapter7(doc):
    h(doc, "第7章 总结与展望")
    h(doc, "7.1 研究工作总结", 2)
    p(doc, "本文围绕多专科医疗健康咨询及辅助分诊系统，完成了需求分析、分层架构、数据库、接口协议、多智能体工作流、RAG 检索、危险信号规则、医生复核、知识治理、Trace 监控和权限安全模块的设计与实现。")
    h(doc, "7.2 系统主要贡献", 2)
    p(doc, "系统主要贡献包括：将多智能体协同落实为可测试的职责节点和条件有向图；将危险信号规则前置于模型调用；将医学证据、科室投票和引用快照结合；将医生复核、医疗关系、MFA、知识版本和审计纳入同一工程闭环；将测试结果与临床效果边界明确区分。")
    h(doc, "7.3 系统不足", 2)
    p(doc, "系统缺少临床专家盲评、独立检索金标准、真实世界前瞻验证和多机构数据；规则覆盖和同义词扩展仍有限；监控查询在大数据量下需要进一步数据库分页和索引优化；界面截图和部分运行指标需要结合最终部署环境补充。")
    h(doc, "7.4 未来展望", 2)
    p(doc, "未来可建立专家标注的多专科检索数据集，开展 Recall@K、MRR、引用审查和医生一致性评估；研究隐私合规的多模态附件确认流程；完善机构、科室和院区级授权；引入模型与知识版本的自动回归门禁，并在伦理审批和临床合作基础上开展前瞻性验证。")


def references_and_appendix(doc):
    h(doc, "参考文献")
    refs = [
        "[1] Semigran H L, Linder J A, Gidengil C, et al. Evaluation of symptom checkers for self diagnosis and triage: audit study[J]. BMJ, 2015, 351: h3480.",
        "[2] Chambers D, Cantrell A J, Johnson M, et al. Digital and online symptom checkers and health assessment/triage services for urgent health problems: systematic review[J]. BMJ Open, 2019, 9(8): e027743.",
        "[3] Wallace W, Chan C, Chidambaram S, et al. The diagnostic and triage accuracy of digital and online symptom checker tools: a systematic review[J]. npj Digital Medicine, 2022, 5: 118.",
        "[4] Schmieding M L, Kopka M, Schmidt K, et al. Triage Accuracy of Symptom Checker Apps: 5-Year Follow-up Evaluation[J]. Journal of Medical Internet Research, 2022, 24(5): e31810.",
        "[5] Fraser H, Crossland D, Bacher I, et al. Comparison of Diagnostic and Triage Accuracy of Ada Health and WebMD Symptom Checkers, ChatGPT, and Physicians[J]. JMIR mHealth and uHealth, 2023, 11: e49995.",
        "[6] Singhal K, Azizi S, Tu T, et al. Large language models encode clinical knowledge[J]. Nature, 2023, 620: 172-180.",
        "[7] Hager P, Jungmann F, Holland R, et al. Evaluation and mitigation of the limitations of large language models in clinical decision-making[J]. Nature Medicine, 2024, 30: 2613-2622.",
        "[8] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]. NeurIPS, 2020: 9459-9474.",
        "[9] Zakka C, Shad R, Chaurasia A, et al. Almanac—Retrieval-Augmented Language Models for Clinical Medicine[J]. NEJM AI, 2024, 1(2).",
        "[10] Xiong G, Jin Q, Lu Z, et al. Benchmarking Retrieval-Augmented Generation for Medicine[C]. Findings of ACL, 2024: 6233-6251.",
        "[11] Yao S, Zhao J, Yu D, et al. ReAct: Synergizing Reasoning and Acting in Language Models[C]. ICLR, 2023.",
        "[12] Chen J, Xiao S, Zhang P, et al. M3-Embedding: Multi-Linguality, Multi-Functionality, Multi-Granularity Text Embeddings Through Self-Knowledge Distillation[C]. Findings of ACL, 2024: 2318-2335.",
        "[13] Qwen Team. Qwen2.5 Technical Report[EB/OL]. arXiv:2412.15115, 2024.",
        "[14] Lu M, Ho B, Ren D, Wang X. TriageAgent: Towards Better Multi-Agents Collaborations for Large Language Model-Based Clinical Triage[C]. Findings of EMNLP, 2024: 5747-5764.",
    ]
    for ref in refs:
        pp = doc.add_paragraph()
        pp.paragraph_format.left_indent = Pt(20)
        pp.paragraph_format.first_line_indent = Pt(-20)
        pp.paragraph_format.line_spacing = Pt(18)
        rr = pp.add_run(ref)
        base.set_font(rr, "Times New Roman", 10.5)
    doc.add_page_break()
    h(doc, "致谢")
    p(doc, "感谢指导教师在课题选题、系统设计、算法实现和论文撰写过程中提供的指导，感谢同学在测试、界面检查和复核流程讨论中提供的帮助。本文病例均为人工构造或脱敏数据，MedPilot 仅作为证据增强、规则护栏下的辅助分诊原型。")
    doc.add_page_break()
    h(doc, "附录")
    h(doc, "附录A 系统截图替换清单", 2)
    p(doc, "以下截图由用户在本地运行系统后提供，放入 outputs/thesis-images/user-screenshots 目录，再替换正文中的占位框。截图需要隐藏姓名、手机号、身份证号、病历号、token 和数据库密钥。")
    table(doc, "表A-1 系统截图替换清单", ["序号", "建议文件名", "截图内容"], [(1, "01_登录页面.png", "登录、角色和 MFA"), (2, "02_患者智能问诊.png", "患者端问诊输入"), (3, "03_正常问诊结果.png", "分诊结果和引用"), (4, "04_主动追问.png", "信息不足追问"), (5, "05_高风险快速通道.png", "危险信号提示"), (6, "06_问诊记录.png", "记录和健康档案"), (7, "07_医生复核队列.png", "复核列表"), (8, "08_医生复核详情.png", "复核决定"), (9, "09_医学知识库管理.png", "知识审核和索引版本"), (10, "10_智能体运行监控.png", "Trace 监控"), (11, "11_用户权限管理.png", "用户和角色"), (12, "12_审计日志.png", "审计记录")], [18, 65, 72])
    h(doc, "附录B 核心接口示例", 2)
    p(doc, "咨询接口示例：POST /api/consult，输入 text、session_id 和 history，响应为 NDJSON。事件示例包括 safety_screen completed、extract completed、retrieve completed、classify completed、compose completed 和 done completed。临床复核接口包括 GET /api/clinical-reviews、POST /api/clinical-reviews/{id}/claim 和 POST /api/clinical-reviews/{id}/decision。")
    h(doc, "附录C 评测数据说明", 2)
    p(doc, "当前评测数据包含12条危险信号规则回归样例、5条主动追问样例、4条受控 RAG 排序样例、normal/red_flag/followup 三类 NDJSON 夹具以及自动化测试统计。RAG 的 Recall@K、MRR 和引用可追溯率因缺少 gold_evidence_ids 和人工证据标注而暂未评测。")


def collect_headings(doc):
    return [x.text.strip() for x in doc.paragraphs if x.style.name.startswith("Heading")]


def build():
    doc = Document()
    base.configure(doc)
    doc.core_properties.title = "基于多智能体协同与大语言模型的多专科医疗健康咨询及辅助分诊系统"
    doc.core_properties.subject = "按用户目录重新编写的本科毕业设计论文"
    front_matter(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    chapter6(doc)
    chapter7(doc)
    references_and_appendix(doc)
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    expected = [
        "第1章 绪论", "1.1 研究背景", "1.2 研究目的与意义", "1.3 国内外研究现状", "1.3.1 医疗健康咨询研究现状", "1.3.2 智能辅助分诊研究现状", "1.3.3 多智能体与大语言模型应用现状", "1.3.4 现有研究不足", "1.4 研究内容", "1.5 技术路线与论文组织", "1.6 本章小结",
        "第2章 系统分析与相关技术", "2.1 系统可行性分析", "2.1.1 技术可行性", "2.1.2 经济可行性", "2.1.3 操作可行性", "2.1.4 医疗安全与伦理可行性", "2.2 系统角色与业务流程分析", "2.3 系统功能需求分析", "2.3.1 患者智能问诊需求", "2.3.2 辅助分诊需求", "2.3.3 医生复核需求", "2.3.4 医学知识库管理需求", "2.3.5 监控、审计与治理需求", "2.4 系统非功能需求分析", "2.5 相关技术介绍", "2.5.1 大语言模型与 Ollama", "2.5.2 多智能体协同与 LangGraph", "2.5.3 RAG、Embedding 与 FAISS", "2.5.4 Vue 3、Spring Boot 与 FastAPI", "2.5.5 NDJSON 流式通信与 Trace", "2.6 本章小结",
        "第3章 系统总体设计", "3.1 系统设计目标与原则", "3.2 系统总体架构设计", "3.3 系统功能模块设计", "3.4 多智能体分诊工作流设计", "3.5 用户角色与权限设计", "3.6 医疗关系与医生复核流程设计", "3.7 数据库设计", "3.7.1 问诊记录相关数据表", "3.7.2 患者与医疗关系数据表", "3.7.3 医生复核数据表", "3.7.4 知识治理、审计与监控数据表", "3.8 接口与事件协议设计", "3.9 本章小结",
        "第4章 核心算法设计与实现", "4.1 系统辅助分诊总体算法流程", "4.2 多智能体协同算法", "4.2.1 智能体角色划分", "4.2.2 结构化状态与消息传递", "4.2.3 条件有向图路由机制", "4.2.4 主动追问与信息充分性判断", "4.3 医学知识检索增强生成算法", "4.3.1 医学知识库构建与文本切分", "4.3.2 医学文本向量化与 FAISS 索引", "4.3.3 向量检索与词法匹配融合", "4.3.4 基于证据的科室推荐", "4.3.5 引用可追溯与回答约束", "4.4 医疗安全与兜底机制", "4.4.1 危险信号识别与否定表达处理", "4.4.2 高风险快速通道", "4.4.3 证据不足时的 abstain 机制", "4.4.4 医生复核与人工接管机制", "4.5 算法复杂度与局限性分析", "4.6 本章小结",
        "第5章 系统详细设计与实现", "5.1 患者端智能问诊模块", "5.2 症状抽取与主动追问模块", "5.3 医学知识检索与辅助分诊模块", "5.4 问诊记录、健康档案与复诊提醒模块", "5.5 医生复核队列模块", "5.6 医学知识治理模块", "5.7 智能体运行监控与审计模块", "5.8 用户认证、MFA 与医疗数据安全模块", "5.9 本章小结",
        "第6章 系统测试与实验分析", "6.1 测试环境与测试数据", "6.2 系统功能测试", "6.2.1 用户登录与权限测试", "6.2.2 正常问诊流程测试", "6.2.3 信息不足追问测试", "6.2.4 高风险快速通道测试", "6.2.5 医生复核流程测试", "6.3 辅助分诊算法评测", "6.3.1 安全召回率", "6.3.2 科室 Macro-F1", "6.3.3 风险等级 Macro-F1", "6.3.4 错误率与响应延迟", "6.4 医学 RAG 检索评测", "6.4.1 Recall@K", "6.4.2 MRR", "6.4.3 引用可追溯率", "6.5 对比实验与消融实验", "6.6 性能、安全性与鲁棒性测试", "6.7 实验结果分析", "6.8 本章小结",
        "第7章 总结与展望", "7.1 研究工作总结", "7.2 系统主要贡献", "7.3 系统不足", "7.4 未来展望", "参考文献", "致谢", "附录"
    ]
    actual = collect_headings(doc)
    missing = [item for item in expected if item not in actual]
    if missing:
        raise RuntimeError("missing headings: " + ", ".join(missing))
    print(f"saved={OUT}; headings={len(actual)}; missing=0")


if __name__ == "__main__":
    build()
