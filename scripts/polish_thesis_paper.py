from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"D:\毕设制作")
INPUT = ROOT / "outputs" / "7230264110_窦成皓_23软件1_基于多智能体协同与大语言模型的多专科医疗健康咨询及辅助分诊系统_千例评测增强版.docx"
OUTPUT = ROOT / "outputs" / "7230264110_窦成皓_23软件1_基于多智能体协同与大语言模型的多专科医疗健康咨询及辅助分诊系统_千例评测增强版_paper-polish润色版.docx"


def set_paragraph_text(paragraph, text: str) -> None:
    """Replace paragraph text while retaining paragraph and first-run formatting."""
    p = paragraph._p
    first_run = paragraph.runs[0] if paragraph.runs else None
    rpr = deepcopy(first_run._r.get_or_add_rPr()) if first_run is not None else None
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    run = OxmlElement("w:r")
    if rpr is not None:
        run.append(rpr)
    node = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" ") or "  " in text:
        node.set(qn("xml:space"), "preserve")
    node.text = text
    run.append(node)
    p.append(run)


def replace_in_runs(paragraph, old: str, new: str) -> bool:
    changed = False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            changed = True
    return changed


def replace_paragraph_substrings(paragraph, replacements: list[tuple[str, str]]) -> bool:
    old_text = paragraph.text
    new_text = old_text
    for old, new in replacements:
        new_text = new_text.replace(old, new)
    if new_text == old_text:
        return False
    if len(paragraph.runs) <= 1:
        set_paragraph_text(paragraph, new_text)
    else:
        # Most multi-run paragraphs are title/keyword blocks. Keep their direct
        # formatting and apply only replacements that stay within a run.
        changed = False
        for old, new in replacements:
            changed = replace_in_runs(paragraph, old, new) or changed
        if not changed:
            set_paragraph_text(paragraph, new_text)
    return True


def add_citation(paragraph, marker: str, anchor: str | None = None) -> bool:
    text = paragraph.text
    if marker in text:
        return False
    if anchor and anchor in text:
        return replace_paragraph_substrings(paragraph, [(anchor, anchor + marker)])
    if text.endswith("。"):
        # Insert once, at the paragraph's final sentence boundary.  Replacing
        # every full stop would attach the same citation to unrelated claims.
        return replace_paragraph_substrings(paragraph, [(text, text[:-1] + marker + "。")])
    return replace_paragraph_substrings(paragraph, [(text, text + marker)])


def rewrite_target_paragraphs(doc: Document) -> None:
    # These edits remove field-by-field narration from Chapters 3 and 4 while
    # retaining the data model and implementation meaning.
    rewrites = {
        "AI服务的工作流状态由输入文本、历史消息、健康档案上下文、症状、追问、证据、分诊结果、回答和事件发送器等字段组成。症状、证据、分诊结果、回答和追问均使用不可变Pydantic模型，节点只能生成新的状态更新，不能在后续节点中随意修改已有结果。":
            "AI服务以统一的工作流状态承载本轮输入及各节点产出的中间结果。与医疗业务含义直接相关的对象均采用不可变 Pydantic 模型，节点只能提交新的状态更新，不能在后续节点中改写已形成的结果。",
        "工作流状态还包含会话标识、轮次和历史模式。历史模式区分完整对话和摘要上下文，避免把过长的历史消息无限传给模型。健康档案上下文只包含用户授权的过敏史、既往情况、用药信息和备注，并在抽取提示中明确其仅为背景资料，不能被当作本轮症状或指令执行，各智能体节点的输入、输出及路由条件汇总见表3-3。":
            "为控制上下文规模并保持会话连续性，工作流还记录本轮会话的标识、轮次及历史处理方式。历史处理方式在完整对话和摘要上下文之间切换，避免过长的历史消息持续传入模型。健康档案仅在用户授权后作为背景资料参与抽取，并明确与本轮症状和指令分离。各智能体节点的输入、输出及路由条件汇总见表3-3。",
        "每条证据以RankedEvidence模型返回，包含citation_id、doc_id、chunk_id、department、source、quote、score、index_version、institution、title、url、published_date、license和review_status等字段。":
            "每条证据以 RankedEvidence 模型返回，并携带能够定位来源、分片、科室、分数、索引版本和审核状态的元数据。",
        "users表以id为主键，以username建立唯一约束，保存password_hash、role、active、token_version和created_at等字段，并扩展identity_provider、external_subject、organization_code、campus_code、department_code、patient_mpi_id和mfa_assurance_level等身份上下文。consultation_sessions以session_id作为业务唯一标识，通过user_id关联用户，保存创建时间和最近活跃时间。consultation_messages保存会话中的user或assistant消息，并可关联trace_id。":
            "用户、会话和消息分别承担身份管理、对话组织与过程关联职责。用户实体保存账户状态、角色及必要的身份上下文；会话实体维护用户归属和活跃时间；消息实体按用户或助手角色保存对话内容，并与执行轨迹关联。这样的分层既支持会话所有权校验，也为历史消息回放提供稳定入口。",
        "consultation_records是问诊结果的业务快照，保存user_id、patient_mpi_id、session_id、trace_id、symptoms、department、risk_level、confidence、support_score、abstained、urgency、matched_rule、triage_factors、explanation、answer、citations和conversation_history。trace_id建立唯一约束，避免同一执行轨迹生成多条主记录。症状、回答、引用和历史均采用长文本字段保存结构化JSON或序列化内容，具体解析格式由后端服务统一管理。":
            "consultation_records 是一次咨询完成后的业务快照，关联用户、患者上下文、会话和执行轨迹，同时保存结构化症状、分诊结论、解释、回答、引用以及对话上下文。Trace 标识设置唯一约束，避免同一执行轨迹生成重复主记录；需要保持结构的内容由后端统一序列化和解析。",
        "consultation_traces以trace_id为唯一键，保存session_id、user_id、events_json、citations_json、terminal_phase、followup_pending、failure_code、total_duration_ms和created_at。events_json保存节点事件序列，citations_json保存本次咨询的引用快照，terminal_phase区分completed、escalated、failed等终态，failure_code用于定位失败原因。audit_logs则记录event_id、actor_username、actor_role、method、action、status、success、request_id、ip_hash、duration_ms和created_at，用于审计操作者和接口行为。":
            "consultation_traces 以 trace_id 为唯一键，保存一次咨询的事件序列、引用快照、终态、追问标记、失败原因和耗时等过程信息。它用于还原节点执行顺序并区分完成、升级和失败等结果。audit_logs 则面向受保护接口记录操作者、角色、动作、响应状态和耗时，为权限审计与异常定位提供依据。",
        "health_profiles以user_id建立唯一约束，profile_json保存用户授权的背景信息，consent_granted表示是否允许其进入问诊上下文。follow_up_tasks保存user_id、record_id、title、notes、due_at和status，支持根据问诊记录建立复诊或提醒任务。面向医院扩展的patients以mpi_id标识患者，patient_encounters保存就诊号、组织、院区、科室和责任医生；patient_care_relationships记录医生与患者之间的关系；break_glass_accesses记录紧急授权的目的、理由、有效期和撤销时间。":
            "健康档案与复诊任务分别保存经用户授权的背景信息和后续提醒，并通过问诊记录建立关联。面向医院扩展的患者、就诊和医疗关系实体用于表达组织、院区、科室及责任医生边界；紧急授权则单独保存目的、理由、有效期和撤销记录。这样可以把背景资料、业务任务和临床访问关系分开管理。",
        "knowledge_documents以doc_id为主键，保存title、department、source_type、institution、url、published_date、source_version、license_name、checksum、parsing_status、vector_status、review_status、chunk_count、processing_error、reviewer和reviewed_at等字段。该表保存知识文档的生命周期状态，不直接保存FAISS向量本体。向量文件和分片元数据由AI服务的索引目录管理，索引版本写入检索证据。":
            "knowledge_documents 聚焦知识文档的来源、适用科室、版本、许可、摘要、解析、向量化和审核生命周期，不直接保存 FAISS 向量本体。向量文件与分片元数据由 AI 服务的索引目录管理，检索证据记录所使用的索引版本。",
        "知识库支持结构化文本和文件上传两种方式。文件上传接口读取doc_id、department、institution、title、url、published_date、version、license等元数据，并根据扩展名判断文本和PDF等受支持类型；系统限制文件大小，计算SHA-256摘要并保存来源信息。PDF文本由PDFBox提取，解析失败时将文档置为解析失败并保留错误信息。":
            "知识库支持结构化文本和文件上传两种方式。文件上传时，系统读取文档标识、科室、机构、题名、来源链接、发布日期、版本和许可等必要元数据，并根据扩展名判断受支持类型；同时限制文件大小、计算 SHA-256 摘要并保存来源信息。PDF 文本由 PDFBox 提取，解析失败时保留错误状态和原因。",
        "评测文件同时保存每条样本的Top-3文档标识和分数。":
            "评测文件同时保存每条样本的 Top-3 文档及其排序分数，以便复核检索相关性。",
        "引用测试继续检查citation_id、doc_id、chunk_id、source、url、published_date、version、license和review_status是否完整。":
            "引用测试继续检查引用是否具备完整的来源、分片、链接、日期、版本、许可和审核信息。",
        "每条事件包含protocol_version、trace_id、session_id、sequence、type、status、state和data等字段。node事件描述某个智能体节点的开始、完成或错误，answer_delta事件传递回答增量，done事件标志本轮终态，error事件表示流程失败。state中保存intent、phase、turn_count和history_mode，data中保存症状、追问、证据、分诊或回答等结构化内容。":
            "每条事件都遵循统一协议，事件类型分别表达节点执行、回答增量、正常终止和流程错误。事件同时携带本轮会话的阶段状态及结构化业务结果，使前端能够渐进式展示，后端能够按序校验并保存完整轨迹。",
        "clinical_reviews与consultation_records一对一关联，保存AI原始科室、风险等级、就医时效和Trace标识，同时保存领取人、复核人、最终科室、最终风险等级、最终就医时效、决定理由和处理时间。设计上，AI原始结果不可被医生决定覆盖；医生决定作为独立字段写入复核聚合。复核队列按status和created_at索引，领取动作记录claimed_by_user_id和claimed_at，避免多个复核人员同时处理同一任务。":
            "clinical_reviews 与 consultation_records 一对一关联，既保留 AI 原始建议及其 Trace 标识，也保存领取、复核、最终决定、理由和处理时间。AI 原始结果在复核过程中保持不变，医生决定作为独立结果写入复核聚合；队列状态和领取动作采用并发控制，避免同一任务被重复处理。",
        "筛查结果被转换为`SafetyScreenResult`，包含matched、matched_terms、triage和rule_id等字段。":
            "筛查结果被转换为 `SafetyScreenResult`，用于承载是否命中、命中项、风险路径和规则标识。",
        "抽取提示要求聊天模型只返回JSON，字段包括symptoms、duration、severity和history。服务端使用`_parse_llm_json`截取首个左大括号到最后一个右大括号之间的内容，再交给JSON解析器；如果模型返回无法解析的内容，系统返回空结构而不是把自然语言当作结构化字段。":
            "抽取提示要求聊天模型仅返回结构化 JSON。服务端使用 `_parse_llm_json` 截取 JSON 边界并交给解析器；如果模型输出无法解析，系统返回空结构，而不会把自然语言直接当作结构化字段。",
        "抽取结果经过`normalize_symptoms`归一化，历史字段转换为不可变元组，当前文本再次通过危险信号匹配器生成red_flags。健康档案在进入提示词前经过字段白名单和长度限制，只允许allergies、conditions、medications和notes四类字符串字段。":
            "抽取结果经过 `normalize_symptoms` 归一化，历史信息转为不可变结构，当前文本再次通过危险信号匹配器复核。健康档案进入提示词前经过白名单和长度限制，并明确只作为授权背景资料使用。",
        "低于`MEDPILOT_RAG_MIN_SCORE`的结果被丢弃，剩余证据按综合分排序后截取Top-K。每条结果构造为`RankedEvidence`，包含引用标识、文档标识、分片标识、科室、来源、证据原文、分数、索引版本、机构、题名、URL、发布日期、许可和审核状态。":
            "低于 `MEDPILOT_RAG_MIN_SCORE` 的结果被丢弃，剩余证据按综合分排序后截取 Top-K。每条结果构造为 `RankedEvidence`，并保存可定位来源、科室、证据原文、分数、索引版本及审核状态的元数据。",
        "Trace详情展示trace_id、session_id、意图、阶段和轮次，并按序号渲染事件。节点耗时、失败代码、终态和引用数量用于辅助定位问题。监控图表展示已采集的运行样本，不把样本量或耗时直接解释为医疗质量指标。":
            "Trace 详情按会话、阶段和轮次展示事件序列，并汇总节点耗时、失败原因、终态和引用数量以辅助定位问题。监控图表只反映已采集的运行样本，不把样本量或耗时直接解释为医疗质量指标。",
    }
    for paragraph in doc.paragraphs:
        for old, new in rewrites.items():
            if old in paragraph.text:
                set_paragraph_text(paragraph, paragraph.text.replace(old, new))

    # A few source paragraphs contain additional implementation details after
    # the passages above.  Rewrite them by stable sentence starts so the prose
    # does not become a catalogue of database or protocol column names.
    starts = {
        "系统使用Flyway按版本执行迁移。":
            "系统使用 Flyway 按版本执行迁移，并将账户、会话、问诊、知识、复核和治理等数据按业务边界分开管理。基础迁移先建立核心业务实体，后续迁移再补充消息、知识文档、医院身份关系、临床复核与治理证据。",
        "问诊开始时，后端通过`SessionOwnershipService`认领会话":
            "问诊开始时，后端通过 `SessionOwnershipService` 校验并认领会话，防止不同用户交叉使用同一会话。系统先保存本轮用户消息，再按用户与会话读取历史上下文；AI 服务返回回答或追问后，后端追加助手消息。历史内容按用户和助手角色保存，并可在后续轮次采用完整历史或摘要上下文。",
        "`ConsultationPersistenceService`使用事务保存Trace":
            "`ConsultationPersistenceService` 使用事务保存 Trace、事件和问诊结果，并在写入前检查执行轨迹是否已经存在，以保证重复提交不会生成重复记录。只有在终态明确且业务结果完整时，系统才保存科室、风险、置信度、支持度、解释、回答、引用和对话上下文。",
        "`HealthProfileController`提供档案查询、更新和删除接口":
            "`HealthProfileController` 提供档案查询、更新和删除接口。更新时，系统将用户授权的过敏史、既往情况、用药信息和备注保存为受控背景上下文，并记录是否同意将其用于问诊。读取操作始终依据当前认证身份执行，不能通过请求参数访问其他用户的档案。",
        "复诊任务由`FollowUpTask`":
            "复诊任务由 `FollowUpTask` 与问诊记录关联，系统提供任务列表、单条任务查询、完成和更新状态等操作。到期任务按照当前用户和开放状态筛选，避免把其他用户的提醒混入当前列表。",
        "知识库支持结构化文本和文件上传两种方式":
            "知识库支持结构化文本和文件上传两种方式。上传时系统校验科室、机构、题名、来源链接、发布日期、版本和许可等必要信息，并根据文件类型执行文本解析；同时限制文件大小、计算 SHA-256 摘要并保留来源信息。PDF 文本由 PDFBox 提取，解析失败时保留失败状态和原因。",
        "业务层在转发给AI服务前执行`sanitizeIngestPayload`":
            "业务层在转发给 AI 服务前执行 `sanitizeIngestPayload`，移除客户端不应提交的审核信息，强制新文档进入待审核状态，并重新计算摘要以防止元数据被篡改。文档元数据保存于 MySQL，切片和向量化结果由 AI 服务管理。",
        "文档通过审核后，审核接口通过`X-MedPilot-Reviewer`":
            "文档通过审核后，审核接口通过 `X-MedPilot-Reviewer` 传递当前认证主体，将审核结果同步给 AI 服务，并更新文档的审核人和审核时间。知识版本接口负责构建新的索引版本、切换活动版本和比较版本差异。",
        "检索结果写入`index_version`":
            "检索结果写入索引版本标识，引用快照因此能够保留本次咨询实际使用的知识环境。知识文档的删除、审核、版本切换和索引操作均受角色权限约束；解析状态和向量状态分开记录，避免把“已入库”误认为“可供检索”。",
        "`KnowledgeView.vue`提供知识统计":
            "`KnowledgeView.vue` 提供知识统计、文档列表、版本管理、版本差异和入库表单。入库表单支持文件模式和正文模式，要求填写科室、机构、题名、来源链接、发布日期、版本和许可信息；审核状态由后端审核接口最终控制。页面还提供版本构建、版本激活和文档审核操作，使知识生命周期能够在管理端闭环完成。",
        "临床复核控制器提供记录详情":
            "临床复核控制器提供记录详情、领取和决定接口。复核列表按状态和创建时间查询，领取动作使用版本与状态检查，防止多个复核人员同时处理同一任务；领取成功后，任务进入 IN_REVIEW 状态并记录处理时间。",
        "前端复核表单提供CONFIRM":
            "前端复核表单提供 CONFIRM、MODIFY、REJECT 和 ESCALATE 四类决定。CONFIRM 保留 AI 原始结果；MODIFY 需要提交最终科室、风险等级、就医时效和复核依据；REJECT 表示退回人工处理；ESCALATE 将任务转入急诊人工流程。复核依据为必填项，决定接口同时记录复核人、员工号、理由和处理时间。",
        "治理端还提供模型发布":
            "治理端还提供模型发布、评测、回滚、知识来源登记、变更审批、红队测试、回滚演练、监控快照和安全事件等入口。模型发布需要记录版本、签名、提示词与嵌入配置、知识索引版本、硬件基线和审批证据；高风险动作由具备相应权限的角色执行。",
        "治理变更按角色分配":
            "治理变更按角色分配。模型和知识证据由管理员或知识编辑人员提交，审核人员或医生复核，冻结、执行和回滚等高风险操作限制为管理员。系统不会因为存在某个模型记录就自动切换模型，模型和知识版本必须经过明确审批并激活。",
        "`MonitorView.vue`维护健康状态":
            "`MonitorView.vue` 维护健康状态、历史 Trace、实时 Trace、节点状态和事件列表。页面通过健康接口获取模型服务、知识索引和治理状态，通过 Trace 列表查询历史运行，通过事件流定位当前活动。节点列表展示安全筛查、症状采集、追问、知识检索、辅助分诊和回答编排的状态。",
        "`AuditLogInterceptor`为受保护Web请求":
            "`AuditLogInterceptor` 为受保护 Web 请求记录操作者、角色、HTTP 方法、动作、响应状态、请求标识、地址摘要和耗时，并将成功或失败结果写入审计日志。监控与审计页面分别提供查询和导出入口，不同角色只能读取其权限范围内的记录。",
    }
    # Match on stable technical anchors because Word may normalize punctuation
    # or whitespace around the Chinese lead sentence.
    anchor_rewrites = [
        ("SessionOwnershipService", starts["问诊开始时，后端通过`SessionOwnershipService`认领会话"]),
        ("ConsultationPersistenceService", starts["`ConsultationPersistenceService`使用事务保存Trace"]),
        ("HealthProfileController", starts["`HealthProfileController`提供档案查询、更新和删除接口"]),
        ("FollowUpTask", starts["复诊任务由`FollowUpTask`"]),
        ("知识库支持结构化文本和文件上传两种方式", starts["知识库支持结构化文本和文件上传两种方式"]),
        ("sanitizeIngestPayload", starts["业务层在转发给AI服务前执行`sanitizeIngestPayload`"]),
        ("X-MedPilot-Reviewer", starts["文档通过审核后，审核接口通过`X-MedPilot-Reviewer`"]),
        ("KnowledgeView.vue", starts["`KnowledgeView.vue`提供知识统计"]),
        ("临床复核控制器提供记录详情", starts["临床复核控制器提供记录详情"]),
        ("前端复核表单提供CONFIRM", starts["前端复核表单提供CONFIRM"]),
        ("治理端还提供模型发布", starts["治理端还提供模型发布"]),
        ("治理变更按角色分配", starts["治理变更按角色分配"]),
        ("MonitorView.vue", starts["`MonitorView.vue`维护健康状态"]),
        ("AuditLogInterceptor", starts["`AuditLogInterceptor`为受保护Web请求"]),
    ]
    for paragraph in doc.paragraphs:
        for anchor, new in anchor_rewrites:
            if anchor in paragraph.text and paragraph.text.strip() not in {new, new.rstrip("。")}: 
                set_paragraph_text(paragraph, new)
    force_rewrites = {
        "系统使用Flyway": "系统使用 Flyway 按版本执行数据库迁移，并将账户、会话、问诊、知识、复核和治理等数据按业务边界分开管理。基础迁移先建立核心业务实体，后续迁移再补充消息、知识文档、医院身份关系、临床复核与治理证据。",
        "model_releases": "治理数据按照模型发布、临床评测、知识来源、变更审批、红队测试、回滚演练、运行监测和安全事件等业务聚合分别保存。各聚合承担版本追踪、验证、审批或事件处置职责，避免在同一记录中混合模型状态与治理结论。",
        "知识库入库支持": "知识库入库支持结构化文本和文件上传两种方式。上传时系统校验科室、机构、题名、来源链接、发布日期、版本和许可等必要信息，并根据文件类型执行文本解析；同时限制文件大小、计算 SHA-256 摘要并保留来源信息。PDF 文本由 PDFBox 提取，解析失败时保留失败状态和原因。",
        "业务层在转发入库请求前": "业务层在转发入库请求前执行 `sanitizeIngestPayload`，移除客户端不应提交的审核信息，强制新文档进入待审核状态，并重新计算摘要以防止元数据被篡改。文档元数据保存于 MySQL，切片和向量化结果由 AI 服务管理。",
        "文档入库后必须经过审核": "文档入库后必须经过审核。审核接口通过 `X-MedPilot-Reviewer` 传递当前认证主体，将审核结果同步给 AI 服务，并更新文档的审核人和审核时间。知识版本接口负责构建新的索引版本、切换活动版本和比较版本差异。",
        "检索结果写入index_version": "检索结果写入索引版本标识，引用快照因此能够保留本次咨询实际使用的知识环境。知识文档的删除、审核、版本切换和索引操作均受角色权限约束；解析状态和向量状态分开记录，避免把“已入库”误认为“可供检索”。",
        "治理控制器提供": "治理控制器提供模型发布、评测、回滚、知识来源登记、变更审批、红队测试、回滚演练、监控快照和安全事件等入口。模型发布需要记录版本、签名、提示词与嵌入配置、知识索引版本、硬件基线和审批证据；高风险动作由具备相应权限的角色执行。",
        "治理模块采用角色分离": "治理模块采用角色分离。模型和知识证据由管理员或知识编辑人员提交，审核人员或医生复核，冻结、执行和回滚等高风险操作限制为管理员。系统不会因为存在某个模型记录就自动切换模型，模型和知识版本必须经过明确审批并激活。",
        "KnowledgeView.vue": "`KnowledgeView.vue` 提供知识统计、文档列表、版本管理、版本差异和入库表单。入库表单支持文件模式和正文模式，要求填写科室、机构、题名、来源链接、发布日期、版本和许可信息；审核状态由后端审核接口最终控制。页面还提供版本构建、版本激活和文档审核操作，使知识生命周期能够在管理端闭环完成。",
        "MonitorView.vue": "`MonitorView.vue` 维护健康状态、历史 Trace、实时 Trace、节点状态和事件列表。页面通过健康接口获取模型服务、知识索引和治理状态，通过 Trace 列表查询历史运行，通过事件流定位当前活动。节点列表展示安全筛查、症状采集、追问、知识检索、辅助分诊和回答编排的状态。",
        "AuditLogInterceptor": "`AuditLogInterceptor` 为受保护 Web 请求记录操作者、角色、HTTP 方法、动作、响应状态、请求标识、地址摘要和耗时，并将成功或失败结果写入审计日志。监控与审计页面分别提供查询和导出入口，不同角色只能读取其权限范围内的记录。",
    }
    for paragraph in doc.paragraphs:
        for anchor, new in force_rewrites.items():
            if anchor in paragraph.text and paragraph.text.strip() != new:
                set_paragraph_text(paragraph, new)


def add_references_and_polish(doc: Document) -> None:
    abstract_rewrites = {
        46: "随着互联网医疗、医学知识库和大语言模型技术的发展，自然语言交互已成为健康信息服务的重要入口。然而，医疗健康咨询具有症状表述不完整、专业知识密集和错误风险不对称等特点。若直接依赖大语言模型生成回答，容易出现危险信号识别不稳定、证据来源难以追溯以及咨询与诊断边界不清等问题。针对上述问题，本文设计并实现了基于多智能体协同与大语言模型的多专科医疗健康咨询及辅助分诊系统 MedPilot。",
        47: "系统采用 Vue 3、Spring Boot 和 FastAPI 构建前后端分离架构，AI 服务基于 LangGraph 组织安全筛查、症状抽取、信息充分性判断、主动追问、医学知识检索、辅助分诊和回答编排等固定职责节点。系统在调用大语言模型和向量检索前执行危险信号筛查，并根据信息完整程度选择主动追问或进入分诊流程；普通场景使用 bge-m3 生成查询向量，结合 FAISS 与词法重合度检索医学证据，再根据证据所属科室进行加权分析。系统还实现了 NDJSON 流式事件、问诊记录、授权健康档案、医生复核、知识审核与索引版本治理、Trace 监控、审计日志及基于医疗关系、组织范围和多因素认证的权限控制。",
        48: "本文使用固定种子构建并执行了 1000 条人工构造工程测试。危险信号召回率为 80.00%，四个受支持科室的 Macro-F1 为 0.9585；加入 200 条知识库外症状后，五分类 Macro-F1 降至 0.5861。检索 Recall@3 为 71.88%，主动追问触发率为 100.00%，低证据回退率为 0.00%。结果表明，当前系统在封闭科室分诊和结构化追问方面表现较稳定，但对未登录危险表达和知识库外输入的识别与回退仍需改进。上述结果仅用于说明原型系统的工程行为，不作为临床准确率或医疗效果结论。",
        56: "Natural-language interaction has become an important entry point for digital health services as Internet healthcare, medical knowledge bases, and large language models continue to develop. However, healthcare consultation involves incomplete symptom descriptions, specialized knowledge, and asymmetric consequences of errors. Directly generating responses with a large language model may therefore lead to unstable danger-signal detection, untraceable evidence, and an unclear boundary between consultation and diagnosis. To address these problems, this study designs and implements MedPilot, a multi-specialty healthcare consultation and auxiliary triage system based on multi-agent collaboration and large language models.",
        57: "MedPilot uses Vue 3, Spring Boot, and FastAPI in a separated front-end and back-end architecture. Its AI service uses LangGraph to coordinate fixed-responsibility nodes for safety screening, symptom extraction, information-sufficiency assessment, active follow-up questioning, medical knowledge retrieval, auxiliary triage, and answer composition. Danger signals are screened before large-language-model inference or vector retrieval. When information is insufficient, the system asks a targeted follow-up question; otherwise, bge-m3 generates query embeddings, and FAISS retrieval is combined with lexical overlap to obtain evidence for department-level weighted analysis. The system also provides NDJSON streaming events, consultation records, authorized health-profile context, physician review, knowledge and index-version governance, trace monitoring, audit logs, and access control based on care relationships, organizational scope, and multi-factor authentication.",
        58: "A fixed-seed engineering test set containing 1,000 manually constructed cases was executed. Danger-signal recall was 80.00%, and Macro-F1 across the four supported departments was 0.9585; after 200 out-of-knowledge-base cases were added, five-class Macro-F1 decreased to 0.5861. Retrieval Recall@3 was 71.88%, the follow-up trigger rate was 100.00%, and the low-evidence fallback rate was 0.00%. These results indicate stable behavior in closed-set department triage and structured follow-up, while exposing limitations in unseen danger expressions and out-of-distribution fallback. The findings characterize engineering behavior only and should not be interpreted as evidence of clinical accuracy or effectiveness.",
        59: "Key Words: Multi-Agent Collaboration; Large Language Models; Retrieval-Augmented Generation; Healthcare Consultation; Auxiliary Triage",
    }
    for index, value in abstract_rewrites.items():
        set_paragraph_text(doc.paragraphs[index], value)

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "参考文献":
            break
        text = paragraph.text
        replacements = [
            ("系统MedPilot", "系统 MedPilot"),
            ("MedPilot系统", "MedPilot 系统"),
            ("工程Demo", "工程 Demo"),
            ("工程测试集共1000条，具体分布如表5-2。", "工程测试集共 1000 条，具体分布如表 5-2 所示。"),
            ("图5-1按扰动等级比较", "图 5-1 按扰动等级比较"),
            ("本章使用固定种子生成并执行1000条", "本章使用固定种子生成并执行 1000 条"),
            ("不能将曲线间高低直接解释为同一临床能力", "不能将曲线间的高低直接解释为同一临床能力"),
            ("不止于一次性聊天页面", "不止停留在一次性聊天页面"),
            ("运行界面如图4-7所示", "运行界面如图 4-7 所示"),
            ("对应结果页面如图4-4所示", "对应结果页面如图 4-4 所示"),
        ]
        if text.strip():
            replace_paragraph_substrings(paragraph, replacements)

    citation_rules = [
        ("在线医疗咨询研究大体可以分为三条路线。", "[1-5]", None),
        ("其中，急诊分诊更强调时间敏感性", "[2-5]", "时间敏感性"),
        ("大语言模型为自然语言交互提供了新的实现路径。", "[6-9]", None),
        ("医疗场景对错误的容忍度低。", "[10-14]", None),
        ("知识来源是另一个需要解决的问题。", "[15-23]", None),
        ("多智能体协同提供了另一种组织思路。", "[10,24]", None),
        ("张建同等", "[1]", "张建同等"),
        ("杨霞等", "[3]", "杨霞等"),
        ("徐伟等", "[4]", "徐伟等"),
        ("Tam等提出的QUEST框架", "[13]", "Tam等"),
        ("TRIPOD-LLM相关研究", "[14]", "TRIPOD-LLM相关研究"),
        ("报告规范和评价框架研究也在快速完善", "[14,29-30]", "报告规范和评价框架研究也在快速完善"),
        ("Hager等", "[12]", "Hager等"),
        ("Qiu等", "[24]", "Qiu等"),
        ("Chen等", "[10]", "Chen等"),
        ("Amugongo等的系统综述", "[22]", "Amugongo等"),
        ("Almanac等工作", "[23]", "Almanac等"),
        ("医疗数据脱敏研究", "[28]", "医疗数据脱敏研究"),
        ("医疗不良事件分类研究", "[27]", "医疗不良事件分类研究"),
        ("这类工作说明，医疗系统中的模型评价必须关联数据安全、任务标准和人工复核", "[25-28,32,36-37,40]", "人工复核"),
        ("医疗大语言模型的研究重点已经从", "[6-9,25-26,33-35,38-39]", "应用综述"),
        ("知识来源和输入环境本身可能成为攻击面", "[31]", "攻击面"),
        ("RAG研究则解决外部知识如何进入生成流程的问题。", "[15-23]", None),
    ]
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "参考文献":
            break
        for anchor_text, marker, anchor in citation_rules:
            if anchor_text in paragraph.text and marker not in paragraph.text:
                add_citation(paragraph, marker, anchor)

    for paragraph in doc.paragraphs:
        if "医疗大语言模型的数据投毒研究" in paragraph.text and "[41-44]" not in paragraph.text:
            replace_paragraph_substrings(
                paragraph,
                [("因此，本文在系统中", "与此同时，医疗智能化还涉及临终关怀伦理、健康治理、算法偏见和安全治理等制度性问题[41-44]。因此，本文在系统中")],
            )

    # Ensure every figure has a readable lead-in immediately before its caption.
    figure_captions = []
    for index, paragraph in enumerate(doc.paragraphs):
        t = paragraph.text.strip()
        if re.match(r"^图\d+-\d+\s", t) and len(t) < 80:
            figure_captions.append((index, t.split()[0]))
    for index, label in figure_captions:
        if index == 0:
            continue
        previous = index - 1
        while previous >= 0 and not doc.paragraphs[previous].text.strip():
            previous -= 1
        if previous < 0:
            continue
        lead = f"如{label}所示"
        previous_normalized = re.sub(r"\s+", "", doc.paragraphs[previous].text)
        lead_normalized = re.sub(r"\s+", "", lead)
        if lead_normalized not in previous_normalized:
            sentence = doc.paragraphs[previous].text
            if sentence.endswith("。"):
                sentence = sentence[:-1] + f"，{lead}。"
            else:
                sentence += f"，{lead}。"
            set_paragraph_text(doc.paragraphs[previous], sentence)

    table_captions = []
    for index, paragraph in enumerate(doc.paragraphs):
        t = paragraph.text.strip()
        if re.match(r"^表\d+-\d+\s", t) and len(t) < 100:
            table_captions.append((index, t.split()[0]))
    for index, label in table_captions:
        previous = index - 1
        while previous >= 0 and not doc.paragraphs[previous].text.strip():
            previous -= 1
        if previous < 0:
            continue
        paragraph = doc.paragraphs[previous]
        if paragraph.text.strip().startswith(tuple(str(i) for i in range(1, 7))):
            continue
        normalized = re.sub(r"\s+", "", paragraph.text)
        if label not in normalized:
            sentence = paragraph.text
            if sentence.endswith("。"):
                sentence = sentence[:-1] + f"，相关内容如{label}所示。"
            else:
                sentence += f"，相关内容如{label}所示。"
            set_paragraph_text(paragraph, sentence)

    # The source document keeps the references as plain paragraphs. Numbering
    # them here makes the in-text numeric citations resolvable without changing
    # their order or bibliographic content.
    in_references = False
    reference_number = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "参考文献":
            in_references = True
            continue
        if text == "致谢":
            in_references = False
        if in_references and text:
            reference_number += 1
            if not re.match(r"^\[\d+\]\s", text):
                set_paragraph_text(paragraph, f"[{reference_number}] {text}")

    # Final pass for implementation paragraphs whose technical anchors may be
    # split across runs by Word. This is intentionally after citation work so
    # the rewritten prose is not touched again by generic substitutions.
    final_anchor_rewrites = {
        "KnowledgeView.vue": "`KnowledgeView.vue` 提供知识统计、文档列表、版本管理、版本差异和入库表单。入库表单支持文件模式和正文模式，要求填写科室、机构、题名、来源链接、发布日期、版本和许可信息；审核状态由后端审核接口最终控制。页面还提供版本构建、版本激活和文档审核操作，使知识生命周期能够在管理端闭环完成。",
        "MonitorView.vue": "`MonitorView.vue` 维护健康状态、历史 Trace、实时 Trace、节点状态和事件列表。页面通过健康接口获取模型服务、知识索引和治理状态，通过 Trace 列表查询历史运行，通过事件流定位当前活动。节点列表展示安全筛查、症状采集、追问、知识检索、辅助分诊和回答编排的状态。",
        "AuditLogInterceptor": "`AuditLogInterceptor` 为受保护 Web 请求记录操作者、角色、HTTP 方法、动作、响应状态、请求标识、地址摘要和耗时，并将成功或失败结果写入审计日志。监控与审计页面分别提供查询和导出入口，不同角色只能读取其权限范围内的记录。",
    }
    for paragraph in doc.paragraphs:
        for anchor, new in final_anchor_rewrites.items():
            if anchor in paragraph.text and paragraph.text != new:
                set_paragraph_text(paragraph, new)


def main() -> None:
    doc = Document(str(INPUT))
    rewrite_target_paragraphs(doc)
    add_references_and_polish(doc)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
