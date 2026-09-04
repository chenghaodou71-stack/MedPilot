from __future__ import annotations

from copy import deepcopy
import json
from pathlib import Path
import statistics

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"D:\毕设制作")
SOURCE = ROOT / "outputs" / (
    "7230264110_窦成皓_23软件1_基于多智能体协同与大语言模型的"
    "多专科医疗健康咨询及辅助分诊系统_第四章图文增强版.docx"
)
OUTPUT = ROOT / "outputs" / (
    "7230264110_窦成皓_23软件1_基于多智能体协同与大语言模型的"
    "多专科医疗健康咨询及辅助分诊系统_千例评测增强版.docx"
)
RESULTS_DIR = ROOT / "ai-service" / "evaluation-results-1000"
FIGURE_DIR = ROOT / "outputs" / "thesis-evaluation-1000"


FIGURES = {
    "图5-1 各类危险信号召回率": (
        "图5-1 不同扰动强度下的任务性能",
        FIGURE_DIR / "figure5-1_perturbation_robustness.png",
        5.80,
    ),
    "图5-2 各科室分诊F1值": (
        "图5-2 科室与回退类别混淆矩阵",
        FIGURE_DIR / "figure5-2_department_confusion_matrix.png",
        4.80,
    ),
    "图5-3 各风险等级F1值": (
        "图5-3 不同测试路径的离线组件耗时分布",
        FIGURE_DIR / "figure5-3_offline_latency_boxplot.png",
        5.80,
    ),
}


def exact_paragraph(doc: Document, text: str) -> Paragraph:
    matches = [paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"expected one paragraph {text!r}, found {len(matches)}")
    return matches[0]


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_run_font(run, size: float = 12.0) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")


def replace_paragraph(doc: Document, old: str, new: str) -> Paragraph:
    paragraph = exact_paragraph(doc, old)
    run_properties = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run_properties = deepcopy(paragraph.runs[0]._r.rPr)
    clear_paragraph(paragraph)
    run = paragraph.add_run(new)
    if run_properties is not None:
        run._r.insert(0, run_properties)
    set_run_font(run, 12.0)
    return paragraph


def replace_caption_text(paragraph: Paragraph, new: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new
        for run in paragraph.runs[1:]:
            run.text = ""
        set_run_font(paragraph.runs[0], 10.5)
    else:
        run = paragraph.add_run(new)
        set_run_font(run, 10.5)


def replace_figure(doc: Document, old_caption: str, new_caption: str, image: Path, width: float) -> None:
    if not image.is_file():
        raise FileNotFoundError(image)
    caption = exact_paragraph(doc, old_caption)
    image_element = caption._p.getprevious()
    while image_element is not None and not image_element.xpath(".//w:drawing"):
        image_element = image_element.getprevious()
    if image_element is None:
        raise RuntimeError(f"no figure found before {old_caption}")
    image_paragraph = Paragraph(image_element, caption._parent)
    clear_paragraph(image_paragraph)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.first_line_indent = Pt(0)
    image_paragraph.paragraph_format.keep_with_next = True
    run = image_paragraph.add_run()
    shape = run.add_picture(str(image), width=Inches(width))
    shape._inline.docPr.set("descr", new_caption)
    replace_caption_text(caption, new_caption)


def set_cell_text(cell, text: str, *, center: bool) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    run_properties = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run_properties = deepcopy(paragraph.runs[0]._r.rPr)
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)
    set_run_font(run, 10.5)


def update_distribution_table(doc: Document) -> None:
    matches = [
        table for table in doc.tables
        if table.rows and table.rows[0].cells[0].text.strip() == "测试类别"
    ]
    if len(matches) != 1:
        raise RuntimeError(f"expected one test distribution table, found {len(matches)}")
    table = matches[0]
    rows = (
        ("高风险危险信号", "320", "8类危险信号，每类40条；5级扰动各64条"),
        ("证据支持的中风险", "320", "4个受支持科室，每科80条；5级扰动各64条"),
        ("低证据回退", "200", "10类知识库外症状；5级扰动各40条"),
        ("信息不足追问", "160", "无症状或单症状缺字段；5级扰动各32条"),
        ("合计", "1000", "固定种子20260831，逐条保存预测、证据与耗时"),
    )
    if len(table.rows) != len(rows) + 1:
        raise RuntimeError("unexpected test distribution table row count")
    for row, values in zip(table.rows[1:], rows):
        for index, value in enumerate(values):
            set_cell_text(row.cells[index], value, center=index < 2)


def update_figure_list(doc: Document) -> None:
    replacements = {
        "各类危险信号召回率": "不同扰动强度下的任务性能",
        "各科室分诊F1值": "科室与回退类别混淆矩阵",
        "各风险等级F1值": "不同测试路径的离线组件耗时分布",
    }
    table = doc.tables[0]
    found = set()
    for row in table.rows[1:]:
        old = row.cells[1].text.strip()
        if old in replacements:
            set_cell_text(row.cells[1], replacements[old], center=False)
            found.add(old)
    if found != set(replacements):
        raise RuntimeError(f"figure list replacements incomplete: {found}")


def enable_field_updates(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def ratio(value: float) -> str:
    return f"{value:.2%}"


def percentile(values: list[float], level: float) -> float:
    ordered = sorted(values)
    position = (len(ordered) - 1) * level
    lower = int(position)
    upper = min(lower + 1, len(ordered) - 1)
    weight = position - lower
    return ordered[lower] + (ordered[upper] - ordered[lower]) * weight


def latency_summary(outcomes: list[dict], predicate) -> tuple[int, float, float]:
    values = [float(row["latency_ms"]) for row in outcomes if predicate(row) and not row["error"]]
    return len(values), statistics.median(values), percentile(values, 0.95)


def build_replacements(metrics: dict, outcomes: list[dict]) -> dict[str, str]:
    safety = metrics["safety"]
    department = metrics["department"]
    supported = department["supported_only"]
    risk = metrics["risk"]
    retrieval = metrics["retrieval"]
    fallback = metrics["fallback_and_followup"]
    perturbation = metrics["perturbation_metrics"]
    matched_count, matched_p50, matched_p95 = latency_summary(
        outcomes,
        lambda row: row["task"] == "high_risk" and row["predicted_route"] == "high_risk_fast_path",
    )
    missed_count, missed_p50, missed_p95 = latency_summary(
        outcomes,
        lambda row: row["task"] == "high_risk" and row["predicted_route"] != "high_risk_fast_path",
    )
    evidence_latency = metrics["latency_by_task"]["evidence_triage"]
    low_latency = metrics["latency_by_task"]["low_evidence"]
    followup_latency = metrics["latency_by_task"]["insufficient_followup"]

    return {
        "测试输入分为高风险、证据支持、低证据和信息不足四类。高风险输入用于验证模型调用前的确定性安全筛查；证据支持输入用于验证bge-m3、FAISS和科室加权分诊；低证据输入用于验证没有有效知识证据时的线下分诊回退；信息不足输入用于验证系统是否提出追问而不是强行分诊。每条输入先独立记录金标准，再运行系统，避免用系统输出反向构造测试标签。":
            "测试输入分为高风险、证据支持、低证据和信息不足四类。测试集由固定种子20260831生成并在运行前冻结，逐条记录金标准、扰动等级和预期路径。高风险文本调用生产安全筛查；证据支持与低证据文本调用现有bge-m3、FAISS和证据加权分类；信息不足样本调用结构化症状充分性与追问规则。自由文本回答生成不计入分类指标，避免把随机生成质量混入确定性组件评测。",
        "信息不足样本不纳入科室和风险F1计算，而单独统计追问触发率。检索模块额外统计是否返回可用证据和证据引用标识。上述指标只描述当前Demo在构造工程测试集上的行为，不用于推断临床准确率。":
            "信息不足样本不纳入科室和风险F1，而单独统计追问触发率。证据支持样本统计Recall@3与MRR；科室指标分别报告四个受支持科室的Macro-F1，以及加入全科/线下回退类别后的五分类Macro-F1。组件耗时从函数实际开始执行到返回结果为止，不包含并发队列、浏览器、业务后端、数据库和网络传输，因此不等同于线上端到端延迟。上述指标只描述当前Demo在人工构造工程测试集上的行为，不用于推断临床准确率。",
        "本次扩展工程测试集共240条，具体分布如表5-2。":
            "本次固定种子工程测试集共1000条，具体分布如表5-2。",
        "高风险样本覆盖胸痛、呼吸困难、气促、咯血、便血、晕厥、意识不清和大出血。每类样本同时加入口语化表达、症状组合、时间描述和否定词，测试规则是否能区分“出现危险信号”和“否认危险信号”。中风险样本围绕当前知识库中的四个证据科室编写，输入包含至少两个症状，或包含一个症状及持续时间、严重程度信息。低证据和信息不足样本用于验证系统的保守分支。":
            "高风险样本覆盖胸痛、呼吸困难、气促、咯血、便血、晕厥、意识不清和大出血，每类40条。全部样本按P0至P4划分为规范表达、口语与同义表达、组合与背景干扰、否定与跨症状干扰、错别字与未登录改写五级。证据支持样本围绕心血管内科、呼吸内科、消化内科和皮肤科构造；低证据样本来自当前知识库未覆盖的骨科、耳鼻喉、眼科、口腔、泌尿等症状；信息不足样本只保留空症状或单症状且不提供持续时间和严重程度。",
        "安全筛查节点在所有普通模型调用之前执行。测试重点包括规范词匹配、同义词归一化、标点切分、转折词和否定状态继承。测试结果显示，胸痛、呼吸困难、气促、咯血、晕厥、意识不清和大出血在本次样本中均被正确识别；便血类别有一条输入使用“鲜红色血液”表达，当前词表未将其映射为便血，因而形成一条漏检记录。该结果反映的是规则词表覆盖范围，而不是系统对所有出血表现的医学识别能力。":
            f"安全筛查节点在普通模型调用前执行。320条高风险样本中，P0至P3共256条全部命中，P4的64条错别字或未登录改写全部漏检，总体召回率为{ratio(safety['recall'])}。其余680条非高风险样本没有触发危险信号，工程测试特异度为{ratio(safety['specificity'])}。该结果说明否定与转折处理在已登记词形上保持稳定，但规则对词表外语义改写缺少泛化能力。",
        "图5-1展示各类危险信号召回率。除便血类别为91.67%外，其余类别为100%，总体召回率为98.96%。这个结果支持两点判断：前置规则能够稳定覆盖已纳入词表的主要表达，同时规则系统仍需要通过同义词扩展和人工审核不断补充词表。论文不将该结果写成临床急症筛查能力。":
            f"图5-1按扰动等级比较四条测试路径。P0至P3的危险信号召回率均为100%，P4下降为0；证据分诊准确率在P3为{ratio(perturbation[3]['department_accuracy'])}，其余等级为{ratio(perturbation[0]['department_accuracy'])}；主动追问始终为100%，低证据回退始终为0。不同曲线的分母分别为每级64、64、32和40条，不能将曲线间高低直接解释为同一临床能力。",
        "信息充分性规则根据症状数量、持续时间和严重程度决定是否进入检索。信息不足样本全部触发追问，追问触发率为100%。系统能够返回“还缺少持续时间”“还需要说明其他症状”或“需要补充严重程度”等具体缺失字段，前端据此显示追问状态。":
            f"信息充分性规则根据症状数量、持续时间和严重程度决定是否进入检索。160条信息不足样本全部触发追问，触发率为{ratio(fallback['followup_trigger_rate'])}。该结果验证的是结构化症状对象上的纯规则行为；本轮没有把qwen2.5:7b自由抽取结果并入该指标，因此不能据此声称任意口语输入都能被正确抽取。",
        "追问测试还验证了一个重要边界：没有危险信号不等于信息已经充分。输入“最近不舒服”不会被当作低风险完成，而是要求用户补充主要不适；输入“最近咳嗽”在缺少持续时间和程度时也不会直接生成科室建议。该机制减少了因输入过短而产生的无依据分诊。":
            "追问测试验证了一个工程边界：没有危险信号不等于信息已经充分。空症状对象会询问主要不适；单症状且缺少持续时间和程度时，会优先询问持续时间或其他症状。该分支在固定结构化输入上能够阻止提前分诊，但真实使用效果仍取决于上游症状抽取是否完整。",
        "64条证据支持样本均返回至少一条当前知识索引证据，检索命中率为100%。检索器将向量相似度和字符二元组重合度结合，输出Top-K证据，并把科室、来源、URL、发布日期、版本和审核状态保存在证据对象中。分诊函数对四个受支持科室进行加权投票，最终回答使用结构化字段确定性渲染。":
            f"320条证据支持样本均返回至少一条索引证据，证据返回率为{ratio(retrieval['evidence_return_rate'])}；但Top-3中包含预先标注的可接受文档的比例为{ratio(retrieval['recall_at_3'])}，MRR@3为{retrieval['mrr_at_3']:.4f}。这说明“返回了证据”与“返回了目标证据”不是同一指标。检索结果继续按科室分数加权，回答层只渲染结构化分诊字段与通过治理检查的引用。",
        "科室分诊F1如图5-2所示。皮肤科、神经内科和急诊科在本次标签分布中达到100%，呼吸内科F1为98.11%，心血管内科为97.50%，全科/建议线下分诊台为98.97%，消化内科为94.34%，科室Macro-F1为0.9842。消化内科得分较低，主要与输入表达和检索证据的区分度有关；这也说明单纯增加证据数量不一定能解决相近症状之间的分类边界问题。":
            f"图5-2给出320条受支持科室样本与200条低证据样本的五分类混淆矩阵。仅在四个受支持科室内统计时，准确率为{ratio(supported['accuracy'])}，Macro-F1为{supported['macro_f1']:.4f}；加入回退类别后，准确率降至{ratio(department['accuracy'])}，Macro-F1为{department['macro_f1']:.4f}。200条知识库外症状全部被分配到四个受支持科室，回退类别召回率为0，说明当前相似度阈值没有形成有效的分布外拒答边界。",
        "图5-1的横坐标为8类危险信号，纵坐标为召回率。胸痛、呼吸困难、气促、咯血、晕厥、意识不清和大出血召回率均为1.0000，便血召回率为0.9167。总体召回率由96条高风险样本计算，结果为0.9896。":
            f"图5-1的横坐标为五个扰动等级，纵坐标为各路径的样本级比例。八类危险信号每类40条，均表现为P0至P3命中32条、P4漏检8条，因此各类别召回率均为0.8000，总体召回率也是{safety['recall']:.4f}。一致的80%并不表示类别难度完全相同，而是本测试设计为每类固定配置8条未登录改写。",
        "漏检样本“排便后看到鲜红色血液”没有命中当前便血词表，系统转入全科/线下分诊台低风险回退。该样本被保留在结果中，原因是测试集需要暴露规则边界，而不是只保留能够通过的输入。后续词表扩展可以增加“鲜红色血液”等需要结合消化道出血语境判断的表达，但新增词条仍需经过人工审核，避免把其他部位出血误判为便血。":
            "64条漏检样本包括“胸口像被压住一样疼”“神志模糊叫不醒”“伤口一直血流不止”等未包含规范词形的表达。漏检后，这些文本进入普通检索路径，无法恢复为高风险标签。后续不能只机械加入字符串，还需要结合语境约束、同义归一化与人工复核，防止扩大词表后引入新的误报。",
        "5.4.2 风险等级统计": "5.4.2 风险等级与处理耗时统计",
        "风险等级F1如图5-3所示。低风险F1为98.97%，中风险F1为100%，高风险F1为99.48%，风险Macro-F1为0.9948。中风险样本均由受支持科室证据进入检索分诊，低风险样本主要进入无证据回退，高风险样本主要由确定性规则直接处理。":
            f"在排除160条追问样本后，风险标签统计覆盖840条输入。低、中、高风险F1分别为{risk['per_label']['低']['f1']:.4f}、{risk['per_label']['中']['f1']:.4f}和{risk['per_label']['高']['f1']:.4f}，Macro-F1为{risk['macro_f1']:.4f}，准确率为{ratio(risk['accuracy'])}。低风险F1为0源于低证据回退未触发，高风险召回损失源于64条P4改写漏检。图5-3进一步比较不同路径的离线组件耗时。",
        "风险结果说明，当前系统的风险标签与工程路径具有较强对应关系：危险信号触发高风险，证据支持触发中风险，证据不足触发低风险或暂缓判断。风险等级并非疾病严重程度的医学量表，而是系统用于安排回答时效和路径的工程标签。第4章中的前置规则、证据投票和abstained字段共同决定了这一标签。":
            f"图5-3显示，{matched_count}条规则命中样本的P50和P95耗时分别为{matched_p50:.3f} ms和{matched_p95:.3f} ms；{missed_count}条高危漏检样本进入检索后，P50和P95增至{missed_p50:.3f} ms和{missed_p95:.3f} ms。证据支持与低证据测试的P50分别为{evidence_latency['p50_ms']:.3f} ms和{low_latency['p50_ms']:.3f} ms，主动追问规则P50为{followup_latency['p50_ms']:.3f} ms。图中使用对数坐标，数据只表示本机离线组件处理时间。",
        "检索测试检查空查询、索引缺失、索引损坏、低分过滤、Top-K截取、同义词扩展和活动版本切换。空查询直接返回空列表，避免对无意义文本调用嵌入模型；索引缺失或损坏时健康接口返回不可用状态，普通咨询不会伪造证据；综合分低于阈值的分片被过滤，避免无关文档进入分诊。":
            f"检索测试检查空查询、索引缺失、低分过滤、Top-K截取、同义词扩展和活动版本切换。空查询与索引异常的失败关闭行为通过既有自动化测试验证；千例评测则暴露了阈值边界：200条知识库外症状均获得高于当前阈值的候选证据，低证据回退率为{ratio(fallback['low_evidence_abstain_rate'])}。因此，阈值过滤在本次负样本上没有实现预期的无关证据拒绝。",
        "引用测试检查citation_id、doc_id、chunk_id、source、url、published_date、version、license和review_status是否完整。证据引用必须来自审核通过文档，URL必须使用HTTPS，发布日期必须能够按ISO格式解析；包含确诊、处方、用药和剂量内容的证据会被二次过滤。记录详情页面展示引用来源、原文片段、分片标识和索引版本，监控页面也能通过Trace查看本次引用数量。":
            f"引用测试继续检查citation_id、doc_id、chunk_id、source、url、published_date、version、license和review_status是否完整。320条证据支持样本均返回结构化引用，但Recall@3为{ratio(retrieval['recall_at_3'])}，表明元数据完整不能替代相关性判断。记录详情页和Trace可以追溯本次引用来源与索引版本，评测文件同时保存每条样本的Top-3文档标识和分数。",
        "科室Macro-F1达到0.9842，说明在当前扩展测试集上，知识证据投票与危险信号快速通道能够覆盖主要科室路径。皮肤科的证据表达较集中，分诊F1为1.0000；消化内科的F1为0.9434，是七个标签中最低的类别。消化道症状在输入文本中可能与一般腹部不适、饮食相关不适或出血表现相互接近，检索结果的科室分布会受到查询词和分片内容共同影响。":
            f"在预先限定为四个受支持科室的320条样本上，科室Macro-F1为{supported['macro_f1']:.4f}。呼吸内科F1最高，为{supported['per_label']['呼吸内科']['f1']:.4f}；心血管内科F1最低，为{supported['per_label']['心血管内科']['f1']:.4f}。然而加入200条知识库外症状后，五分类Macro-F1降至{department['macro_f1']:.4f}。前一结果描述封闭类别内的区分能力，后一结果描述开放输入下的回退能力，两者不能相互替代。",
        "风险Macro-F1达到0.9948，低、中、高三个等级均获得较高F1。高风险F1略低于1.0000，原因是个别危险信号表达没有覆盖在当前规则词表内；低风险F1略低于1.0000，原因是部分一般不适文本与全科回退路径存在证据边界。中风险样本全部按预期进入证据支持路径，但这不代表系统能够对所有中风险疾病做出可靠判断。":
            f"风险Macro-F1为{risk['macro_f1']:.4f}，没有达到旧测试中的近满分水平。高风险F1为{risk['per_label']['高']['f1']:.4f}，反映未登录表达的漏检；中风险F1为{risk['per_label']['中']['f1']:.4f}，其精度受到低证据样本被错误接纳的影响；低风险F1为0。该结果表明当前风险标签高度依赖规则词表与证据阈值，尚不能在开放输入上形成稳定的高、中、低路径划分。",
        "测试结果最直接的价值是验证系统设计是否按预定路径运行。高风险样本是否绕过普通模型流程，信息不足样本是否触发追问，证据支持样本是否返回引用，低证据样本是否保守回退，权限和异常状态是否能被持久化，均可以通过本次结果检查。统计图因此用于展示工程行为的稳定性和当前规则覆盖情况。":
            "千例评测同时验证了成功路径和失败路径。已登记危险词形能够绕过普通模型流程，结构化信息不足能够触发追问，受支持科室在封闭标签内保持较高区分度；但未登录危险表达不会触发快速通道，知识库外症状也没有进入预期回退。统计图用于展示这些工程行为及其边界，不把失败样本删除，也不把组件指标解释为医疗效果。",
        "结果也暴露出明确的改进方向。第一，危险信号词表需要继续补充口语表达和语境约束。第二，消化内科相关知识分片需要进一步评估查询扩展和科室权重。第三，信息不足场景可以继续优化追问排序，让问题更快获得对分诊有价值的字段。第四，知识版本、引用质量和人工复核应纳入后续持续回归，而不能只依靠一次测试集结果。":
            "结果给出四项优先改进方向。第一，为危险信号增加经审核的同义归一化与语境识别，并用未登录表达持续回归。第二，使用知识库外硬负样本重新标定检索阈值，或增加独立的分布外检测与拒答判据。第三，针对心血管内科易混样本和Recall@3未命中样本检查查询扩展、分片内容和排序权重。第四，将上游症状抽取、知识版本、引用相关性和人工复核纳入后续端到端评测。",
        "本章建立了覆盖高风险规则、证据检索、低证据回退和主动追问的扩展工程测试集，并从功能、流程、安全、权限、异常和知识追踪等方面验证系统。240条测试输入中，高风险样本总体危险信号召回率为98.96%，科室Macro-F1为0.9842，风险Macro-F1为0.9948，信息不足样本追问触发率为100%，证据支持样本检索命中率为100%。测试还确认了事件断流、权限越界、密文篡改、索引损坏和AI服务不可用等异常路径能够进入可追踪状态。上述结果用于说明MedPilot当前Demo的工程实现情况，不作为临床准确率或医疗效果结论。":
            f"本章使用固定种子生成并执行1000条人工构造工程测试。危险信号召回率为{ratio(safety['recall'])}，四个受支持科室的Macro-F1为{supported['macro_f1']:.4f}，加入回退类别后的五分类Macro-F1为{department['macro_f1']:.4f}，风险Macro-F1为{risk['macro_f1']:.4f}，Recall@3为{ratio(retrieval['recall_at_3'])}，追问触发率为{ratio(fallback['followup_trigger_rate'])}，低证据回退率为{ratio(fallback['low_evidence_abstain_rate'])}。结果说明当前Demo在封闭科室分诊和结构化追问上较稳定，但在未登录危险表达与知识库外拒答方面存在明确缺口。上述结果只用于工程验证，不作为临床准确率或医疗效果结论。",
    }


def build() -> None:
    if not SOURCE.is_file():
        raise FileNotFoundError(SOURCE)
    metrics = json.loads((RESULTS_DIR / "metrics.json").read_text(encoding="utf-8"))
    outcomes = json.loads((RESULTS_DIR / "case_outcomes.json").read_text(encoding="utf-8"))
    if metrics["benchmark"]["case_count"] != 1000 or len(outcomes) != 1000:
        raise RuntimeError("benchmark outputs are incomplete")
    if metrics["benchmark"]["seed"] != 20260831 or metrics["errors"]["count"] != 0:
        raise RuntimeError("benchmark seed or error count is unexpected")

    doc = Document(str(SOURCE))
    source_shapes = len(doc.inline_shapes)
    replacements = build_replacements(metrics, outcomes)
    for old, new in replacements.items():
        replace_paragraph(doc, old, new)

    for old_caption, (new_caption, image, width) in FIGURES.items():
        replace_figure(doc, old_caption, new_caption, image, width)
    update_distribution_table(doc)
    update_figure_list(doc)
    enable_field_updates(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)

    check = Document(str(OUTPUT))
    if len(check.inline_shapes) != source_shapes:
        raise RuntimeError(f"inline shape count changed: {len(check.inline_shapes)} != {source_shapes}")
    for _old, (caption, _image, _width) in FIGURES.items():
        exact_paragraph(check, caption)
    combined = "\n".join(paragraph.text for paragraph in check.paragraphs)
    stale = ("共240条", "98.96%", "0.9842", "0.9948", "各风险等级F1值")
    leftovers = [value for value in stale if value in combined]
    if leftovers:
        raise RuntimeError(f"stale chapter 5 values remain: {leftovers}")
    print(f"saved={OUTPUT}; replacements={len(replacements)}; inline_shapes={len(check.inline_shapes)}")


if __name__ == "__main__":
    build()
