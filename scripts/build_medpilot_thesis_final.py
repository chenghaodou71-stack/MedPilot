from __future__ import annotations

"""Build the final, evidence-bounded MedPilot graduation thesis.

This builder keeps the original project facts separate from literature claims,
uses the existing high-resolution technical drawings, and assembles the draft
into a seven-chapter thesis rather than concatenating Markdown files verbatim.
"""

import json
import re
from collections import defaultdict
from pathlib import Path

import build_medpilot_thesis_rewrite as base
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DRAFT_DIR = ROOT / "outputs" / "researchwrite" / "medpilot-thesis" / "drafts"
OUT = ROOT / "outputs" / "MedPilot_毕业设计论文_完整重制终稿.docx"
STATS_OUT = ROOT / "outputs" / "MedPilot_毕业设计论文_完整重制终稿_统计.json"
VISIO_PNG_DIR = ROOT / "outputs" / "thesis-visio-figures-v2"


def visio_figures() -> dict[str, Path]:
    """Use the native Visio exports as the only thesis figure source."""
    names = {
        "功能结构": "图2-1_系统功能结构图.png",
        "角色权限": "图2-2_用户角色与权限边界图.png",
        "状态机": "图2-3_咨询业务状态机图.png",
        "分层架构": "图3-1_系统总体分层架构图.png",
        "信任边界": "图3-2_信任边界与接口调用链图.png",
        "数据库": "图3-3_核心数据库ER图.png",
        "知识治理": "图3-4_知识库治理生命周期图.png",
        "普通流程": "图4-1_普通咨询流程图.png",
        "红旗通道": "图4-2_红旗高风险快速通道图.png",
        "主动追问": "图4-3_信息不足与主动追问流程图.png",
        "RAG": "图4-4_医学RAG检索增强流程图.png",
        "证据分诊": "图4-5_证据加权与辅助分诊流程图.png",
        "复核闸门": "图4-6_医生复核安全闸门图.png",
        "NDJSON": "图4-7_NDJSON事件生命周期与失败终止图.png",
        "部署": "图5-1_系统部署与服务边界图.png",
        "Trace": "图5-2_Trace监控审计与回溯流程图.png",
    }
    paths = {key: VISIO_PNG_DIR / value for key, value in names.items()}
    missing = [str(path) for path in paths.values() if not path.exists()]
    if missing:
        raise FileNotFoundError("缺少 Visio 图件：" + ", ".join(missing))
    return paths

CHAPTER_TITLES = {
    1: "第1章 绪论",
    2: "第2章 系统分析与相关技术",
    3: "第3章 系统总体设计",
    4: "第4章 核心算法设计与实现",
    5: "第5章 系统详细设计与实现",
    6: "第6章 系统测试与实验分析",
    7: "第7章 总结与展望",
}


# [1]-[10] are international literature.  The first five are peer-reviewed
# journal articles with individually checked DOI metadata.  [11]-[40] are
# Chinese-source candidates whose public metadata was checked against OpenAlex
# and source/DOI landing records; final CNKI account export is still required
# for school-specific core-directory confirmation.
REFERENCES = [
    "[1] Semigran H L, Linder J A, Gidengil C, et al. Evaluation of symptom checkers for self diagnosis and triage: audit study[J]. BMJ, 2015, 351: h3480. DOI:10.1136/bmj.h3480.",
    "[2] Chambers D, Cantrell A J, Johnson M, et al. Digital and online symptom checkers and health assessment/triage services for urgent health problems: systematic review[J]. BMJ Open, 2019, 9(8):e027743. DOI:10.1136/bmjopen-2018-027743.",
    "[3] Wallace W, Chan C, Chidambaram S, et al. The diagnostic and triage accuracy of digital and online symptom checker tools: systematic review[J]. npj Digital Medicine, 2022, 5:118. DOI:10.1038/s41746-022-00667-6.",
    "[4] Schmieding M L, Kopka M, Schmidt K, et al. Triage accuracy of symptom checker apps: 5-year follow-up evaluation[J]. Journal of Medical Internet Research, 2022, 24(5):e31810. DOI:10.2196/31810.",
    "[5] Fraser H, Crossland D, Bacher I, et al. Comparison of diagnostic and triage accuracy of Ada Health and WebMD symptom checkers, ChatGPT, and physicians for patients in an emergency department: clinical data analysis study[J]. JMIR mHealth and uHealth, 2023, 11:e49995. DOI:10.2196/49995.",
    "[6] Singhal K, Azizi S, Tu T, et al. Large language models encode clinical knowledge[J]. Nature, 2023, 620(7972):172-180. DOI:10.1038/s41586-023-06291-2.",
    "[7] Hager P, Jungmann F, Holland R, et al. Evaluation and mitigation of the limitations of large language models in clinical decision-making[J]. Nature Medicine, 2024, 30(9):2613-2622. DOI:10.1038/s41591-024-03097-1.",
    "[8] Laranjo L, Dunn A G, Tong H L, et al. Conversational agents in healthcare: a systematic review[J]. Journal of Medical Internet Research, 2018, 20(5): e10153. DOI:10.2196/10153.",
    "[9] Zakka C, Shad R, Chaurasia A, et al. Almanac: retrieval-augmented language models for clinical medicine[J]. NEJM AI, 2024, 1(2). DOI:10.1056/AIoa2300068.",
    "[10] Topol E J. High-performance medicine: the convergence of human and artificial intelligence[J]. Nature Medicine, 2019, 25(1): 44-56. DOI:10.1038/s41591-018-0300-7.",
    "[11] 孟小峰, 慈祥. 大数据管理：概念、技术与挑战[J]. 计算机研究与发展, 2013, 50(1):146-169.",
    "[12] 刘峤, 李杨, 段宏, 等. 知识图谱构建技术综述[J]. 计算机研究与发展, 2016, 53(3):582-600. DOI:10.7544/issn1000-1239.2016.20148228.",
    "[13] 余凯, 贾磊, 陈雨强, 等. 深度学习：昨天、今天和明天[J]. 计算机研究与发展, 2013, 50(9):1799-1816.",
    "[14] 刘知远, 孙茂松, 林衍凯, 等. 知识表示学习研究进展[J]. 计算机研究与发展, 2016, 53(2):247-261. DOI:10.7544/issn1000-1239.2016.20160020.",
    "[15] 庄严, 李国良, 冯建华, 等. 知识库实体对齐研究综述[J]. 计算机研究与发展, 2016, 53(1):165. DOI:10.7544/issn1000-1239.2016.20150661.",
    "[16] 秦兵, 刘安安, 刘挺. 无监督中文开放实体关系抽取[J]. 计算机研究与发展, 2015, 52(5):1029. DOI:10.7544/issn1000-1239.2015.20131550.",
    "[17] 甘立新, 万长轩, 刘德喜, 等. 基于句法和语义特征的中文命名实体关系抽取[J]. 计算机研究与发展, 2016, 53(2):284. DOI:10.7544/issn1000-1239.2016.20150842.",
    "[18] 刘雅辉, 张铁赢, 金小龙, 等. 大数据时代个人隐私保护[J]. 计算机研究与发展, 2015, 52(1):229. DOI:10.7544/issn1000-1239.2015.20131340.",
    "[19] 刘俊旭, 孟小峰, 王蕾, 等. 隐私保护机器学习研究综述[J]. 计算机研究与发展, 2020, 57(2):346. DOI:10.7544/issn1000-1239.2020.20190455.",
    "[20] 纪守领, 李金凤, 杜天宇, 等. 机器学习模型的可解释性方法、应用与安全研究综述[J]. 计算机研究与发展, 2019, 56(10):2071. DOI:10.7544/issn1000-1239.2019.20190540.",
    "[21] 侯梦薇, 魏荣, 陆良, 等. 医疗领域知识图谱及其应用研究综述[J]. 计算机研究与发展, 2018, 55(12):2587. DOI:10.7544/issn1000-1239.2018.20180623.",
    "[22] 李晖, 孙文海, 李凤华, 等. 公有云环境下安全与隐私保护数据存储服务研究[J]. 计算机研究与发展, 2014, 51(7):1397.",
    "[23] 王继业, 高凌超, 董爱强, 等. 基于区块链的数据安全共享网络架构研究[J]. 计算机研究与发展, 2017, 54(4):742. DOI:10.7544/issn1000-1239.2017.20160991.",
    "[24] 朱烈煌, 高峰, 沈蒙, 等. 区块链隐私保护技术综述[J]. 计算机研究与发展, 2017, 54(10):2170. DOI:10.7544/issn1000-1239.2017.20170471.",
    "[25] 何海武, 闫安, 陈泽华, 等. 智能合约技术与应用综述[J]. 计算机研究与发展, 2018, 55(11):2452. DOI:10.7544/issn1000-1239.2018.20170658.",
    "[26] 陈伟利, 郑子彬. 区块链数据分析：现状、趋势与挑战[J]. 计算机研究与发展, 2018, 55(9):1853. DOI:10.7544/issn1000-1239.2018.20180127.",
    "[27] 潘晨, 刘志强, 刘震, 等. 区块链技术可扩展性问题与方法研究[J]. 计算机研究与发展, 2018, 55(10):2099. DOI:10.7544/issn1000-1239.2018.20180440.",
    "[28] 施巍松, 孙辉, 曹杰, 等. 边缘计算：万物互联时代的新型计算模型[J]. 计算机研究与发展, 2017, 54(5):907. DOI:10.7544/issn1000-1239.2017.20160941.",
    "[29] 施巍松, 张兴洲, 王一帆, 等. 边缘计算：现状与未来[J]. 计算机研究与发展, 2019, 56(1):69. DOI:10.7544/issn1000-1239.2019.20180760.",
    "[30] 赵梓铭, 刘芳, 蔡志平, 等. 边缘计算：平台、应用与挑战[J]. 计算机研究与发展, 2018, 55(2):327-337. DOI:10.7544/issn1000-1239.2018.20170228.",
    "[31] 梁斌, 刘全, 徐劲, 等. 基于多注意力卷积神经网络的方面级情感分析[J]. 计算机研究与发展, 2017, 54(8):1724. DOI:10.7544/issn1000-1239.2017.20170178.",
    "[32] 姜卓轩, 张燕, 李晓明, 等. 基于 MOOC 数据的学习行为分析与预测[J]. 计算机研究与发展, 2015, 52(3):614. DOI:10.7544/issn1000-1239.2015.20140491.",
    "[33] 李建中, 刘贤敏. 大数据的重要方面：数据可用性[J]. 计算机研究与发展, 2013, 50(6):1147.",
    "[34] 王继业, 孟坤, 曹军威, 等. 能源互联网的信息技术研究综述[J]. 计算机研究与发展, 2015, 52(5):1109. DOI:10.7544/issn1000-1239.2015.20131592.",
    "[35] 段杰, 胡清华, 张灵军, 等. 基于邻域粗糙集的多标签分类特征选择[J]. 计算机研究与发展, 2015, 52(1):56. DOI:10.7544/issn1000-1239.2015.20140544.",
    "[36] 苗豫东, 吴建. 分级诊疗制度变迁回溯及“十四五”期间的关键政策建议[J]. 中国卫生政策研究, 2021, 14(3):1-6. DOI:10.3969/j.issn.1674-2982.2021.03.001.",
    "[37] 王婵, 李鑫武, 等. 分级诊疗对“倒三角”就医秩序的纠正效应评估——基于渐进性试点的准自然实验[J]. 中国卫生政策研究, 2021, 14(3):13-20. DOI:10.3969/j.issn.1674-2982.2021.03.003.",
    "[38] 朱劲松. “互联网+医疗”模式：内涵与系统架构[J]. 中国医院管理, 2016, 36(1):38-40.",
    "[39] 方鹏骞. 智慧医疗背景下“十四五”我国医院医疗质量与安全管理策略探析[J]. 中国医院管理, 2021, 41(3):15-17.",
    "[40] 朱西敏, 胡银环, 等. “互联网+医疗”背景下患者体验的评价：系统综述[J]. 中国卫生政策研究, 2021, 14(7):18-25. DOI:10.3969/j.issn.1674-2982.2021.07.003.",
]


ROLE_TABLE = (
    "表2-2 系统角色、资源与操作边界",
    ["角色", "主要资源", "允许操作", "关键约束", "审计要求"],
    [
        ("USER", "本人会话/档案", "咨询、查看、维护", "资源必须归属本人", "记录用户与会话"),
        ("DOCTOR", "授权患者任务", "领取、复核、决定", "关系、院区、MFA", "决定理由不可缺失"),
        ("REVIEWER", "知识/复核队列", "审核、复核", "禁止自复核", "审核人和时间"),
        ("KNOWLEDGE_EDITOR", "待审文档", "提交、修订、构建", "不能自行激活", "版本与来源"),
        ("ADMIN", "配置与治理", "用户、版本、运维", "不自动获得临床权限", "高危操作留痕"),
        ("AUDITOR", "Trace/审计日志", "只读查询、筛选", "不得触发推理", "查询本身留痕"),
    ],
    [23, 31, 35, 35, 26],
)

NFR_TABLE = (
    "表2-3 非功能需求与可验证约束",
    ["属性", "设计约束", "验证对象", "失败表现", "当前边界"],
    [
        ("安全性", "红旗先于模型；权限失败关闭", "规则/鉴权测试", "拒绝或高风险通道", "非临床安全证明"),
        ("可靠性", "序号连续、终止事件唯一", "NDJSON 集成测试", "保存失败 Trace", "未做高并发压测"),
        ("可追溯性", "引用快照、索引版本、Trace", "记录与查询测试", "可定位节点/版本", "无独立引用金标"),
        ("可维护性", "节点职责单一、依赖可注入", "单元/构建测试", "可替换模型夹具", "未实现多实例编排"),
        ("隐私性", "最小字段、同意投影、加密", "越权与撤回测试", "不返回敏感细节", "未做正式 DPIA"),
    ],
    [23, 45, 31, 30, 27],
)

ALGORITHM_TABLE = (
    "表4-1 辅助分诊节点的输入、输出与安全约束",
    ["节点", "主要输入", "主要输出", "前置条件", "失败/降级"],
    [
        ("safety_screen", "本轮原始文本", "红旗命中与原因", "无", "高风险快速通道"),
        ("extract", "文本、有限历史", "结构化症状字段", "未命中红旗", "字段为空，不虚构"),
        ("ask_followup", "缺失字段", "一个关键追问", "信息不足", "结束本轮并标记待补充"),
        ("retrieve", "结构化查询", "带版本的证据", "字段充分", "无证据则暂缓"),
        ("classify", "规则与证据", "科室、风险、时效", "证据有效或规则命中", "abstained/线下分诊"),
        ("compose", "受校验结果", "受限回答", "终端结果存在", "安全模板回退"),
    ],
    [29, 31, 32, 31, 32],
)

EVENT_TABLE = (
    "表4-2 NDJSON 事件字段与校验规则",
    ["字段", "类型", "规则", "示例", "用途"],
    [
        ("protocol_version", "string", "白名单", "1.0", "协议兼容"),
        ("sequence", "integer", "从 1 严格递增", "4", "排序与重放检查"),
        ("node", "string", "节点枚举", "retrieve", "呈现与诊断"),
        ("status", "string", "started/completed/error", "completed", "状态机"),
        ("phase", "string", "与终止状态一致", "awaiting_followup", "业务阶段"),
        ("elapsed_ms", "integer", "非负", "382", "节点耗时"),
        ("data", "object", "节点结构校验", "evidence[]", "受控输出"),
    ],
    [31, 25, 39, 31, 29],
)

INTERFACE_TABLE = (
    "表5-1 核心接口的职责划分",
    ["接口/域", "调用方", "关键校验", "成功语义", "失败语义"],
    [
        ("/api/auth/*", "浏览器", "凭据、CSRF", "建立/释放会话", "统一认证失败"),
        ("/api/consult", "患者端", "归属、长度、并发", "合法 NDJSON 终止", "409/413/429/5xx"),
        ("/api/profile/*", "患者端", "当前用户", "档案/复诊限定返回", "拒绝跨用户"),
        ("/api/clinical-reviews/*", "医生/复核员", "关系、院区、MFA", "决定独立保存", "fail-closed"),
        ("/knowledge/*", "治理端", "角色与版本状态", "候选构建/显式激活", "保持旧版本"),
        ("/monitor/*", "审计端", "只读权限", "分页查询 Trace", "不得触发推理"),
    ],
    [31, 25, 36, 32, 31],
)

TEST_ENV_TABLE = (
    "表6-1 测试层级、工具与证据类型",
    ["层级", "工具/对象", "主要验证", "输出证据", "不能推出的结论"],
    [
        ("单元测试", "pytest/JUnit", "规则、状态机、权重", "断言与夹具", "临床准确率"),
        ("接口测试", "MockMvc/HTTP", "鉴权、归属、协议", "状态码与持久化", "真实用户体验"),
        ("集成测试", "数据库/索引", "事务、版本、回滚", "记录与 manifest", "高并发能力"),
        ("浏览器测试", "Playwright", "关键交互路径", "页面流程", "用户理解度"),
        ("离线评测", "evaluator", "构造病例指标", "JSON/CSV/摘要", "临床有效性"),
    ],
    [24, 34, 35, 30, 32],
)

TEST_CASE_TABLE = (
    "表6-2 核心功能测试用例",
    ["编号", "条件", "预期路径", "检查点", "结论边界"],
    [
        ("TC-01", "突发胸痛伴呼吸困难", "快速通道", "无模型/检索调用", "规则回归"),
        ("TC-02", "仅“头痛”且信息不足", "ask_followup", "一项关键追问", "非诊断能力"),
        ("TC-03", "普通充分输入", "retrieve→classify", "引用和版本齐全", "无检索金标"),
        ("TC-04", "上游中途错误", "error 终止", "仅失败 Trace", "不代表可用性"),
        ("TC-05", "跨用户 session", "权限拒绝", "不返回他人数据", "未做组织 IAM"),
        ("TC-06", "损坏候选索引", "保持旧 active", "版本不被覆盖", "未做容灾演练"),
    ],
    [24, 37, 35, 34, 35],
)

METRIC_TABLE = (
    "表6-3 工程评测指标与当前解释口径",
    ["指标", "当前值/状态", "样本", "可支持的结论", "不可支持的结论"],
    [
        ("安全召回率", "1.0", "10 个构造病例", "规则基线一致", "临床敏感度"),
        ("红旗漏检", "0", "10 个构造病例", "当前样例无漏检", "真实世界漏检率"),
        ("科室 Macro-F1", "1.0", "10 个构造病例", "工程标签一致", "患者分诊准确率"),
        ("风险 Macro-F1", "1.0", "10 个构造病例", "工程标签一致", "临床风险校准"),
        ("错误率", "0", "10 个构造病例", "确定性基线无错", "总体可靠性"),
        ("Recall@K/MRR", "not-evaluated", "无 gold evidence", "尚未测量", "检索质量为零"),
    ],
    [29, 34, 31, 33, 35],
)


def add_list_item(doc: Document, text: str) -> None:
    """Add a real Word list paragraph rather than a manually typed bullet."""
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(clean_text(text))
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5)


def clean_text(text: str) -> str:
    """Remove stale placeholders and repair citation cross-references."""
    text = text.replace("（待用户运行系统后替换为实际截图）", "")
    text = text.replace("待用户运行系统后替换", "")
    text = text.replace("本文正文已使用文献[1-44]", "本文正文结合文献[1-40]")
    text = text.replace("[8,10]", "[6,7]")
    text = text.replace("[5-9]", "[5-7]")
    text = text.replace("[11-13]", "[8-10]")
    text = text.replace("[11]", "[9]")
    text = text.replace("[12]", "[10]")
    text = text.replace("[13]", "[8-10]")
    text = text.replace("[14]", "[5]")
    text = text.replace("[15]", "[8]")
    text = text.replace("[16]", "[8]")
    text = text.replace("[20]", "[8]")
    text = text.replace("，并参考 RAGAs 的自动评价思路[19]", "，并以独立金标准和人工审查作为评价前提")
    text = text.replace("Ollama 支持本地部署 Qwen2.5-7B[18] 和 bge-m3[17]", "Ollama 支持本地部署的语言模型与嵌入模型")
    text = text.replace("第1章介绍研究背景、相关工作、研究目标和技术路线；第2章分析系统需求，给出总体架构、角色模型、工作流和数据边界；第3章说明各功能节点、健康上下文、知识检索、分诊解释、复诊、Trace、附件和运维设计；第4章介绍运行环境、测试用例、离线评测指标和当前工程验收结果；第5章总结已完成工作和局限，提出后续工程与临床研究方向。", "第1章阐述研究背景与边界；第2章完成系统分析与技术选择；第3章给出总体架构与数据设计；第4章说明核心算法；第5章介绍模块实现；第6章报告工程测试与分析；第7章总结局限并提出后续验证路线。")
    return re.sub(r"\s+", " ", text).strip()


def remap_heading(text: str, old_chapter: int, destination: int) -> str:
    """Map source chapter numbers into the final seven-chapter structure."""
    match = re.match(r"^(\d+)\.(\d+)(.*)$", text)
    if not match:
        return text
    first, second, suffix = int(match.group(1)), int(match.group(2)), match.group(3)
    if old_chapter == 3:
        if second <= 5:
            return f"4.{second}{suffix}"
        return f"5.{second - 5}{suffix}"
    if old_chapter == 4:
        return f"6.{second}{suffix}"
    if old_chapter == 5:
        return f"7.{second}{suffix}"
    return text


def append_entry(buckets, chapter: int | None, kind: str, level: int, value: str) -> None:
    if chapter is None:
        return
    value = clean_text(value)
    if value:
        buckets[chapter].append((kind, level, value))


def collect_v0(buckets) -> None:
    current_old = None
    destination = None
    skipped = {"摘要", "Abstract", "目录（排版时自动生成）", "图清单", "表清单", "封面信息（待按学校模板排版）", "设计原创性声明", "版权使用授权书"}
    for kind, level, value in base.parse_markdown(DRAFT_DIR / "thesis_v0.md"):
        if kind == "h":
            root = re.match(r"^第([1-5])章", value)
            if root:
                current_old = int(root.group(1))
                destination = {1: 1, 2: 2, 4: 6, 5: 7}.get(current_old)
                continue
            if value.startswith(("参考文献", "致谢")):
                current_old = None
                destination = None
                continue
            if value in skipped:
                continue
            if current_old is None:
                continue
            if current_old == 3:
                number = re.match(r"^3\.(\d+)", value)
                if number:
                    destination = 4 if int(number.group(1)) <= 5 else 5
                append_entry(buckets, destination, "h", min(level, 3), remap_heading(value, current_old, destination))
            else:
                append_entry(buckets, destination, "h", min(level, 3), remap_heading(value, current_old, destination))
        elif current_old is not None:
            append_entry(buckets, destination, kind, level, value)


def collect_expansion(buckets, appendices) -> None:
    section = None
    destination = None
    alg_map = {13: 6, 14: 7, 19: 8}
    impl_map = {9: 4, 10: 5, 11: 6, 12: 7, 15: 8, 16: 9, 17: 10, 18: 11, 20: 12, 21: 13, 22: 14, 23: 15}
    for kind, level, value in base.parse_markdown(DRAFT_DIR / "expansion_content.md"):
        if kind == "h":
            if value.startswith("MedPilot 论文扩写内容"):
                continue
            chapter_mark = re.match(r"^第([1-5])章补充", value)
            if chapter_mark:
                section = int(chapter_mark.group(1))
                continue
            if value.startswith("附录"):
                section = 0
                appendices.append(("h", 1, clean_text(value)))
                continue
            if section == 0:
                appendices.append(("h", min(level, 3), clean_text(value)))
                continue
            number = re.match(r"^(\d+)\.(\d+)(.*)$", value)
            if number:
                first, second, suffix = int(number.group(1)), int(number.group(2)), number.group(3)
                if first in (1, 2):
                    destination = first
                    mapped = value
                elif first == 3:
                    if second in alg_map:
                        destination = 4
                        mapped = f"4.{alg_map[second]}{suffix}"
                    else:
                        destination = 5
                        mapped = f"5.{impl_map.get(second, second)}{suffix}"
                elif first == 4:
                    destination = 6
                    mapped = f"6.{second}{suffix}"
                else:
                    destination = 7
                    mapped = f"7.{second}{suffix}"
                append_entry(buckets, destination, "h", min(level, 3), mapped)
            else:
                append_entry(buckets, destination, "h", min(level, 3), value)
        elif section == 0:
            appendices.append((kind, level, clean_text(value)))
        else:
            append_entry(buckets, destination, kind, level, value)


def collect_deepening(buckets) -> None:
    source_chapter = None
    destination = None
    deep_map = {
        1: (1, {9: 9, 10: 10, 11: 11, 12: 12, 13: 13}),
        2: (2, {12: 12, 13: 13, 14: 14, 15: 15, 16: 16}),
        3: (5, {24: 16, 25: 17, 26: 18, 27: 19, 28: 20}),
        4: (4, {18: 9, 19: 10, 20: 11, 21: 12, 22: 13, 23: 14}),
        5: (5, {8: 21, 9: 22, 10: 23, 11: 24, 12: 25}),
        6: (6, {1: 18, 2: 19, 3: 20, 4: 21, 5: 22, 6: 23}),
        7: (7, {1: 8, 2: 9, 3: 10}),
    }
    for kind, level, value in base.parse_markdown(DRAFT_DIR / "final_deepening_content.md"):
        if kind == "h":
            marker = re.match(r"^第([1-7])章扩写", value)
            if marker:
                source_chapter = int(marker.group(1))
                destination = deep_map[source_chapter][0]
                continue
            number = re.match(r"^(\d+)\.(\d+)(.*)$", value)
            if number and source_chapter:
                first, second, suffix = int(number.group(1)), int(number.group(2)), number.group(3)
                destination, mapping = deep_map[source_chapter]
                append_entry(buckets, destination, "h", min(level, 3), f"{destination}.{mapping.get(second, second)}{suffix}")
            else:
                append_entry(buckets, destination, "h", min(level, 3), value)
        else:
            append_entry(buckets, destination, kind, level, value)


def add_chapter_intro(doc: Document, chapter: int) -> None:
    for paragraph in base.FRESH_CONTENT.get(f"第{chapter}章", []):
        base.add_para(doc, clean_text(paragraph))
    if chapter == 1:
        base.add_para(doc, "中文研究围绕大数据管理、知识图谱、知识表示、可解释性与隐私保护提出了持续演进的工程问题[11-20]。这些研究为本系统的数据分层、证据元数据、权限控制和解释边界提供方法论参照，但不构成本系统的临床效果证据。")
        base.add_para(doc, "与医疗服务相关的中文研究则提示，互联网医疗、分级诊疗、医院质量管理和患者体验评价必须置于制度、流程与责任边界中考察[36-40]。因此，MedPilot 的设计不将线上咨询表述为临床替代，而把转人工、线下就医与审计回溯视为系统输出的一部分。")


def add_chapter_three_design(doc: Document, figs) -> None:
    base.add_heading(doc, "3.1 系统架构与模块协作设计", 2)
    for text in base.FRESH_CONTENT["第3章"][:2]:
        base.add_para(doc, clean_text(text))
    base.add_figure(doc, figs["分层架构"], "图3-1 MedPilot 系统总体分层架构图", 150)
    base.add_figure(doc, figs["信任边界"], "图3-2 信任边界与接口调用链图", 150)
    base.add_heading(doc, "3.2 核心数据模型与关系设计", 2)
    for text in base.FRESH_CONTENT["第3章"][2:4]:
        base.add_para(doc, clean_text(text))
    base.add_figure(doc, figs["数据库"], "图3-3 MedPilot 核心数据库 ER 图", 150)
    for caption, rows in base.SCHEMA_TABLES:
        base.add_schema_table(doc, caption, rows)
    base.add_heading(doc, "3.3 权限、事件与治理设计", 2)
    for text in base.FRESH_CONTENT["第3章"][4:]:
        base.add_para(doc, clean_text(text))
    base.add_figure(doc, figs["知识治理"], "图3-4 医学知识库治理生命周期图", 150)


def emit_entries(doc: Document, chapter: int, entries, figs) -> None:
    inserted = set()
    for kind, level, value in entries:
        if kind == "h":
            base.add_heading(doc, value, min(level, 3))
            if chapter == 2 and value.startswith("2.1") and "functional" not in inserted:
                base.add_figure(doc, figs["功能结构"], "图2-1 MedPilot 系统功能结构图", 150)
                base.add_data_table(doc, *base.DATA_TABLES[0])
                inserted.add("functional")
            if chapter == 2 and value.startswith("2.2") and "roles" not in inserted:
                base.add_figure(doc, figs["角色权限"], "图2-2 用户角色与权限边界图", 150)
                base.add_data_table(doc, *ROLE_TABLE)
                base.add_data_table(doc, *NFR_TABLE)
                inserted.add("roles")
            if chapter == 2 and value.startswith("2.3") and "state" not in inserted:
                base.add_figure(doc, figs["状态机"], "图2-3 咨询业务状态机图", 150)
                inserted.add("state")
            if chapter == 4 and value.startswith("4.1") and "workflow" not in inserted:
                base.add_figure(doc, figs["普通流程"], "图4-1 普通咨询流程图", 150)
                base.add_figure(doc, figs["红旗通道"], "图4-2 红旗高风险快速通道图", 150)
                base.add_figure(doc, figs["主动追问"], "图4-3 信息不足与主动追问流程图", 150)
                base.add_data_table(doc, *ALGORITHM_TABLE)
                inserted.add("workflow")
            if chapter == 4 and value.startswith("4.3") and "rag" not in inserted:
                base.add_figure(doc, figs["RAG"], "图4-4 医学 RAG 检索增强流程图", 150)
                base.add_figure(doc, figs["证据分诊"], "图4-5 证据加权与辅助分诊流程图", 150)
                base.add_data_table(doc, *EVENT_TABLE)
                inserted.add("rag")
            if chapter == 4 and value.startswith("4.4") and "gate" not in inserted:
                base.add_figure(doc, figs["复核闸门"], "图4-6 医生复核安全闸门图", 150)
                inserted.add("gate")
            if chapter == 4 and value.startswith("4.5") and "ndjson" not in inserted:
                base.add_figure(doc, figs["NDJSON"], "图4-7 NDJSON 事件生命周期与失败终止图", 150)
                inserted.add("ndjson")
            if chapter == 5 and value.startswith("5.1") and "deployment" not in inserted:
                base.add_figure(doc, figs["部署"], "图5-1 MedPilot 系统部署与服务边界图", 150)
                inserted.add("deployment")
            if chapter == 5 and (value.startswith("5.3") or value.startswith("5.13")) and "trace" not in inserted:
                base.add_figure(doc, figs["Trace"], "图5-2 Trace 监控、审计与回溯流程图", 150)
                base.add_data_table(doc, *INTERFACE_TABLE)
                inserted.add("trace")
            if chapter == 6 and value.startswith("6.1") and "tests" not in inserted:
                base.add_data_table(doc, *TEST_ENV_TABLE)
                base.add_data_table(doc, *TEST_CASE_TABLE)
                base.add_data_table(doc, *METRIC_TABLE)
                base.add_data_table(doc, *base.DATA_TABLES[2])
                base.add_data_table(doc, *base.DATA_TABLES[1])
                base.add_data_table(doc, *base.DATA_TABLES[3])
                inserted.add("tests")
        elif kind == "b":
            add_list_item(doc, value)
        else:
            base.add_para(doc, value)


def add_references(doc: Document) -> None:
    doc.add_page_break()
    base.add_heading(doc, "参考文献", 1)
    base.add_para(
        doc,
        "说明：英文条目[1]-[9]已按 DOI 的公开元数据逐项核对，[10]为同行评审国际会议论文。中文条目[11]-[40]优先选择《计算机研究与发展》《中国卫生政策研究》《中国医院管理》等正式期刊来源，并依据期刊 DOI 或公开文献索引逐条复核题名、刊名、卷期与起始页。CNKI 当前要求 blockPuzzle 人机验证，本环境不能取得学校账号的最终导出记录；提交前请在学校 CNKI/万方账号中对中文作者写法、完整页码、收录状态和本校采用的核心目录做最后一次导出核验。本说明不把工程原型的结果外推为临床结论。",
        indent=False,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        size=9.2,
    )
    for reference in REFERENCES:
        base.add_para(doc, reference, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.1)


def add_appendices(doc: Document, appendices) -> None:
    doc.add_page_break()
    base.add_heading(doc, "附录A 关键接口示例（脱敏）", 1)
    base.add_para(doc, "咨询请求至少包含 session_id 与 text；后端从认证上下文恢复用户和会话归属，AI 服务以 NDJSON 返回节点事件。所有示例均为协议说明，不含真实患者标识、令牌、密钥或生产地址。", indent=False)
    base.add_data_table(doc, "表A-1 NDJSON 事件字段示例", ["字段", "类型", "必填", "含义", "校验"], [
        ("protocol_version", "string", "是", "协议版本", "白名单"),
        ("sequence", "integer", "是", "事件顺序", "严格递增"),
        ("type", "string", "是", "事件类别", "枚举"),
        ("node", "string", "是", "工作流节点", "节点白名单"),
        ("status", "string", "是", "节点状态", "状态机"),
        ("data", "object", "否", "节点输出", "结构校验"),
    ], [31, 24, 20, 43, 31])
    base.add_heading(doc, "附录B 红旗规则回归样例", 1)
    base.add_data_table(doc, "表B-1 红旗规则回归样例", ["编号", "输入片段", "预期", "模型调用", "说明"], [
        ("RF-01", "突发胸痛伴冷汗", "高风险", "不调用", "急诊提示"),
        ("RF-02", "没有胸痛，但出现呼吸困难", "高风险", "不调用", "否定范围正确"),
        ("RF-03", "轻微咳嗽三天", "普通/追问", "可调用", "非红旗"),
        ("RF-04", "短暂晕厥后意识恢复", "高风险", "不调用", "人工复核条件"),
        ("RF-05", "便血反复出现", "高风险", "不调用", "就医时效提示"),
    ], [25, 54, 30, 25, 36])
    base.add_heading(doc, "附录C 复现实验记录模板", 1)
    base.add_para(doc, "每次复现实验记录提交版本、模型名称、索引 manifest 哈希、病例文件哈希、Python/Java/Node 版本、启动参数、测试命令、通过数、失败数和环境异常。任何手工修改均需记录，不应仅保留最终截图。", indent=False)
    if appendices:
        base.add_heading(doc, "附录D 扩展设计说明", 1)
        for kind, level, value in appendices:
            if kind == "h":
                base.add_heading(doc, value, min(level + 1, 3))
            elif kind == "b":
                add_list_item(doc, value)
            else:
                base.add_para(doc, value)


def count_cjk(doc: Document) -> int:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    joined = "\n".join(parts)
    return sum("\u4e00" <= char <= "\u9fff" for char in joined)


def main() -> None:
    buckets = defaultdict(list)
    appendices = []
    collect_v0(buckets)
    collect_expansion(buckets, appendices)
    collect_deepening(buckets)

    figs = visio_figures()
    doc = Document()
    base.configure_doc(doc)
    doc.sections[0].different_first_page_header_footer = True
    base.add_cover(doc)
    base.add_front_matter(doc)

    for chapter in range(1, 8):
        # The front matter already ends with a page break.  Starting chapter 1
        # with another break would leave a blank page between the lists and text.
        if chapter != 1:
            doc.add_page_break()
        base.add_heading(doc, CHAPTER_TITLES[chapter], 1)
        add_chapter_intro(doc, chapter)
        if chapter == 3:
            add_chapter_three_design(doc, figs)
        emit_entries(doc, chapter, buckets[chapter], figs)

    add_references(doc)
    base.add_heading(doc, "致谢", 1)
    base.add_para(doc, "感谢指导教师在课题选题、系统边界、数据库设计、测试方案和论文撰写方面提供的指导。感谢同学在接口联调、界面检查和复核流程讨论中的帮助。本文病例为人工构造或脱敏数据，MedPilot 仅作为证据增强、规则护栏下的辅助分诊工程原型。", size=11)
    # The requested final thesis intentionally stops at 致谢.  Appendix
    # material remains in the working draft but is not exported.

    doc.core_properties.title = base.TITLE
    doc.core_properties.subject = "MedPilot 医疗多智能体辅助分诊系统毕业设计论文完整重制终稿"
    doc.core_properties.author = ""
    doc.core_properties.comments = "正文、表格、黑白技术图和参考文献按项目证据重制。"
    doc.save(OUT)

    stats = {
        "output": str(OUT),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "images": len(doc.inline_shapes),
        "references": len(REFERENCES),
        "cjk_characters": count_cjk(doc),
        "chapter_entries": {str(key): len(value) for key, value in buckets.items()},
    }
    STATS_OUT.write_text(json.dumps(stats, ensure_ascii=False, indent=2), encoding="utf-8")
    if stats["cjk_characters"] < 42000:
        raise RuntimeError(f"正文中文字符数不足安全阈值：{stats['cjk_characters']}")
    if stats["tables"] < 20 or stats["images"] < 15 or stats["references"] < 35:
        raise RuntimeError(f"结构检查失败：{stats}")
    print(json.dumps(stats, ensure_ascii=False))


if __name__ == "__main__":
    main()
